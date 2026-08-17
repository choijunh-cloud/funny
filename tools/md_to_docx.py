#!/usr/bin/env python3
"""Convert the daily note markdown into a Word document.

Supports the subset of markdown used in notes/: ATX headings, unordered and
ordered lists (one nesting level), pipe tables, blockquotes, fenced code
blocks, horizontal rules and inline **bold** / *italic* / `code`.

Usage:
    python3 tools/md_to_docx.py notes/2026-08-18-non-samjeonnix.md "출력 파일.docx"
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

KOREAN_FONT = "맑은 고딕"
INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`.+?`)")


def set_korean_font(document: Document) -> None:
    for style_name in ("Normal", "List Bullet", "List Number", "Quote"):
        try:
            style = document.styles[style_name]
        except KeyError:
            continue
        style.font.name = KOREAN_FONT
        style.font.size = Pt(10.5)
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
            KOREAN_FONT,
        )


def add_inline(paragraph, text: str) -> None:
    for chunk in INLINE_PATTERN.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            paragraph.add_run(chunk[2:-2]).bold = True
        elif chunk.startswith("*") and chunk.endswith("*"):
            paragraph.add_run(chunk[1:-1]).italic = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(chunk)


def is_table_separator(line: str) -> bool:
    return bool(re.fullmatch(r"\|[\s:\-|]+\|", line.strip()))


def split_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            if c >= len(rows[0]):
                continue
            cell = table.cell(r, c)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            add_inline(paragraph, cell_text)
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if r == 0:
                    run.bold = True
    document.add_paragraph()


def add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(18)
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)


def convert(md_path: Path, docx_path: Path) -> None:
    document = Document()
    set_korean_font(document)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        if stripped.startswith("```"):
            block: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                block.append(lines[index])
                index += 1
            index += 1
            add_code_block(document, block)
            continue

        if stripped.startswith("|") and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            header = split_row(stripped)
            rows = [header]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(split_row(lines[index]))
                index += 1
            add_table(document, rows)
            continue

        if re.fullmatch(r"-{3,}", stripped):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run("· · ·")
            run.font.size = Pt(9)
            index += 1
            continue

        heading = re.match(r"(#{1,4})\s+(.*)", stripped)
        if heading:
            level = len(heading.group(1))
            document.add_heading(heading.group(2), level=min(level, 4))
            index += 1
            continue

        if stripped.startswith("> "):
            paragraph = document.add_paragraph(style="Quote")
            add_inline(paragraph, stripped[2:])
            index += 1
            continue

        bullet = re.match(r"(\s*)[-*]\s+(?:\[[ x]\]\s+)?(.*)", line)
        if bullet:
            nested = len(bullet.group(1)) >= 2
            paragraph = document.add_paragraph(style="List Bullet")
            if nested:
                paragraph.paragraph_format.left_indent = Pt(48)
            add_inline(paragraph, bullet.group(2))
            index += 1
            continue

        numbered = re.match(r"(\s*)\d+\.\s+(.*)", line)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            if len(numbered.group(1)) >= 2:
                paragraph.paragraph_format.left_indent = Pt(48)
            add_inline(paragraph, numbered.group(2))
            index += 1
            continue

        paragraph = document.add_paragraph()
        add_inline(paragraph, stripped)
        index += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_path)


def main() -> int:
    if len(sys.argv) != 3:
        print(__doc__)
        return 1
    convert(Path(sys.argv[1]), Path(sys.argv[2]))
    print(f"saved: {sys.argv[2]}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
