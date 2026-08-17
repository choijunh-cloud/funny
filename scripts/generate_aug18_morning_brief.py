#!/usr/bin/env python3
"""8월 18일 오전 Quick 코멘트 → 강의노트(.docx) + 시각화 슬라이드(.html)."""

from __future__ import annotations

import io
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import font_manager
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Mm, Pt, RGBColor

from aug18_morning_data import (
    FX_KRW,
    HYNIX,
    HYNIX_TP,
    KOREA_2D,
    MICRON,
    MONITOR,
    SAMSUNG,
    SAMSUNG_TP,
    SANDISK,
    SEAGATE,
    US,
    WD,
    adr_to_local_man,
    assert_all,
    qoq_fy_eps,
)

OUT_DIR = Path("/workspace/lectures")
DOCX_PATH = OUT_DIR / "8월 18일 오전 시장 브리핑 (유가·금리·메모리).docx"
HTML_PATH = OUT_DIR / "8월 18일 오전 시장 브리핑.html"

KR_FONT = "맑은 고딕"
CHART_FONT = "WenQuanYi Micro Hei"
FONT_PATH = "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc"

NAVY = RGBColor(0x0F, 0x20, 0x43)
NAVY2 = RGBColor(0x1E, 0x40, 0x7C)
GOLD = RGBColor(0xB8, 0x94, 0x3A)
GRAY = RGBColor(0x4B, 0x55, 0x63)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x16, 0x65, 0x34)
RED = RGBColor(0x99, 0x1B, 0x1B)
AMBER = RGBColor(0x7A, 0x5C, 0x12)

NAVY_HEX = "0F2043"
NAVY2_HEX = "1E407C"
GOLD_HEX = "B8943A"
LIGHT_HEX = "EEF2F8"
GREEN_HEX = "E8F5E9"
RED_HEX = "FDECEA"
AMBER_HEX = "FFF8E7"
BLUE_HEX = "E8F1FB"
ROW_HEX = "F7F9FC"
WHITE_HEX = "FFFFFF"


def _setup_chart_font():
    if Path(FONT_PATH).exists():
        font_manager.fontManager.addfont(FONT_PATH)
        plt.rcParams["font.family"] = CHART_FONT
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["figure.facecolor"] = "white"
    plt.rcParams["savefig.facecolor"] = "white"
    plt.rcParams["axes.facecolor"] = "white"


def set_run_font(run, size=11, bold=False, color=DARK, italic=False, font=KR_FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), KR_FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill}"/>'))


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(
        parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'<w:top w:w="{top}" w:type="dxa"/>'
            f'<w:left w:w="{left}" w:type="dxa"/>'
            f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
            f'<w:right w:w="{right}" w:type="dxa"/>'
            f"</w:tcMar>"
        )
    )


def set_table_borders(table, color="D0D7E2", sz="4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tbl_pr.append(
        parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f"</w:tblBorders>"
        )
    )


def set_left_accent(cell, color=NAVY_HEX, sz="24"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(
        parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="nil"/>'
            f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="nil"/>'
            f'<w:right w:val="nil"/>'
            f"</w:tcBorders>"
        )
    )


def prevent_row_split(row):
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()
    tr_pr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))


def cell_text(cell, text, size=10, bold=False, color=DARK, align="left", font=KR_FONT):
    cell.text = ""
    align_enum = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    lines = str(text).split("\n")
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = align_enum
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1 if i < len(lines) - 1 else 0)
        p.paragraph_format.line_spacing = 1.12
        run = p.add_run(line)
        set_run_font(run, size=size, bold=bold, color=color, font=font)
    set_cell_margins(cell)


def add_runs(paragraph, parts, size=11, color=DARK):
    for part in parts:
        if isinstance(part, str):
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)
        else:
            text, bold, *rest = part
            c = rest[0] if rest else color
            run = paragraph.add_run(text)
            set_run_font(run, size=size, bold=bold, color=c)


class Notes:
    def __init__(self):
        self.doc = Document()
        self._setup()

    def _setup(self):
        sec = self.doc.sections[0]
        sec.page_width = Mm(210)
        sec.page_height = Mm(297)
        sec.left_margin = Mm(16)
        sec.right_margin = Mm(16)
        sec.top_margin = Mm(16)
        sec.bottom_margin = Mm(16)
        sec.header_distance = Mm(8)
        sec.footer_distance = Mm(8)

        normal = self.doc.styles["Normal"]
        normal.font.name = KR_FONT
        normal.font.size = Pt(11)
        normal.font.color.rgb = DARK
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), KR_FONT)
        pf = normal.paragraph_format
        pf.space_after = Pt(6)
        pf.space_before = Pt(0)
        pf.line_spacing = 1.18

        header = sec.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = hp.add_run("8/18 오전  ·  유가·장기금리 vs 메모리  ·  강의노트")
        set_run_font(r, size=8.5, color=GRAY)

        footer = sec.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run("영상 녹화용 정리  ·  숫자·시나리오는 공개 코멘트 기준  ·  ")
        set_run_font(r, size=8, color=GRAY)
        fp._p.append(
            parse_xml(
                f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE ">'
                f'<w:r><w:rPr><w:sz w:val="16"/><w:color w:val="4B5563"/>'
                f'<w:rFonts w:ascii="{KR_FONT}" w:hAnsi="{KR_FONT}" w:eastAsia="{KR_FONT}"/></w:rPr>'
                f"<w:t></w:t></w:r></w:fldSimple>"
            )
        )

        core = self.doc.core_properties
        core.title = "8월 18일 오전 시장 브리핑 (유가·금리·메모리)"
        core.author = "준혁"
        core.subject = "8/17 미국장 유가·장기금리, 메모리 밸류, 8/18 한국장 기준점"

    def p(self, text, size=11, bold=False, color=DARK, space_after=6, space_before=0, align="left"):
        para = self.doc.add_paragraph()
        para.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }[align]
        para.paragraph_format.space_after = Pt(space_after)
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.line_spacing = 1.18
        run = para.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
        return para

    def rich(self, parts, size=11, space_after=6, space_before=0):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_after = Pt(space_after)
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.line_spacing = 1.18
        add_runs(para, parts, size=size)
        return para

    def h1(self, text, num=None):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing = 1.1
        if num:
            run = para.add_run(f"{num}  ")
            set_run_font(run, size=16, bold=True, color=GOLD)
        run = para.add_run(text)
        set_run_font(run, size=16, bold=True, color=NAVY)
        pPr = para._p.get_or_add_pPr()
        pPr.append(
            parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="12" w:space="4" w:color="{NAVY_HEX}"/>'
                f"</w:pBdr>"
            )
        )
        return para

    def h2(self, text):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(text)
        set_run_font(run, size=13, bold=True, color=NAVY2)
        return para

    def bullet(self, text, level=0, bold_lead=None, size=11):
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.55 + level * 0.45)
        para.paragraph_format.first_line_indent = Cm(-0.35)
        para.paragraph_format.space_after = Pt(2.5)
        para.paragraph_format.line_spacing = 1.15
        mark = "• " if level == 0 else "– "
        run = para.add_run(mark)
        set_run_font(run, size=size, color=NAVY2 if level == 0 else GRAY)
        if bold_lead:
            run = para.add_run(bold_lead)
            set_run_font(run, size=size, bold=True, color=DARK)
            run = para.add_run(text)
            set_run_font(run, size=size, color=DARK)
        else:
            run = para.add_run(text)
            set_run_font(run, size=size, color=DARK)
        return para

    def flow(self, items, size=11):
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing = 1.2
        for i, item in enumerate(items):
            if i:
                run = para.add_run("   →   ")
                set_run_font(run, size=size, bold=True, color=GOLD)
            run = para.add_run(item)
            set_run_font(run, size=size, bold=True, color=NAVY)
        return para

    def callout(self, title, body, kind="key"):
        palette = {
            "key": (NAVY_HEX, LIGHT_HEX, NAVY),
            "bull": ("166534", GREEN_HEX, GREEN),
            "bear": ("991B1B", RED_HEX, RED),
            "note": (GOLD_HEX, AMBER_HEX, AMBER),
            "blue": (NAVY2_HEX, BLUE_HEX, NAVY2),
        }
        accent, fill, title_color = palette[kind]
        table = self.doc.add_table(rows=1, cols=1)
        table.autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        shade_cell(cell, fill)
        set_left_accent(cell, accent, sz="28")
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        r = p1.add_run(title)
        set_run_font(r, size=10, bold=True, color=title_color)
        if isinstance(body, str):
            body = [body]
        for line in body:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(line)
            set_run_font(r, size=10.5, color=DARK)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def table(self, headers, rows, col_widths=None, first_col_bold=True):
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        set_table_borders(table, color="D5DCE6", sz="4")
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            shade_cell(cell, NAVY_HEX)
            cell_text(cell, h, size=9.5, bold=True, color=WHITE, align="center")
        prevent_row_split(table.rows[0])
        for r_i, row in enumerate(rows):
            for c_i, val in enumerate(row):
                cell = table.rows[r_i + 1].cells[c_i]
                shade_cell(cell, ROW_HEX if r_i % 2 == 1 else WHITE_HEX)
                align = "left" if c_i == 0 else "center"
                cell_text(cell, str(val), size=9.5, bold=first_col_bold and c_i == 0, color=DARK, align=align)
            prevent_row_split(table.rows[r_i + 1])
        if col_widths:
            for row in table.rows:
                for i, w in enumerate(col_widths):
                    row.cells[i].width = Cm(w)
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(8)
        return table

    def picture(self, png_bytes: bytes, width_cm=17.6):
        stream = io.BytesIO(png_bytes)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run()
        run.add_picture(stream, width=Cm(width_cm))
        return p

    def save(self, path: Path):
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))


def fig_to_png(fig) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def chart_per_memory() -> bytes:
    labels = [
        "마이크론 FWD 12M",
        "마이크론 CY27",
        "Sandisk FY27",
        "하이닉스 ADR 26년",
        "하이닉스 ADR 27년",
        "하이닉스 본주 26년",
        "하이닉스 본주 27년",
        "삼성전자 26년",
        "삼성전자 27년",
    ]
    values = [
        MICRON["fwd12_per"],
        MICRON["cy27_per"],
        SANDISK["fy27_per"],
        HYNIX["adr_per_26"],
        HYNIX["adr_per_27"],
        HYNIX["per_26"],
        HYNIX["per_27"],
        SAMSUNG["per_26"],
        SAMSUNG["per_27"],
    ]
    colors = [
        "#1E407C",
        "#1E407C",
        "#3F6F9C",
        "#B8943A",
        "#B8943A",
        "#0F2043",
        "#0F2043",
        "#4B5563",
        "#4B5563",
    ]
    fig, ax = plt.subplots(figsize=(10.4, 5.2))
    y = list(range(len(labels)))[::-1]
    ax.barh(y, values, color=colors, height=0.62, zorder=3)
    for yi, v in zip(y, values):
        ax.text(v + 0.12, yi, f"{v:.2f}배", va="center", ha="left", fontsize=9, color="#1A1A1A")
    ax.set_yticks(y)
    ax.set_yticklabels(labels, fontsize=10)
    ax.set_xlim(0, 10.2)
    ax.set_xlabel("PER (배)", fontsize=10)
    ax.set_title("Memory PER 비교  ·  8/17 종가 · 공개 코멘트 기준", fontsize=13, color="#0F2043", pad=10, loc="left")
    ax.axvline(MICRON["cy27_per"], color="#1E407C", ls="--", lw=0.8, alpha=0.55, zorder=2)
    ax.text(MICRON["cy27_per"] + 0.08, -0.85, "마이크론 CY27 6.75배", fontsize=8, color="#1E407C")
    ax.grid(axis="x", color="#D5DCE6", lw=0.6, zorder=0)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("#D0D7E2")
    ax.spines["bottom"].set_color("#D0D7E2")
    fig.tight_layout()
    return fig_to_png(fig)


def chart_targets() -> bytes:
    fig, axes = plt.subplots(2, 1, figsize=(10.4, 5.6), sharex=False)

    def draw(ax, title, current, marks, xmax):
        ax.barh([0], [xmax], color="#EEF2F8", height=0.38, zorder=1)
        ax.barh([0], [current], color="#0F2043", height=0.18, zorder=3)
        palette = ["#B8943A", "#1E407C", "#166534", "#7A5C12"]
        for i, (label, lo, hi) in enumerate(marks):
            ax.plot([lo, hi], [0.42 + i * 0.28, 0.42 + i * 0.28], color=palette[i % 4], lw=6, solid_capstyle="round", zorder=3)
            ax.text(hi + xmax * 0.012, 0.42 + i * 0.28, f"{label}  {lo:g}~{hi:g}" if lo != hi else f"{label}  {lo:g}", va="center", fontsize=9, color="#1A1A1A")
        ax.axvline(current, color="#0F2043", lw=1, ls=":", zorder=2)
        ax.text(current, -0.42, f"현재 {current:g}", ha="center", fontsize=8.5, color="#0F2043")
        ax.set_xlim(0, xmax)
        ax.set_ylim(-0.7, 0.42 + len(marks) * 0.28 + 0.15)
        ax.set_yticks([])
        ax.set_title(title, fontsize=12, color="#0F2043", loc="left", pad=6)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.spines["left"].set_visible(False)
        ax.spines["bottom"].set_color("#D0D7E2")
        ax.grid(axis="x", color="#E5EAF1", lw=0.6, zorder=0)

    draw(
        axes[0],
        "SK하이닉스 본주 (만원)",
        HYNIX["local_man"],
        [
            ("최근 프리미엄 30~35%", HYNIX_TP["recent_premium"][0], HYNIX_TP["recent_premium"][1]),
            ("정상 프리미엄 +20%", HYNIX_TP["normal_premium"], HYNIX_TP["normal_premium"]),
            ("보수 PER 6~7배", HYNIX_TP["per6"], HYNIX_TP["per7"]),
        ],
        260,
    )
    draw(
        axes[1],
        "삼성전자 본주 (만원)",
        SAMSUNG["local_man"],
        [
            ("같은 흐름 (최근 프리미엄)", SAMSUNG_TP["same_flow"], SAMSUNG_TP["same_flow"]),
            ("보수 PER 6~7배", SAMSUNG_TP["per6"], SAMSUNG_TP["per7"]),
        ],
        38,
    )
    fig.suptitle("목표가 시나리오 밴드  ·  현재가 대비", fontsize=13, color="#0F2043", x=0.02, ha="left")
    fig.tight_layout(rect=(0, 0, 1, 0.96))
    return fig_to_png(fig)


def chart_korea_drivers() -> bytes:
    labels = ["EWY\n(2일)", "하이닉스 ADR\n(2일)", "SOX\n(2일)", "코스피\n기준점"]
    values = [KOREA_2D["ewy"], KOREA_2D["hynix_adr"], KOREA_2D["sox"], KOREA_2D["kospi_baseline"]]
    colors = ["#1E407C", "#B8943A", "#3F6F9C", "#0F2043"]
    fig, ax = plt.subplots(figsize=(10.4, 4.2))
    bars = ax.bar(labels, values, color=colors, width=0.58, zorder=3)
    for b, v in zip(bars, values):
        ax.text(b.get_x() + b.get_width() / 2, v + 0.08, f"+{v:.1f}%", ha="center", va="bottom", fontsize=11, color="#0F2043")
    ax.set_ylim(0, 4.4)
    ax.set_ylabel("%", fontsize=10)
    ax.set_title("한국시장 영향 3가지 (8/14+8/17)  →  코스피 +1% 중반 기준점", fontsize=13, color="#0F2043", loc="left", pad=10)
    ax.axhline(KOREA_2D["kospi_baseline"], color="#0F2043", ls="--", lw=0.9, alpha=0.7, zorder=2)
    ax.grid(axis="y", color="#D5DCE6", lw=0.6, zorder=0)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("#D0D7E2")
    ax.spines["bottom"].set_color("#D0D7E2")
    fig.tight_layout()
    return fig_to_png(fig)


def chart_storage_vs_memory() -> bytes:
    labels = ["마이크론\nCY27", "Sandisk\nFY27*", "하이닉스ADR\n27년", "WD\nFY27*", "Seagate\nFY27*"]
    values = [MICRON["cy27_per"], SANDISK["fy27_per"], HYNIX["adr_per_27"], WD["fy27_per"], SEAGATE["fy27_per"]]
    colors = ["#1E407C", "#3F6F9C", "#B8943A", "#991B1B", "#991B1B"]
    fig, ax = plt.subplots(figsize=(10.4, 4.4))
    bars = ax.bar(labels, values, color=colors, width=0.58, zorder=3)
    for b, v in zip(bars, values):
        ax.text(b.get_x() + b.get_width() / 2, v + 0.4, f"{v:.1f}배", ha="center", va="bottom", fontsize=10, color="#1A1A1A")
    ax.set_ylabel("PER (배)", fontsize=10)
    ax.set_title("같은 보수 성장 가정(*)에서도 스토리지는 30배대  ·  메모리는 한 자릿수", fontsize=12, color="#0F2043", loc="left", pad=10)
    ax.grid(axis="y", color="#D5DCE6", lw=0.6, zorder=0)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("#D0D7E2")
    ax.spines["bottom"].set_color("#D0D7E2")
    fig.tight_layout()
    return fig_to_png(fig)


def build_docx(charts: dict[str, bytes]) -> Path:
    n = Notes()
    adr_man = adr_to_local_man(HYNIX["adr_usd"])

    n.p("2026. 8. 18. 오전  ·  영상 녹화 / 장전·오프닝 브리핑", size=10.5, color=GRAY, align="center", space_after=4)
    n.p("Quick 코멘트  06:45–07:01", size=13, bold=True, color=GOLD, align="center", space_after=2)
    n.p("유가↑ · 장기금리↑  vs  메모리", size=22, bold=True, color=NAVY, align="center", space_after=2)
    n.p("8/17 미국 종가  →  8/18 한국장 기준점", size=14, bold=True, color=NAVY2, align="center", space_after=8)
    n.p("브렌트 $90.87  ·  미 30년물 5.31%  ·  하이닉스 ADR 171.38달러  ·  코스피 +1% 중반", size=11, color=GRAY, align="center", space_after=10)

    n.callout(
        "오늘 한 장으로 보면",
        [
            "증시 전체를 누르는 변수는 AI CAPEX가 아니라 유가와 장기금리다. 미·이란 교착 → 호르무즈 → 유가 $90 돌파 → 인플레 우려 → 30년물 2007년 이후 최고.",
            "그 안에서도 반도체·메모리는 강하다. Anthropic 연환산 $65B+, 엔비디아·OpenAI 데이터센터 최대 $105B, SOX +1.6%, 마이크론 +4%.",
            "한국장은 8/14+8/17 이틀 EWY +3.6% · 하이닉스 ADR +3.4% · SOX +1.3%. Non-삼전닉스(유가·금리)를 감안하면 코스피 +1% 중반이 기준점.",
        ],
        kind="key",
    )

    n.h2("브리핑 구성")
    n.table(
        ["파트", "원문", "오늘 확인할 것"],
        [
            ["1", "06:45", "유가·장기금리 인과사슬 — 고PER 부담의 경로"],
            ["2", "06:52", "AI·반도체 상대 강세, 차별화 표, 모니터링 4개"],
            ["3", "06:54", "마이크론·샌디스크·하이닉스 ADR PER · 본주 환산"],
            ["4", "06:55", "삼전닉스 본주 시나리오 + WD/STX 30배대"],
            ["5", "07:01", "한국장 3가지 변동률 → 코스피 +1% 중반"],
        ],
        col_widths=[2.0, 2.4, 13.2],
    )

    n.h2("오프닝 멘트 (녹화용)")
    n.p(
        "오늘은 미국장 핵심 변수와 메모리 밸류입니다. "
        "지수는 내렸지만 반도체는 올랐습니다. 유가와 장기금리가 고PER를 누르는 구간이고, "
        "실적이 실제로 나오는 메모리·HBM은 그 안에서 차별화됩니다. "
        "한국장은 이틀간 EWY와 하이닉스 ADR이 3%대였으니, 유가·금리 변수를 빼면 코스피 +1% 중반을 기준점으로 보겠습니다."
    )

    n.h1("오늘 숫자 한 장", num="0.")
    n.table(
        ["구분", "핵심 숫자", "한 줄"],
        [
            ["미국 지수", f"다우 {US['dow']}% / S&P {US['spx']}% / 나스닥 {US['ndx']}%", "유가·금리에 전체 하락"],
            ["유가·금리", f"브렌트 ${US['brent']} (+{US['brent_chg']}%)\n30년 {US['ust30']}% · 10년 {US['ust10']}%", "30년물 2007년 이후 최고"],
            ["AI/반도체", f"SOX +{US['sox_1d']}% · 마이크론 +{US['micron_chg']}%", "AI CAPEX는 아직 안 꺾임"],
            ["마이크론", f"${MICRON['px']}  FWD 8.1배 / CY27 6.75배", "메모리 밸류의 앵커"],
            ["하이닉스 ADR", f"${HYNIX['adr_usd']} → {HYNIX['adr_krw_man_stated']:.0f}만원\n26년 7.0배 · 27년 5.56배", "마이크론 대비 −18%, 본주 대비 +48%"],
            ["하이닉스 본주", f"{HYNIX['local_man']}만원  26년 {HYNIX['per_26']}배 / 27년 {HYNIX['per_27']}배", "시나리오 180~203, 보수 208~242"],
            ["삼성전자", f"{SAMSUNG['local_man']}만원  26년 {SAMSUNG['per_26']}배 / 27년 {SAMSUNG['per_27']}배", "같은 흐름 30만, 보수 28.7~33.5"],
            ["한국장", f"EWY +{KOREA_2D['ewy']}% · ADR +{KOREA_2D['hynix_adr']}% · SOX +{KOREA_2D['sox']}%", "코스피 +1% 중반 기준점"],
        ],
        col_widths=[3.2, 7.4, 7.0],
    )

    n.h1("미국 증시: 유가↑ · 장기금리↑", num="1.")
    n.callout(
        "한 줄 결론",
        [
            "하락의 핵심은 실적 쇼크가 아니라 지정학 → 유가 → 인플레 → 장기금리다.",
            "고PER 자산이 먼저 맞는다. 그래서 지수와 반도체 온도가 갈린다.",
        ],
        kind="key",
    )
    n.h2("1) 시장 상황 (8/17 종가)")
    n.table(
        ["항목", "숫자", "의미"],
        [
            ["다우 / S&P500 / 나스닥", f"{US['dow']}% / {US['spx']}% / {US['ndx']}%", "지수 전부 약세"],
            ["브렌트유", f"${US['brent']}  (+{US['brent_chg']}%)", "$90 돌파"],
            ["미 30년물", f"{US['ust30']}%", "2007년 이후 최고"],
            ["미 10년물", f"{US['ust10']}%", "장기물 동반 상승"],
        ],
        col_widths=[4.4, 5.2, 8.0],
    )

    n.h2("2) 하락의 핵심 — 인과사슬")
    n.flow(["미·이란 협상 교착", "호르무즈 리스크", f"유가 ${US['brent']}", "인플레 우려 ↑", f"30년물 {US['ust30']}%", "고PER 부담 ↑"])
    n.p("이 경로가 오늘의 시각화 포인트입니다. AI CAPEX 둔화 논쟁보다, 유가와 할인율이 지수 전체를 먼저 재가격합니다.")
    n.callout(
        "왜 고PER가 먼저 맞나",
        [
            "장기금리 ↑  =  미래 이익의 현재가치 ↓.  이익이 멀리 있는 소프트웨어·고배수 성장주가 먼저 눌린다.",
            "메모리처럼 올해·내년 EPS가 이미 크게 잡혀 있는 주식은 같은 금리에도 배수 부담이 작다.",
        ],
        kind="note",
    )

    n.h1("그런데 AI · 반도체는 강함", num="2.")
    n.callout(
        "한 줄 결론",
        ["지수 약세와 AI 실수요는 별개다. CAPEX 자체는 아직 꺾이지 않았다."],
        kind="bull",
    )
    n.bullet("Anthropic 연환산 매출 $65B+")
    n.bullet("Nvidia, OpenAI 데이터센터에 최대 $105B 지원")
    n.bullet(f"SOX +{US['sox_1d']}%  (당일)  ·  마이크론 +{US['micron_chg']}%")
    n.p("→ AI CAPEX 자체는 아직 꺾이지 않음.")

    n.h2("시장의 차별화")
    n.table(
        ["상대 강세", "상대 약세"],
        [
            ["반도체 · 메모리", "고평가 소프트웨어"],
            ["AI 실수요 직접 수혜", "AI 수익화 의문"],
            ["HBM · 메모리 가격 상승", "높은 PER 부담"],
        ],
        col_widths=[8.8, 8.8],
        first_col_bold=True,
    )

    n.h2("투자 포인트")
    n.callout(
        "지금 증시의 최대 변수",
        ["AI CAPEX 둔화보다  유가 + 장기금리가 현재 증시의 최대 변수"],
        kind="blue",
    )
    n.h2("핵심 모니터링")
    n.table(
        ["항목", "체크"],
        [[k, v] for k, v in MONITOR],
        col_widths=[3.6, 14.0],
    )
    n.callout(
        "결론",
        ["금리 상승은 증시 전체에는 부담이지만, AI 실적이 실제 발생하는 반도체·메모리는 상대적으로 차별화될 가능성"],
        kind="key",
    )

    n.h1("Memory / Storage 주가 · 밸류에이션", num="3.")
    n.p("8/17 종가. 메모리는 한 자릿수 PER, 스토리지(HDD)는 같은 보수 성장 가정에도 30배대.")
    n.picture(charts["per"])
    n.h2("1) 마이크론 · 샌디스크")
    n.table(
        ["종목", "가격", "밸류에이션"],
        [
            ["마이크론", f"{MICRON['px']}달러", f"Forward 12개월 PER {MICRON['fwd12_per']}배  /  CY27 EPS {MICRON['cy27_eps']}달러 → PER {MICRON['cy27_per']}배"],
            ["Sandisk", f"{SANDISK['px']}달러", f"FY27.1Q 가이던스 EPS {SANDISK['q1_eps']:.0f}달러, QoQ +10%/+5%/+5% → FY27 EPS {SANDISK['fy27_eps']}달러, PER {SANDISK['fy27_per']}배"],
        ],
        col_widths=[3.2, 3.6, 10.8],
    )

    n.h2("2) SK하이닉스 ADR vs 본주")
    n.table(
        ["항목", "숫자"],
        [
            ["SK하이닉스 ADR", f"{HYNIX['adr_usd']}달러  ({HYNIX['adr_krw_man_stated']:.0f}만원 = {FX_KRW}원/달러, ADR 10주=본주 1주)"],
            ["검산 환산", f"{adr_man:.1f}만원  (원문 243만원과 일치)"],
            ["ADR 26년 / 27년 PER", f"{HYNIX['adr_per_26']}배 / {HYNIX['adr_per_27']}배"],
            ["마이크론 대비", f"{HYNIX['adr_vs_micron_pct']}%  (과거 −20~−50%)"],
            ["본주 대비 프리미엄", f"{HYNIX['adr_vs_local_premium_pct']}%"],
            ["정상 프리미엄 +20% (TSMC 15% 수준)", f"본주는 현시점 {HYNIX_TP['normal_premium']}만원은 되어야"],
            ["최근 실제 30% 이상, 30~35% 적용", f"본주 {HYNIX_TP['recent_premium'][0]}~{HYNIX_TP['recent_premium'][1]}만원"],
        ],
        col_widths=[8.4, 9.2],
    )
    n.callout(
        "읽는 법",
        [
            "ADR이 본주보다 48% 비싸다. 정상 프리미엄을 20%로 보면 갭의 상당 부분은 본주가 못 따라간 것이다.",
            "최근처럼 30~35%를 ‘인정’하면 본주 180~187만원이 먼저 보이는 구간이다.",
            "마이크론 대비 −18%는 과거 할인 밴드(−20~−50%)의 상단 근처. 할인이 이미 많이 줄었다.",
        ],
        kind="note",
    )

    n.h2("3) 스토리지 3사 — 같은 보수 사다리")
    n.p("Sandisk · Western Digital · Seagate. FY27.1Q 가이던스에 QoQ +10%, +5%, +5%를 얹어 FY27 EPS를 만들고 PER을 계산. 이번 분기~1년 성장률을 한 자릿수로 가정한 수치.")
    n.picture(charts["storage"])
    n.table(
        ["종목", "가격", "1Q EPS", "FY27 EPS", "PER"],
        [
            ["Sandisk", f"{SANDISK['px']}", f"{SANDISK['q1_eps']:.0f}", f"{SANDISK['fy27_eps']}", f"{SANDISK['fy27_per']}배"],
            ["Western Digital", f"{WD['px']}", f"{WD['q1_eps']:.0f}", f"{WD['fy27_eps']}", f"{WD['fy27_per']}배"],
            ["Seagate", f"{SEAGATE['px']}", f"{SEAGATE['q1_eps']}", f"{SEAGATE['fy27_eps']}", f"{SEAGATE['fy27_per']}배"],
        ],
        col_widths=[3.8, 3.4, 3.2, 3.4, 3.8],
    )
    n.callout(
        "함의",
        [
            f"검산: 1Q×(1+1.10+1.155+1.21275). Sandisk {qoq_fy_eps(SANDISK['q1_eps']):.1f} / WD {qoq_fy_eps(WD['q1_eps']):.1f} / Seagate {qoq_fy_eps(SEAGATE['q1_eps']):.1f}.",
            "성장률 높은 전망이 다수여서, 실제 밸류에이션 배수는 더 낮을 가능성.",
            "그래도 메모리 한 자릿수 vs HDD 30배대라는 레벨 차이는 남는다.",
        ],
        kind="blue",
    )

    n.h1("삼전닉스 본주 시나리오", num="4.")
    n.callout(
        "한 줄 결론",
        [
            "최근 ADR 프리미엄을 인정하는 정도면 하이닉스 180만원대, 삼전 30만원 수준.",
            "정상 프리미엄(+20%)이면 하이닉스 203만원. 27년 성장을 아예 빼고 26년 PER 6~7배를 적용하면 하이닉스 208~242만, 삼전 28.7~33.5만.",
        ],
        kind="key",
    )
    n.h2("1) 컨센 실적 · 본주 PER")
    n.table(
        ["종목", "가격", "26년 PER", "27년 PER", "26년 OP / EPS", "27년 OP / EPS"],
        [
            ["SK하이닉스", f"{HYNIX['local_man']}만원", f"{HYNIX['per_26']}배", f"{HYNIX['per_27']}배", f"{HYNIX['op_26t']}조 / {HYNIX['eps_26k']}K", f"{HYNIX['op_27t']}조 / {HYNIX['eps_27k']}K"],
            ["삼성전자", f"{SAMSUNG['local_man']}만원", f"{SAMSUNG['per_26']}배", f"{SAMSUNG['per_27']}배", f"{SAMSUNG['op_26t']}조 / {SAMSUNG['eps_26k']}K", f"{SAMSUNG['op_27t']}조 / {SAMSUNG['eps_27k']}K"],
        ],
        col_widths=[2.8, 2.6, 2.4, 2.4, 3.8, 3.8],
    )
    n.picture(charts["targets"])

    n.h2("2) 접근 A — ADR 프리미엄")
    n.bullet("최근 프리미엄을 인정한다면 감안 정도만: 하이닉스 180만원대. 같은 흐름이면 삼전 30만원 수준.")
    n.bullet("정상적 프리미엄(+20%, TSMC는 15% 수준)이면 하이닉스 203만원.")
    n.h2("3) 접근 B — 밸류에이션 / 성장률 (보수)")
    n.p("26년 실적 대비 27년 성장이 없다고 단순화. 26년 PER 6~7배 (과거 사이클 주식 4~8배) 적용.")
    n.table(
        ["종목", "PER 6배", "PER 7배"],
        [
            [f"SK하이닉스 (26년 EPS {HYNIX['eps_26k']}K)", f"{HYNIX_TP['per6']}만원", f"{HYNIX_TP['per7']}만원"],
            [f"삼성전자 (26년 EPS {SAMSUNG['eps_26k']}K)", f"{SAMSUNG_TP['per6']}만원", f"{SAMSUNG_TP['per7']}만원"],
        ],
        col_widths=[7.6, 5.0, 5.0],
    )
    n.callout(
        "두 접근을 겹치면",
        [
            "하이닉스: 최근 인정 180~187  /  정상 203  /  보수 성장 제거 208~242.",
            "삼성전자: 같은 흐름 30  /  보수 28.7~33.5.",
            "본주 배수(4~6배)는 사이클 밴드 하단~중간. 관건은 26·27 이익이 유지되느냐이지, 배수를 더 깎을 여지는 작다.",
        ],
        kind="note",
    )

    n.h1("한국장 오프닝", num="5.")
    n.callout(
        "한 줄 결론",
        ["8/14+8/17 이틀, 한국시장에 영향이 큰 3가지: EWY +3.6%, 하이닉스 ADR +3.4%, SOX +1.3%. Non-삼전닉스(유가·금리)를 감안하면 코스피 +1% 중반이 기준점."],
        kind="key",
    )
    n.picture(charts["korea"])
    n.p("삼전닉스가 한국 시총을 끌어올리는 날입니다. 다만 유가 $90·장기금리 5.3%는 Non-삼전닉스 — 자동차·화학·내수·고PER — 를 눌러 지수 전체를 3%대까지 밀어 올리지는 못하게 합니다. 그래서 기준점을 +1% 중반에 둡니다.")
    n.bullet("위에서 열리면: 삼전닉스 갭 + 환율/ADR 프리미엄 축소 여부.")
    n.bullet("아래에서 열리면: 유가·30년물 오버나이트가 할증을 먹었는지.")
    n.bullet("오후 강의(NON-삼전닉스)와 연결: 장비·SK·Atlas는 이 기준점 위에서 수주·할인율·배치를 보면 된다.")

    n.h1("클로징", num="6.")
    n.callout(
        "오늘 가져갈 세 문장",
        [
            "1) 지수 하락의 경로는 호르무즈 → 유가 $90 → 30년물 5.31% → 고PER 부담. 모니터링은 $100 / 5.5% / 7월 FOMC 의사록 / 월마트.",
            "2) AI CAPEX는 안 꺾였다. 메모리는 마이크론 6.75~8.1배 옆에 하이닉스 ADR 5.56~7.0배. 본주는 더 싸다.",
            "3) 한국장 기준점은 코스피 +1% 중반. 하이닉스 시나리오는 180 / 203 / 208~242, 삼전은 30 / 28.7~33.5.",
        ],
        kind="key",
    )
    n.h2("클로징 멘트 (녹화용)")
    n.p(
        "금리 상승은 증시 전체에는 부담입니다. 하지만 AI 실적이 실제로 발생하는 반도체·메모리는 상대적으로 차별화될 수 있습니다. "
        "오늘은 그 차별화를 숫자로만 보겠습니다. 배수를 올려 잡는 이야기가 아니라, 이미 나와 있는 26·27 EPS와 ADR 프리미엄을 맞춰 보는 이야기입니다."
    )
    n.p("— 8월 18일 오전 브리핑. 원문 퀵코멘트(06:45~07:01)를 강의 순서로 재구성.", size=9.5, color=GRAY, align="right")

    n.save(DOCX_PATH)
    return DOCX_PATH


def _esc(s: str) -> str:
    return (
        s.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def build_html() -> Path:
    slides = []

    def slide(inner: str, title: str, kicker: str = ""):
        slides.append(
            f'<section class="slide" data-title="{_esc(title)}">'
            f'<div class="top"><span class="kicker">{_esc(kicker)}</span>'
            f'<span class="brand">8/18 오전 · Quick 코멘트</span></div>'
            f"{inner}"
            f'<div class="foot">숫자·시나리오는 공개 코멘트 기준 · 8/17 종가</div>'
            f"</section>"
        )

    slide(
        """
        <div class="hero">
          <p class="eyebrow">2026. 8. 18  오전 브리핑</p>
          <h1>유가↑ · 장기금리↑<br><em>vs 메모리</em></h1>
          <p class="sub">8/17 미국 종가 → 8/18 한국장 기준점</p>
          <div class="hero-pills">
            <div><b>브렌트</b>$90.87 <small>+2.7%</small></div>
            <div><b>30년물</b>5.31% <small>2007 이후 최고</small></div>
            <div><b>하이닉스 ADR</b>$171.38 <small>243만원</small></div>
            <div><b>코스피 기준점</b>+1% 중반</div>
          </div>
        </div>
        """,
        "표지",
        "MARKET BRIEF",
    )

    slide(
        """
        <h2>오늘 한 장</h2>
        <div class="grid-3">
          <article class="card bear">
            <h3>지수 · 매크로</h3>
            <ul>
              <li>다우 −0.51 · S&P −0.52 · 나스닥 −0.32</li>
              <li>브렌트 $90.87 (+2.7%)</li>
              <li>30년 5.31% · 10년 4.725%</li>
            </ul>
            <p class="tag">고PER 부담</p>
          </article>
          <article class="card bull">
            <h3>AI · 반도체</h3>
            <ul>
              <li>Anthropic ARR $65B+</li>
              <li>NVDA·OpenAI DC 최대 $105B</li>
              <li>SOX +1.6% · 마이크론 +4%</li>
            </ul>
            <p class="tag">CAPEX 안 꺾임</p>
          </article>
          <article class="card navy">
            <h3>한국장 기준점</h3>
            <ul>
              <li>EWY +3.6% (2일)</li>
              <li>하이닉스 ADR +3.4%</li>
              <li>SOX +1.3% → 코스피 +1% 중반</li>
            </ul>
            <p class="tag">Non-삼전닉스 차감</p>
          </article>
        </div>
        <p class="one-liner">최대 변수는 AI CAPEX가 아니라 <b>유가 + 장기금리</b>. 그 안에서 실적 나는 메모리는 차별화.</p>
        """,
        "오늘 한 장",
        "0. DASHBOARD",
    )

    slide(
        """
        <h2>하락의 핵심 — 인과사슬</h2>
        <div class="chain">
          <div class="node">미·이란<br>협상 교착</div>
          <div class="arr">→</div>
          <div class="node">호르무즈<br>리스크</div>
          <div class="arr">→</div>
          <div class="node hot">유가<br>$90.87</div>
          <div class="arr">→</div>
          <div class="node">인플레<br>우려 ↑</div>
          <div class="arr">→</div>
          <div class="node hot">30년물<br>5.31%</div>
          <div class="arr">→</div>
          <div class="node">고PER<br>증시 부담</div>
        </div>
        <div class="split">
          <div>
            <h3>지수</h3>
            <p>다우 −0.51% · S&P −0.52% · 나스닥 −0.32%</p>
            <p class="muted">장기금리 ↑ = 먼 미래 이익의 현재가치 ↓</p>
          </div>
          <div>
            <h3>그래서 갈라진다</h3>
            <p>고평가 소프트웨어 · AI 수익화 의문 종목이 먼저 맞음</p>
            <p class="muted">26·27 EPS가 이미 큰 메모리는 배수 부담이 작다</p>
          </div>
        </div>
        """,
        "인과사슬",
        "1. 시각화",
    )

    slide(
        """
        <h2>시장의 차별화</h2>
        <div class="vs">
          <div class="col up">
            <h3>상대 강세</h3>
            <ul>
              <li>반도체 · 메모리</li>
              <li>AI 실수요 직접 수혜</li>
              <li>HBM · 메모리 가격 상승</li>
            </ul>
          </div>
          <div class="col down">
            <h3>상대 약세</h3>
            <ul>
              <li>고평가 소프트웨어</li>
              <li>AI 수익화 의문</li>
              <li>높은 PER 부담</li>
            </ul>
          </div>
        </div>
        <div class="monitor">
          <div><span>유가</span>$90 → $100</div>
          <div><span>30년물</span>5.3% → 5.5%</div>
          <div><span>FOMC</span>7월 의사록</div>
          <div><span>소비</span>월마트 실적</div>
        </div>
        """,
        "차별화",
        "2. AI · 반도체",
    )

    bars = [
        ("마이크론 FWD 12M", 8.1, "navy"),
        ("마이크론 CY27", 6.75, "navy"),
        ("Sandisk FY27", 8.9, "blue"),
        ("하이닉스 ADR 26년", 7.0, "gold"),
        ("하이닉스 ADR 27년", 5.56, "gold"),
        ("하이닉스 본주 26년", 4.8, "ink"),
        ("하이닉스 본주 27년", 3.8, "ink"),
        ("삼성전자 26년", 5.7, "gray"),
        ("삼성전자 27년", 4.1, "gray"),
    ]
    bar_html = "".join(
        f'<div class="bar"><span class="bl">{lab}</span>'
        f'<div class="track"><i class="{cls}" style="width:{v / 10 * 100:.1f}%"></i></div>'
        f'<span class="bv">{v:.2f}배</span></div>'
        for lab, v, cls in bars
    )
    slide(
        f"""
        <h2>Memory PER 비교 <small>8/17 종가</small></h2>
        <div class="bars">{bar_html}</div>
        <p class="caption">하이닉스 ADR 27년 5.56배 = 마이크론 CY27 6.75배 대비 <b>−18%</b> (과거 −20~−50%). 본주는 그보다 더 낮음.</p>
        """,
        "Memory PER",
        "3. 밸류에이션",
    )

    slide(
        """
        <h2>하이닉스 ADR → 본주</h2>
        <div class="grid-2">
          <article class="card navy">
            <h3>환산</h3>
            <p class="big">$171.38 → <b>243만원</b></p>
            <p>1,417원/달러 · ADR 10주 = 본주 1주<br>검산 242.8만원</p>
            <p>본주 164.5만원 대비 <b>+48%</b> 프리미엄</p>
          </article>
          <article class="card">
            <h3>프리미엄을 얼마로 볼 것인가</h3>
            <table class="mini">
              <tr><th>가정</th><th>본주</th></tr>
              <tr><td>최근 30~35% 인정</td><td><b>180~187만원</b></td></tr>
              <tr><td>정상 +20% (TSMC ~15%)</td><td><b>203만원</b></td></tr>
              <tr><td>현재 ADR 48% 유지</td><td>243만원 (ADR)</td></tr>
            </table>
          </article>
        </div>
        <p class="caption">같은 흐름이면 삼성전자는 <b>30만원 수준</b>. 최근 프리미엄 ‘감안 정도’만 반영한 숫자.</p>
        """,
        "ADR 프리미엄",
        "3. 하이닉스",
    )

    slide(
        """
        <h2>삼전닉스 목표가 밴드</h2>
        <div class="bands">
          <div class="band-card">
            <h3>SK하이닉스 <small>현재 164.5만원 · 축 140~260</small></h3>
            <div class="axis">
              <div class="axis-track">
                <i class="mark now" style="left:20.4%"></i>
                <i class="seg gold" style="left:33.3%;width:5.8%"></i>
                <i class="seg navy" style="left:51.7%;width:2.2%"></i>
                <i class="seg green" style="left:56.7%;width:28.3%"></i>
              </div>
              <div class="axis-labels"><span>140</span><span>164.5 현재</span><span>203</span><span>242</span><span>260</span></div>
            </div>
            <ul class="legend">
              <li><i class="g"></i> 최근 프리미엄 30~35% → <b>180~187</b></li>
              <li><i class="n"></i> 정상 +20% → <b>203</b></li>
              <li><i class="a"></i> 26년 PER 6~7배 (27년 성장 0) → <b>208~242</b></li>
            </ul>
          </div>
          <div class="band-card">
            <h3>삼성전자 <small>현재 27.45만원 · 축 25~36</small></h3>
            <div class="axis">
              <div class="axis-track">
                <i class="mark now" style="left:22.3%"></i>
                <i class="seg green" style="left:33.6%;width:43.6%"></i>
                <i class="seg gold" style="left:43.6%;width:2.8%"></i>
              </div>
              <div class="axis-labels"><span>25</span><span>27.45 현재</span><span>30</span><span>33.5</span><span>36</span></div>
            </div>
            <ul class="legend">
              <li><i class="g"></i> 같은 흐름 → <b>30만원</b></li>
              <li><i class="a"></i> 26년 PER 6~7배 → <b>28.7~33.5</b></li>
            </ul>
            <p class="muted">26년 OP/EPS  하이닉스 266조 / 346K · 삼성 391조 / 47.9K</p>
          </div>
        </div>
        """,
        "목표가",
        "4. 본주 시나리오",
    )

    slide(
        """
        <h2>Storage 3사 — 한 자릿수 성장 가정</h2>
        <p class="lead">FY27.1Q 가이던스 + QoQ +10% / +5% / +5% → FY27 EPS. 성장 전망이 더 높으면 실제 PER은 더 낮아짐.</p>
        <table class="wide">
          <thead><tr><th>종목</th><th>가격</th><th>1Q EPS</th><th>FY27 EPS</th><th>PER</th></tr></thead>
          <tbody>
            <tr><td>Sandisk</td><td>1,786.85</td><td>45</td><td>201</td><td>8.9배</td></tr>
            <tr><td>Western Digital</td><td>536.01</td><td>4.0</td><td>17.9</td><td>30배</td></tr>
            <tr><td>Seagate</td><td>994.79</td><td>7.3</td><td>32.6</td><td>31배</td></tr>
          </tbody>
        </table>
        <div class="compare-note">
          <span>메모리 한 자릿수</span>
          <span class="arr">vs</span>
          <span>HDD 30배대</span>
        </div>
        """,
        "Storage",
        "4. WD · STX · SNDK",
    )

    slide(
        """
        <h2>한국장 오프닝 기준점</h2>
        <div class="k-bars">
          <div><em>EWY</em><strong>+3.6%</strong><small>2일</small></div>
          <div class="gold"><em>하이닉스 ADR</em><strong>+3.4%</strong><small>2일</small></div>
          <div><em>SOX</em><strong>+1.3%</strong><small>2일</small></div>
          <div class="ink"><em>코스피</em><strong>+1% 중반</strong><small>기준점</small></div>
        </div>
        <p class="lead">Non-삼전닉스 — 유가 · 금리 영향 등 변수 — 를 빼면 지수 전체를 3%대로 밀어 올리기 어렵다. 그래서 <b>+1% 중반</b>.</p>
        <ul class="plain">
          <li>위에서 열리면: 삼전닉스 갭 + ADR 프리미엄 축소 여부</li>
          <li>아래에서 열리면: 유가·30년물 오버나이트가 할증을 먹었는지</li>
          <li>오후 NON-삼전닉스 강의: 이 기준점 위에서 수주 · 할인율 · 배치</li>
        </ul>
        """,
        "한국장",
        "5. 8/14 + 8/17",
    )

    slide(
        """
        <h2>가져갈 세 문장</h2>
        <ol class="takeaway">
          <li>지수 하락 경로 = 호르무즈 → 유가 $90 → 30년물 5.31% → 고PER 부담. 체크는 <b>$100 / 5.5% / FOMC 의사록 / 월마트</b>.</li>
          <li>AI CAPEX는 안 꺾였다. 메모리 PER은 마이크론 6.75~8.1배 옆 하이닉스 ADR 5.56~7.0배. <b>본주는 더 싸다.</b></li>
          <li>한국장 기준점 <b>코스피 +1% 중반</b>. 하이닉스 180 / 203 / 208~242, 삼전 30 / 28.7~33.5.</li>
        </ol>
        <p class="close">금리 상승은 증시 전체에는 부담이지만, AI 실적이 실제 발생하는 반도체·메모리는 상대적으로 차별화될 가능성.</p>
        """,
        "클로징",
        "6. TAKEAWAY",
    )

    html = f"""<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1"/>
<title>8월 18일 오전 시장 브리핑 — 유가·금리·메모리</title>
<style>
  :root {{
    --navy:#0F2043; --navy2:#1E407C; --gold:#B8943A; --ink:#1A1A1A;
    --gray:#4B5563; --line:#D5DCE6; --bg:#F4F6FA; --card:#FFFFFF;
    --red:#991B1B; --green:#166534; --soft:#EEF2F8;
  }}
  * {{ box-sizing:border-box; margin:0; padding:0; }}
  html,body {{ height:100%; background:#0B1528; color:var(--ink);
    font-family:"Malgun Gothic","Apple SD Gothic Neo","Noto Sans KR",sans-serif; }}
  .deck {{ height:100%; display:flex; align-items:center; justify-content:center; }}
  .frame {{ width:min(1280px,100vw); height:min(720px,100vh); background:var(--bg);
    border-radius:12px; overflow:hidden; position:relative; box-shadow:0 20px 60px rgba(0,0,0,.35); }}
  .slide {{ display:none; position:absolute; inset:0; padding:28px 40px 36px; }}
  .slide.on {{ display:flex; flex-direction:column; }}
  .top {{ display:flex; justify-content:space-between; align-items:center; margin-bottom:10px;
    font-size:12px; letter-spacing:.04em; color:var(--gray); }}
  .kicker {{ color:var(--gold); font-weight:700; }}
  h2 {{ font-size:30px; color:var(--navy); margin:4px 0 16px; letter-spacing:-.02em; }}
  h2 small {{ font-size:14px; color:var(--gray); font-weight:600; margin-left:8px; }}
  h3 {{ font-size:16px; color:var(--navy2); margin-bottom:8px; }}
  .foot {{ position:absolute; left:40px; right:40px; bottom:12px; font-size:11px; color:#8A93A3; }}
  .hero {{ flex:1; display:flex; flex-direction:column; justify-content:center; }}
  .eyebrow {{ color:var(--gold); font-weight:700; letter-spacing:.12em; font-size:13px; }}
  .hero h1 {{ font-size:52px; line-height:1.15; color:var(--navy); margin:12px 0 8px; }}
  .hero h1 em {{ font-style:normal; color:var(--gold); }}
  .sub {{ color:var(--gray); font-size:18px; margin-bottom:28px; }}
  .hero-pills {{ display:grid; grid-template-columns:repeat(4,1fr); gap:12px; }}
  .hero-pills div {{ background:#fff; border:1px solid var(--line); border-top:3px solid var(--navy);
    padding:14px 14px 12px; border-radius:8px; font-size:20px; font-weight:700; color:var(--navy); }}
  .hero-pills b {{ display:block; font-size:11px; color:var(--gray); font-weight:600; margin-bottom:4px; }}
  .hero-pills small {{ display:block; font-size:12px; color:var(--gray); font-weight:500; margin-top:2px; }}
  .grid-3 {{ display:grid; grid-template-columns:1fr 1fr 1fr; gap:14px; }}
  .grid-2 {{ display:grid; grid-template-columns:1fr 1fr; gap:16px; }}
  .card {{ background:#fff; border:1px solid var(--line); border-radius:10px; padding:16px 18px; }}
  .card.bear {{ border-top:4px solid var(--red); }}
  .card.bull {{ border-top:4px solid var(--green); }}
  .card.navy {{ border-top:4px solid var(--navy); }}
  .card ul {{ padding-left:18px; line-height:1.7; font-size:15px; }}
  .tag {{ margin-top:10px; font-size:12px; font-weight:700; color:var(--navy2); }}
  .one-liner {{ margin-top:16px; background:var(--navy); color:#fff; padding:12px 16px; border-radius:8px; font-size:16px; }}
  .one-liner b {{ color:var(--gold); }}
  .chain {{ display:flex; align-items:center; justify-content:space-between; gap:6px; margin:8px 0 22px; }}
  .node {{ flex:1; background:#fff; border:1px solid var(--line); border-radius:10px; padding:14px 8px;
    text-align:center; font-size:14px; font-weight:700; color:var(--navy); line-height:1.35;
    box-shadow:0 4px 0 #E4E9F2; }}
  .node.hot {{ background:var(--navy); color:#fff; border-color:var(--navy); }}
  .arr {{ color:var(--gold); font-weight:800; font-size:20px; }}
  .split {{ display:grid; grid-template-columns:1fr 1fr; gap:16px; }}
  .split > div {{ background:#fff; border:1px solid var(--line); border-radius:10px; padding:16px 18px; }}
  .muted {{ color:var(--gray); font-size:14px; margin-top:6px; }}
  .vs {{ display:grid; grid-template-columns:1fr 1fr; gap:16px; flex:0; }}
  .col {{ border-radius:10px; padding:20px 22px; min-height:220px; }}
  .col.up {{ background:#E8F5E9; }}
  .col.down {{ background:#FDECEA; }}
  .col h3 {{ font-size:18px; }}
  .col.up h3 {{ color:var(--green); }}
  .col.down h3 {{ color:var(--red); }}
  .col ul {{ margin-top:10px; padding-left:18px; font-size:18px; line-height:1.8; }}
  .monitor {{ display:grid; grid-template-columns:repeat(4,1fr); gap:10px; margin-top:18px; }}
  .monitor div {{ background:#fff; border:1px solid var(--line); border-radius:8px; padding:12px; font-size:16px; font-weight:700; color:var(--navy); }}
  .monitor span {{ display:block; font-size:11px; color:var(--gray); font-weight:600; margin-bottom:4px; }}
  .bars {{ display:flex; flex-direction:column; gap:7px; }}
  .bar {{ display:grid; grid-template-columns:170px 1fr 64px; align-items:center; gap:10px; }}
  .bl {{ font-size:13px; color:var(--gray); text-align:right; }}
  .track {{ height:16px; background:#E5EAF1; border-radius:8px; overflow:hidden; }}
  .track i {{ display:block; height:100%; border-radius:8px; }}
  .track i.navy {{ background:#0F2043; }}
  .track i.blue {{ background:#1E407C; }}
  .track i.gold {{ background:#B8943A; }}
  .track i.ink {{ background:#2C3A58; }}
  .track i.gray {{ background:#6B7280; }}
  .bv {{ font-size:13px; font-weight:700; color:var(--navy); }}
  .caption {{ margin-top:14px; font-size:14px; color:var(--gray); }}
  .big {{ font-size:28px; color:var(--navy); margin:8px 0; }}
  .mini {{ width:100%; border-collapse:collapse; font-size:15px; }}
  .mini th,.mini td {{ border-bottom:1px solid var(--line); padding:8px 4px; text-align:left; }}
  .wide {{ width:100%; border-collapse:collapse; background:#fff; font-size:16px; }}
  .wide th {{ background:var(--navy); color:#fff; padding:10px 12px; text-align:center; }}
  .wide td {{ padding:11px 12px; text-align:center; border-bottom:1px solid var(--line); }}
  .wide td:first-child {{ text-align:left; font-weight:700; }}
  .lead {{ color:var(--gray); margin:-6px 0 14px; font-size:15px; }}
  .compare-note {{ display:flex; gap:16px; align-items:center; justify-content:center; margin-top:22px; font-size:20px; font-weight:800; color:var(--navy); }}
  .compare-note .arr {{ color:var(--gold); }}
  .k-bars {{ display:grid; grid-template-columns:repeat(4,1fr); gap:12px; margin:8px 0 18px; }}
  .k-bars > div {{ background:#fff; border-radius:10px; padding:18px 14px; text-align:center; border-top:4px solid var(--navy2); }}
  .k-bars > div.gold {{ border-top-color:var(--gold); }}
  .k-bars > div.ink {{ border-top-color:var(--navy); background:var(--navy); color:#fff; }}
  .k-bars em {{ display:block; font-style:normal; font-size:13px; color:var(--gray); }}
  .k-bars .ink em {{ color:#C9D3E4; }}
  .k-bars strong {{ display:block; font-size:28px; margin:6px 0 2px; }}
  .k-bars small {{ color:var(--gray); }}
  .k-bars .ink small {{ color:#C9D3E4; }}
  .plain {{ padding-left:20px; line-height:1.8; font-size:16px; }}
  .takeaway {{ padding-left:28px; font-size:18px; line-height:1.65; }}
  .takeaway li {{ margin-bottom:14px; }}
  .close {{ margin-top:auto; background:var(--navy); color:#fff; padding:14px 18px; border-radius:8px; font-size:16px; }}
  .nav {{ position:fixed; left:50%; bottom:18px; transform:translateX(-50%); display:flex; gap:8px; z-index:5; }}
  .nav button {{ background:rgba(255,255,255,.12); color:#fff; border:1px solid rgba(255,255,255,.2);
    padding:6px 12px; border-radius:6px; cursor:pointer; font-size:12px; }}
  .dots {{ position:absolute; right:16px; top:50%; transform:translateY(-50%); display:flex; flex-direction:column; gap:6px; }}
  .dots i {{ width:8px; height:8px; border-radius:50%; background:#C5CDD8; cursor:pointer; }}
  .dots i.on {{ background:var(--gold); }}
  .help {{ position:absolute; right:40px; bottom:12px; font-size:11px; color:#8A93A3; }}
  .bands {{ display:grid; grid-template-columns:1fr 1fr; gap:16px; }}
  .band-card {{ background:#fff; border:1px solid var(--line); border-radius:10px; padding:16px 18px; }}
  .band-card h3 small {{ font-size:12px; color:var(--gray); font-weight:600; margin-left:6px; }}
  .axis {{ margin:14px 0 10px; }}
  .axis-track {{ position:relative; height:18px; background:#EEF2F8; border-radius:9px; }}
  .axis-track .seg {{ position:absolute; top:3px; height:12px; border-radius:6px; }}
  .axis-track .seg.gold {{ background:#B8943A; }}
  .axis-track .seg.navy {{ background:#1E407C; }}
  .axis-track .seg.green {{ background:#166534; }}
  .axis-track .mark.now {{ position:absolute; top:-4px; width:3px; height:26px; background:#0F2043; border-radius:2px; }}
  .axis-labels {{ display:flex; justify-content:space-between; margin-top:6px; font-size:11px; color:var(--gray); }}
  .legend {{ list-style:none; padding:0; margin-top:8px; font-size:15px; line-height:1.7; }}
  .legend i {{ display:inline-block; width:10px; height:10px; border-radius:2px; margin-right:8px; }}
  .legend i.g {{ background:#B8943A; }}
  .legend i.n {{ background:#1E407C; }}
  .legend i.a {{ background:#166534; }}
</style>
</head>
<body>
<div class="deck"><div class="frame" id="frame">
{''.join(slides)}
<div class="dots" id="dots"></div>
<div class="help">← →  또는 클릭 · F 전체화면</div>
</div></div>
<script>
const slides=[...document.querySelectorAll('.slide')];
const dots=document.getElementById('dots');
let i=0;
slides.forEach((_,n)=>{{
  const d=document.createElement('i');
  d.onclick=()=>go(n);
  dots.appendChild(d);
}});
function go(n){{
  i=(n+slides.length)%slides.length;
  slides.forEach((s,k)=>s.classList.toggle('on',k===i));
  [...dots.children].forEach((d,k)=>d.classList.toggle('on',k===i));
}}
go(0);
document.addEventListener('keydown',e=>{{
  if(e.key==='ArrowRight'||e.key===' '||e.key==='PageDown') go(i+1);
  if(e.key==='ArrowLeft'||e.key==='PageUp') go(i-1);
  if(e.key==='f'||e.key==='F') document.documentElement.requestFullscreen?.();
}});
document.getElementById('frame').addEventListener('click',e=>{{
  if(e.target.closest('.dots')) return;
  go(i+1);
}});
</script>
</body>
</html>
"""
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    HTML_PATH.write_text(html, encoding="utf-8")
    return HTML_PATH


def main():
    rows = assert_all()
    _setup_chart_font()
    charts = {
        "per": chart_per_memory(),
        "targets": chart_targets(),
        "korea": chart_korea_drivers(),
        "storage": chart_storage_vs_memory(),
    }
    docx = build_docx(charts)
    html = build_html()
    print(f"Wrote {docx} ({docx.stat().st_size} bytes)")
    print(f"Wrote {html} ({html.stat().st_size} bytes)")
    print(f"{len(rows)} number checks passed")


if __name__ == "__main__":
    main()
