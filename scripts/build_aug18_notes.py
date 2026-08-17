#!/usr/bin/env python3
"""8월 18일 NON-삼전닉스 강의 노트 Word 문서 생성."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
from docx.shared import Cm, Inches, Pt, RGBColor, Emu, Twips

# Palette
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
NAVY_MID = RGBColor(0x2C, 0x4A, 0x7C)
TEAL = RGBColor(0x0F, 0x6B, 0x6B)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
RED = RGBColor(0xA3, 0x2D, 0x2D)
GREEN = RGBColor(0x1F, 0x6B, 0x3A)
GRAY = RGBColor(0x4A, 0x4A, 0x4A)
MUTED = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
SOFT_BG = "F4F6F8"
NAVY_HEX = "1B2A4A"
TEAL_HEX = "0F6B6B"
GOLD_HEX = "F8F1DE"
RED_HEX = "F8E8E8"
GREEN_HEX = "E8F4EC"
BLUE_HEX = "E8EEF6"

FONT = "WenQuanYi Micro Hei"
FONT_EAST = "WenQuanYi Micro Hei"


def set_run_font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color="D0D5DD", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def prevent_row_split(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    trPr.append(cant)


def set_table_autofit(table, autofit=True):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "autofit" if autofit else "fixed")
    tblPr.append(layout)


def add_bottom_border(paragraph, color="1B2A4A", sz="12"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_para_spacing(p, before=0, after=8, line=1.15):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_text(p, text, size=11, bold=False, color=GRAY, italic=False):
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return run


def heading1(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=18, after=8)
    add_bottom_border(p, NAVY_HEX, "16")
    add_text(p, text, size=18, bold=True, color=NAVY)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=14, after=6)
    add_text(p, text, size=14, bold=True, color=NAVY_MID)
    return p


def heading3(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=10, after=4)
    add_text(p, text, size=12, bold=True, color=TEAL)
    return p


def body(doc, text, size=11, bold=False, color=GRAY, after=6):
    p = doc.add_paragraph()
    set_para_spacing(p, before=0, after=after, line=1.2)
    add_text(p, text, size=size, bold=bold, color=color)
    return p


def bullet(doc, text, level=0, bold_prefix=None, rest=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    set_para_spacing(p, before=1, after=3, line=1.15)
    add_text(p, "•  ", size=11, bold=True, color=TEAL)
    if bold_prefix and rest is not None:
        add_text(p, bold_prefix, size=11, bold=True, color=NAVY)
        add_text(p, rest, size=11, color=GRAY)
    else:
        add_text(p, text, size=11, color=GRAY)
    return p


def numbered(doc, n, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    set_para_spacing(p, before=1, after=3, line=1.15)
    add_text(p, f"{n}.  ", size=11, bold=True, color=TEAL)
    add_text(p, text, size=11, color=GRAY)
    return p


def callout(doc, title, lines, fill=GOLD_HEX, title_color=GOLD):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_borders(cell, "E5E0D0" if fill == GOLD_HEX else "D0D5DD")
    set_cell_margins(cell, 80, 80, 120, 120)
    cell.text = ""
    p = cell.paragraphs[0]
    set_para_spacing(p, before=0, after=4)
    add_text(p, title, size=11, bold=True, color=title_color)
    for i, line in enumerate(lines):
        p2 = cell.add_paragraph()
        set_para_spacing(p2, before=0, after=2 if i < len(lines) - 1 else 0, line=1.15)
        add_text(p2, line, size=10.5, color=GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    return table


def flow(doc, steps):
    """Render A → B → C as a compact flow line."""
    p = doc.add_paragraph()
    set_para_spacing(p, before=4, after=8)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for i, step in enumerate(steps):
        if i:
            add_text(p, "  →  ", size=10, bold=True, color=TEAL)
        add_text(p, step, size=10.5, bold=True, color=NAVY)
    return p


def make_table(doc, headers, rows, col_widths=None, header_fill=NAVY_HEX):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, header_fill)
        set_cell_borders(cell, header_fill)
        set_cell_margins(cell)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, before=0, after=0)
        add_text(p, h, size=9.5, bold=True, color=WHITE)

    for r_idx, row in enumerate(rows):
        fill = "FFFFFF" if r_idx % 2 == 0 else SOFT_BG
        prevent_row_split(table.rows[r_idx + 1])
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            shade_cell(cell, fill)
            set_cell_borders(cell, "E4E7EC")
            set_cell_margins(cell)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx else WD_ALIGN_PARAGRAPH.LEFT
            set_para_spacing(p, before=0, after=0)
            add_text(p, str(val), size=9.5, color=GRAY)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)
    return table


def add_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "8월 18일 NON-삼전닉스  ·  ", size=8, color=MUTED)
    # PAGE field
    run = p.add_run()
    set_run_font(run, size=8, color=MUTED)
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    run._r.append(fld1)

    run2 = p.add_run()
    set_run_font(run2, size=8, color=MUTED)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._r.append(instr)

    run3 = p.add_run()
    set_run_font(run3, size=8, color=MUTED)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run3._r.append(fld2)

    add_text(p, "  /  ", size=8, color=MUTED)

    run4 = p.add_run()
    set_run_font(run4, size=8, color=MUTED)
    fld3 = OxmlElement("w:fldChar")
    fld3.set(qn("w:fldCharType"), "begin")
    run4._r.append(fld3)

    run5 = p.add_run()
    set_run_font(run5, size=8, color=MUTED)
    instr2 = OxmlElement("w:instrText")
    instr2.set(qn("xml:space"), "preserve")
    instr2.text = " NUMPAGES "
    run5._r.append(instr2)

    run6 = p.add_run()
    set_run_font(run6, size=8, color=MUTED)
    fld4 = OxmlElement("w:fldChar")
    fld4.set(qn("w:fldCharType"), "end")
    run6._r.append(fld4)


def add_header(section):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(p, "소부장  ·  SK  ·  Atlas    |    영상 녹화용 정리본", size=8, color=MUTED)
    add_bottom_border(p, "E4E7EC", "6")


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    add_header(section)
    add_page_number(section)

    # ---------- Cover ----------
    for _ in range(2):
        doc.add_paragraph()

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(kicker, before=0, after=6)
    add_text(kicker, "QUICK COMMENT 정리  ·  영상 녹화용", size=11, bold=True, color=TEAL)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(title, before=4, after=6)
    add_text(title, "8월 18일 NON-삼전닉스", size=26, bold=True, color=NAVY)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(sub, before=0, after=14)
    add_text(sub, "소부장  ·  SK  ·  자동차", size=16, bold=True, color=NAVY_MID)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(meta, before=0, after=4)
    add_text(meta, "테스  ·  한미반도체  ·  원익IPS  ·  SK(에코플랜트)  ·  Atlas", size=11, color=GRAY)

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(meta2, before=0, after=18)
    add_text(meta2, "영상 녹화 예정  ·  당일 오후~저녁 업로드 예상", size=10, color=MUTED)

    callout(
        doc,
        "오늘 한 줄",
        [
            "삼전·닉스가 아니라, 그 투자가 흘러가는 주변 — 장비(테스·한미·원익IPS), 지주사(SK·에코플랜트), 공장 로봇(Atlas)을 봅니다.",
            "공통 질문: 수요가 진짜로 매출·이익이 되기 전에, 주가가 얼마나 먼저 가 있는가.",
        ],
        fill=BLUE_HEX,
        title_color=NAVY,
    )

    heading2(doc, "구성")
    numbered(doc, 1, "소부장 — 테스 / 한미반도체 / 원익IPS  (+ 리노공업 당일 코멘트)")
    numbered(doc, 2, "SK 지주사 — SOTP, SK에코플랜트 재평가, 실트론 매각")
    numbered(doc, 3, "Atlas — 연구용 로봇에서 공장 투입으로, 모비스·글로비스")
    numbered(doc, 4, "당일 추가 — 메모리 밸류 스케치, Dark GPU(엔비디아 금융) 코멘트")

    # ============================================================
    heading1(doc, "1. 테스  —  수주가 먼저, BSD가 다음")
    # ============================================================

    callout(
        doc,
        "핵심",
        [
            "2Q 매출은 서프라이즈, 이익은 일회성(R&D) 비용으로 컨센 부합.",
            "진짜 트리거는 실적보다 수주잔고. 2Q 신규수주 2,823억(QoQ 2배), 기말 잔고 2,000억+ → 하반기 매출화.",
            "기존 ACL/ARC는 NAND. BSD는 DRAM/HBM으로 고객·공정을 넓히는 성장 옵션.",
            "독점 장비로 단정하면 안 됨. 퀄 통과·양산 확대가 관건.",
        ],
        fill=BLUE_HEX,
        title_color=NAVY,
    )

    heading2(doc, "1-1. 2Q26 실적과 수주")
    make_table(
        doc,
        ["항목", "내용"],
        [
            ["2Q 실적 성격", "매출 서프라이즈 / 영업이익은 일회성 비용으로 부합"],
            ["2Q 신규 수주", "2,823억원  ·  전분기 대비 약 2배"],
            ["2Q 말 수주잔고", "2,000억원+  ·  대부분 연내 매출 인식 가능"],
            ["하반기 경로", "3Q부터 상반기 수주 매출화 → 추가 수주 반영 → 연간 컨센 상향 여지"],
        ],
        col_widths=[4.5, 12.5],
    )

    heading3(doc, "고객·공정 — DRAM + NAND 동시 수혜")
    bullet(doc, "삼성전자: P4 → PH4, NAND는 시안, V9 전환 투자")
    bullet(doc, "SK하이닉스: M15X DRAM 투자 지속")
    bullet(doc, "테스의 강점 = DRAM과 NAND를 같이 먹는 포트폴리오")

    heading2(doc, "1-2. 증권사")
    make_table(
        doc,
        ["증권사", "톤", "TP", "포인트"],
        [
            ["SK증권", "유지", "23만원", "수주 2배, 잔고 흐름이 주가 트리거"],
            ["삼성증권", "상향", "20만원", "안팎으로 좋다. BSD 기여 확대. 27년 PER 30배"],
        ],
        col_widths=[3.2, 2.4, 2.6, 8.8],
    )
    body(
        doc,
        "삼성증권: 글로벌 전공정 장비 상위 5개사 평균 PER 30배를 적용. 과거에는 좁은 포트폴리오·높은 NAND 의존으로 동종 대비 할인 거래.",
        size=10.5,
    )

    heading2(doc, "1-3. BSD — 왜 HBM인가")
    body(doc, "BSD = Backside Deposition. 웨이퍼 뒷면에 막을 입혀 휘어짐(Warpage)을 잡는 장비.")
    bullet(doc, "HBM·고단화 NAND처럼 웨이퍼가 얇아질수록 휘어짐이 커짐")
    bullet(doc, "뒷면에 SiO/SiN 등을 증착 → 앞면과 반대 방향 응력 → 평탄도 확보")
    bullet(doc, "논리: HBM 고단화 → BSD 필요성 증가")
    bullet(doc, "NAND · DRAM · Foundry 모두 적용 가능")
    bullet(doc, "2025년부터 국내 메모리 고객 NAND에 공급")
    bullet(doc, "현재 HBM용 DRAM, 다른 고객 NAND/Foundry 퀄리피케이션 진행")
    body(
        doc,
        "기존 ACL/ARC가 NAND 투자와 묶여 있었다면, BSD는 DRAM/HBM으로 고객과 공정을 확장하는 장비. 2027년까지 실적 레벨업의 신규 성장 옵션 중 BSD가 더 중요.",
    )

    heading3(doc, "“테스만 공급?” — 지금은 단정 금지")
    bullet(doc, "글로벌에서 테스만 만들 수 있는 독점 장비로 보기는 어려움")
    bullet(doc, "중요한 것은 장비 독점성보다 삼성·하이닉스 퀄 통과와 양산 확대")
    bullet(doc, "공개자료 기준: 이미 국내 메모리 NAND 공급, DRAM/HBM 확대 평가 중")

    heading2(doc, "1-4. Tetra — 고처리량 PECVD")
    bullet(doc, "Quad chamber. 한 번에 웨이퍼 4장 → throughput 향상")
    bullet(doc, "SiCN 박막 증착. 국내 메모리 고객 DRAM 퀄 진행")
    bullet(doc, "기존 장비 대비 ASP가 높은 신규 장비로 평가")
    bullet(doc, "독점 근거 없음. AMAT·Lam 등 글로벌 PECVD와 경쟁, 국내에도 증착 업체 다수")
    bullet(doc, "의미는 ACL/ARC 중심에서 DRAM 선단으로 포트폴리오를 넓힌다는 점")

    heading2(doc, "1-5. TP 시나리오")
    body(doc, "테스·한미반도체 TP 시나리오의 상세 숫자는 강의에서 풀 예정. 아래는 당일 코멘트에 적힌 원익IPS 쪽 숫자와 구분.")
    body(
        doc,
        "테스 공개 TP: SK 23만원(유지), 삼성 20만원(상향, 27년 PER 30배).",
        size=10.5,
    )

    # ============================================================
    heading1(doc, "2. 한미반도체  —  실적은 강하다. 질문은 점유율")
    # ============================================================

    callout(
        doc,
        "핵심",
        [
            "2Q는 사상 최대. HBM4 양산 → TC 본더. 관전 포인트는 한화세미텍·마이크론·삼성향.",
            "BofA는 Neutral→Buy, TP 30→42만(보도에 50만). 2028 EPS × 47배. 과거 평균 PER ~60배보다 낮게 적용했다고 설명.",
            "같은 47배를 상상인 이익 추정에 넣으면 약 32만원. 멀티플보다 EPS 가정이 갈림.",
            "신사업(Logic TCB, HBF, MSVP/PLP)이 ‘HBM 원툴’이 아님을 증명하면 프리미엄 논리가 살아남.",
        ],
        fill=BLUE_HEX,
        title_color=NAVY,
    )

    heading2(doc, "2-1. 2Q26 — 숫자")
    make_table(
        doc,
        ["항목", "수치", "의미"],
        [
            ["매출", "2,511억원  (+39.5% YoY)", "HBM4 양산 본격화"],
            ["영업이익", "1,303억원  (+51.0% YoY)", "레버리지 확인"],
            ["영업이익률", "51.9%", "사상 최대 분기"],
            ["전방", "HBM4E, 연말~내년 초 준비", "12단·16단 차세대 TCB 수요 이어질 가능성"],
        ],
        col_widths=[3.5, 6.5, 7.0],
    )

    heading2(doc, "2-2. BofA 시각")
    flow(doc, ["2월 Neutral → Buy, TP 30만", "4월 42만", "이후 보도 50만"])
    bullet(doc, "하반기 매출 약 5,200억원 가정, 실적 회복을 강하게 봄")
    heading3(doc, "BofA 다섯 가지")
    numbered(doc, 1, "TCB 시장 지배력")
    numbered(doc, 2, "고객 다변화: SK하이닉스 → Micron → 중국 OSAT")
    numbered(doc, 3, "삼성전자 공급 가능성")
    numbered(doc, 4, "용인 SK하이닉스 + 삼성 P5 신규 팹")
    numbered(doc, 5, "2028년부터 신규 메가팹 TCB 수요 본격화")

    heading3(doc, "42만원의 산식, 그리고 다른 답")
    bullet(doc, "BofA 42만원 = 2028년 EPS × 47배. 과거 평균 PER ~60배보다 낮게 적용 → 성장은 보되 밸류는 보수적이라는 주장")
    bullet(doc, "상상인 이익 추정에 같은 28년 PER 47배를 적용하면 약 32만원")
    bullet(doc, "상상인(7/10) TP 38만원: 27~28년 평균 EPS 6,057원 × PER 62.4배(직전연도 56.7배 +10%)")

    heading2(doc, "2-3. 변수 세 개")
    body(doc, "한미반도체의 리스크 = 한화세미텍. 과거에는 ‘하이닉스 HBM → 한미 TCB’가 거의 공식.")
    body(doc, "앞으로는 한화세미텍과의 경쟁·공급망 다변화. 시장이 확인하려는 식은 하나.")
    callout(
        doc,
        "확인할 곱셈",
        ["TCB 시장 성장률  ×  한미반도체 점유율"],
        fill=GOLD_HEX,
        title_color=GOLD,
    )
    numbered(doc, 1, "한화세미텍 — 점유율 잠식 속도")
    numbered(doc, 2, "마이크론 발주")
    numbered(doc, 3, "가장 큰 것: 삼성전자향 수주 여부")

    heading2(doc, "2-4. 그래도 한미가 유리한 이유")
    flow(doc, ["HBM3E", "HBM4", "HBM4E", "12단 → 16단 이상"])
    body(doc, "세대·적층이 올라갈수록 본딩 정밀도와 생산성이 중요해짐. 난이도가 한미의 해자.")
    heading3(doc, "작업 가정")
    bullet(doc, "HBM4/HBM4E 수요는 강함")
    bullet(doc, "한화세미텍 점유율은 상승")
    bullet(doc, "한미 TCB 점유율 55~60%")
    bullet(doc, "Micron·중국 고객 증가")
    bullet(doc, "삼성 공급은 일부 장비부터")
    bullet(doc, "영업이익률 45~50% 유지")

    heading2(doc, "2-5. 신사업 — HBM 원툴이 아니라는 증명")
    body(doc, "상상인 논리: HBM 하나에만 의존하지 않는다는 것이 증명되면, 시장이 넓어진다.")
    make_table(
        doc,
        ["축", "내용", "의미"],
        [
            ["Logic TC 본더", "CoWoS 외주화, OSAT의 로직용 TCB", "한미의 신규 시장"],
            ["HBF TC 본더", "High Bandwidth Flash, NAND 기반 AI 메모리", "초기 시장, 장기 옵션"],
            ["MSVP", "대면적 기판 절단. AI 패키징 PLP 확대", "MSVP 매출 빠른 증가 관찰"],
        ],
        col_widths=[3.8, 7.2, 6.0],
    )
    body(doc, "강의에서 물을 질문: Logic·HBF·PLP 성공이라는 성장 옵션을 밸류에 얼마나 넣을 것인가.")

    # ============================================================
    heading1(doc, "3. 원익IPS  —  2Q는 부진, 논쟁은 2027년")
    # ============================================================

    callout(
        doc,
        "핵심",
        [
            "주류: 2Q 실적 부진 ≠ 수주 사이클 훼손. 잔고 증가 → 3Q~4Q 매출화 → 하반기 이익 레벨업.",
            "8/7 전후 대부분 증권사는 TP를 낮춰도 톤은 긍정. BNK만 전방 둔화·주가 정점 통과 가능성을 본다.",
            "논쟁의 핵심은 ‘실적이 좋아지느냐’가 아니라 ‘2027년까지 성장이 이어지느냐’.",
            "장비주는 실적보다 주가가 먼저 움직인다. 실적 정점 ≠ 주가 정점.",
        ],
        fill=BLUE_HEX,
        title_color=NAVY,
    )

    heading2(doc, "3-1. 왜 주류는 긍정인가")
    flow(doc, ["1Q 수주잔고 약 4,000억", "2Q 신규 수주 견조", "3Q~4Q 잔고 매출화", "하반기 매출·OP 증가"])
    bullet(doc, "수주가 아직 꺾이지 않았다")
    bullet(doc, "2Q 실적 부진을 수주 사이클 훼손으로 읽지 않는다")

    heading2(doc, "3-2. 하반기 체크포인트")
    make_table(
        doc,
        ["시점", "무엇이 일어나는가", "확인할 것"],
        [
            ["1H26", "수주는 증가, 실적은 기대 이하", "이미 지나간 사실"],
            ["3Q26", "수주잔고 본격 매출화", "영업이익 증가폭이 확대되는가"],
            ["4Q26", "매출 증가 + 영업 레버리지", "매출↑와 OPM 개선이 동시에 오는가"],
        ],
        col_widths=[3.0, 7.0, 7.0],
    )

    heading2(doc, "3-3. 2027년이 승부처")
    make_table(
        doc,
        ["", "Bull", "Bear (BNK)"],
        [
            ["전방", "DRAM + NAND + Foundry + CXMT", "CAPEX 증가율 둔화"],
            ["수주", "투자 지속 → 신규 수주 증가", "신규 수주 둔화"],
            ["실적", "2027 매출·이익 증가", "2027 이익 성장 제한"],
            ["주가", "현재 밸류 부담 완화", "현재 주가 선반영 부담"],
        ],
        col_widths=[2.8, 7.1, 7.1],
    )

    callout(
        doc,
        "BNK 경고를 왜 남기는가",
        [
            "실적 개선을 부정하는 것이 아니다. 실적 개선이 이미 주가에 상당 부분 반영됐다는 경고다.",
            "주류 경로: 수주 증가 → 실적 증가 → 주가 상승",
            "BNK 경로: 수주 증가 → 실적 증가 기대 → 주가 선반영 → 이후 CAPEX 둔화",
        ],
        fill=RED_HEX,
        title_color=RED,
    )

    heading2(doc, "3-4. 앞으로 딱 3개만 확인")
    numbered(doc, 1, "3Q26 영업이익 증가폭이 확대되는가")
    numbered(doc, 2, "4Q26 매출 증가와 OPM 개선이 동시에 발생하는가")
    numbered(doc, 3, "2027년 신규 수주가 다시 증가하는가")
    body(doc, "3개 모두 YES → 주류 전망.   수주 둔화 + OPM 개선 실패 → BNK 정점론.")

    heading2(doc, "3-5. TP 밴드 (8/8 리포트 기준)")
    make_table(
        doc,
        ["방법", "산출"],
        [
            ["증권사 TP", "8.8만 / 14만 / 15.3만 / 16.7만 / 19만  ·  평균 14.8만"],
            ["평균 × 안전마진 약 30%", "10.4만 ~ 11.8만원"],
            ["BNK 제외, 26년 50% + 27년 50% 이익에 PER 20~25배", "8.4~10.5 / 9.6~12 / 9.8~12.2 / 11.9~14.9만"],
        ],
        col_widths=[7.5, 9.5],
    )
    body(doc, "강의 멘트: 테스·한미 TP 시나리오는 영상에서 별도로 푼다.", size=10.5, color=MUTED)

    # ============================================================
    heading1(doc, "4. 리노공업  —  실적은 반박했다. 남은 것은 파업")
    # ============================================================

    body(doc, "오늘 본편 소부장은 테스·한미·원익IPS. 리노는 당일 오후 코멘트로 같이 묶어 둔다.")

    callout(
        doc,
        "핵심",
        [
            "2Q는 물량·판가 동반 상승, OPM 51.3%, 컨센 상회. 파업 직전까지는 강한 성장 국면.",
            "7/23 창사 이래 첫 무기한 총파업, 8/17 기준 약 4주. 최대 변수.",
            "시장이 더 보는 것은 매출 감소보다 마진. OPM 50%+ → 임금체계가 구조적으로 바뀌면 40%대 중반 우려.",
            "삼성증권: 실적으로 우려를 재차 반박. 다만 TP 10만원(27년 PER 32배)으로 하향.",
            "파업이 잘 끝나면 peers에서 리노로 순환매가 가능한 시기가 올 수 있다.",
        ],
        fill=BLUE_HEX,
        title_color=NAVY,
    )

    heading2(doc, "4-1. 2Q26")
    make_table(
        doc,
        ["항목", "수치"],
        [
            ["매출", "약 1,432억원  ·  YoY +30%"],
            ["영업이익", "약 735억원  ·  YoY +37%"],
            ["영업이익률", "51.3%"],
            ["상반기 영업이익", "1,208억원  ·  +36.7%"],
            ["믹스", "핀·소켓 물량 + ASP 동반 상승. AI/HPC 테스트 수요 견조"],
        ],
        col_widths=[4.5, 12.5],
    )

    heading2(doc, "4-2. 노조 vs 회사")
    make_table(
        doc,
        ["노조 요구", "회사 측 우려"],
        [
            ["연 800% 고정상여", "인건비가 매출액의 약 30%까지 상승 가능"],
            ["영업이익의 15% 성과급", "생산 차질 → 3Q 매출 감소 가능성"],
            ["", "장기화 시 고객의 경쟁사 전환 가능성"],
        ],
        col_widths=[8.5, 8.5],
    )
    body(doc, "삼성증권 톤: 판가 확대가 지속되면 추후 비용 구조 변화 리스크도 상쇄될 수 있다. 동시에 TP는 낮췄다.")

    # ============================================================
    heading1(doc, "5. SK  —  할인율이 좁혀질 이유가 늘었다")
    # ============================================================

    callout(
        doc,
        "핵심",
        [
            "대신: TP 88만원 유지. 2Q 연결 영업이익 4.84조(+2,212.5% YoY). 자회사 이익 체력 확인.",
            "SOTP NAV 약 81.7조 × 할인율 41%. 에코플랜트를 건설사 장부가로 두면 너무 싸다.",
            "에코플랜트 2Q OP 5,340억, 수주잔고 26.8조. 반도체 CAPEX + AI DC + 메모리 유통.",
            "에코플랜트만 재평가해도 TP 96~106만 밴드가 열림. 기본 축은 여전히 스퀘어(하이닉스)·이노베이션.",
            "실트론 70.6%를 두산에 2.3조 매각 → 가정 1.8조 대비 NAV +0.5조, 재무·환원 재원.",
        ],
        fill=BLUE_HEX,
        title_color=NAVY,
    )

    heading2(doc, "5-1. 대신증권 2Q 리뷰")
    make_table(
        doc,
        ["항목", "수치"],
        [
            ["매출", "42조 1,247억원"],
            ["영업이익", "4조 8,412억원  (+2,212.5% YoY)"],
            ["메시지", "포트폴리오 리밸런싱 이후 자회사 이익 체력 확인"],
            ["에코플랜트", "분기 최대 매출, AI DC 실적 반영 본격화"],
        ],
        col_widths=[4.5, 12.5],
    )

    heading3(doc, "SOTP (대신, TP 88만원 유지)")
    make_table(
        doc,
        ["구성", "금액"],
        [
            ["자체사업 + 로열티 영업가치", "13.0조원"],
            ["상장 자회사 지분가치", "73.4조원"],
            ["비상장 자회사 지분가치", "3.8조원"],
            ["별도 순차입금", "△ 8.5조원"],
            ["NAV", "81.7조원"],
            ["Target 할인율", "41%"],
        ],
        col_widths=[8.0, 9.0],
    )
    body(doc, "참고: 시나리오 계산에서는 NAV 81.6조 × (1−41%) = 48.2조로 적어 두었다. 반올림 차이.", size=10, color=MUTED)

    heading2(doc, "5-2. SK에코플랜트 — 왜 따로 보나")
    body(doc, "한 줄: SK하이닉스의 반도체 투자와 AI 데이터센터 투자를 동시에 먹는 회사.")
    body(doc, "반도체 공장 건설 + AI DC + 반도체 소재/가스 + 메모리 유통/재활용. 반도체 관련 매출 ≈ 80%+.")

    heading3(doc, "4대 사업 (26년 상반기 매출 비중)")
    make_table(
        doc,
        ["사업", "비중", "내용"],
        [
            ["Asset Lifecycle", "45%", "메모리 모듈 유통 · IT 자산 재활용 (Essencore 등)"],
            ["Hi-Tech", "33%", "반도체 Fab · AI 데이터센터 건설"],
            ["Solution", "18%", "기존 건설 · 인프라"],
            ["Gas & Material", "4%", "반도체용 가스 · 소재"],
        ],
        col_widths=[4.2, 2.4, 10.4],
    )

    heading3(doc, "2Q26 — 왜 급증했나")
    make_table(
        doc,
        ["항목", "수치"],
        [
            ["매출", "5.15조원  ·  +68.4% YoY  ·  분기 최대"],
            ["영업이익", "5,340억원  ·  +265.7% YoY"],
            ["EBITDA 마진", "11.8%  ·  +5.1%p YoY"],
            ["신규 수주", "7.8조원"],
            ["수주잔고", "26.8조원  →  향후 매출 가시성"],
        ],
        col_widths=[4.5, 12.5],
    )

    heading3(doc, "실적 엔진 3개")
    numbered(doc, 1, "반도체 CAPEX — 용인 클러스터 + 청주 M15X → Hi-Tech EPC")
    numbered(doc, 2, "AI CAPEX — 울산 AI 데이터센터 등, AI DC 공정 본격화")
    numbered(doc, 3, "DRAM·NAND 가격 상승 — Essencore 메모리 유통 매출 증가")
    body(doc, "수주의 핵심은 Hi-Tech 대형 프로젝트 확대.")

    heading2(doc, "5-3. 밸류 — 표의 2.1조는 너무 낮다")
    body(doc, "표: 지분율 71.2%, 장부가 1.271조, Value 2.106조.")
    body(
        doc,
        "주의: 1.271조 × 71.2% = 0.905조라서, 장부가 1.271조가 ‘지분 장부가’인지 ‘100% 장부가’인지 표와 안 맞는다. 강의에서 표 정의를 한 번 짚고 넘어갈 것.",
        size=10.5,
    )
    bullet(doc, "2Q OP 5,340억 × 4 = 연환산 약 2.14조. 수주잔고 26.8조.")
    bullet(doc, "반도체 EPC + AI DC + 메모리 유통으로 이익 체력이 올라왔다면, 기업가치 2.1조는 상당히 낮음.")

    heading3(doc, "시나리오")
    make_table(
        doc,
        ["", "시나리오 A", "시나리오 B"],
        [
            ["이익 가정", "OP 1.6~2조 → 순이익 1.0조", "순이익 1.5조 (성장·수주)"],
            ["PER 10배", "기업가치 약 10조", "기업가치 15조"],
            ["SK 지분 71.2%", "약 7.1조", "약 10.7조"],
            ["표의 2.1조 대비", "약 +5조", "약 +8.6조"],
        ],
        col_widths=[4.2, 6.4, 6.4],
    )

    heading3(doc, "대신 TP를 다시 돌리면")
    bullet(doc, "현재: 81.6조 × 59% = 48.2조")
    bullet(doc, "에코플랜트 A/B 평균 차이 약 8조를 더해 NAV 89.5조")
    bullet(doc, "할인율 그대로면 52.8조 → 현재 TP 대비 약 +10%")
    bullet(doc, "할인율 41% → 35%(지분가치 65%)면 89.5조 × 65% = 58.2조 → 약 +21%")
    bullet(doc, "이 경로면 TP 96~106만원도 산출 가능")

    heading3(doc, "증권사 TP 분포 (코멘트 기준)")
    bullet(doc, "최고 100만원 — 6/24 흥국. 이후 97만 1개, 88만 1개, 85만 2개, 73.8만 1개")
    bullet(doc, "안전마진을 충분히 넣어도 긍정 시각 가능")
    callout(
        doc,
        "기본은 잊지 말 것",
        ["SK의 축은 SK스퀘어(즉 하이닉스)와 SK이노베이션. 에코플랜트는 할인율을 좁히는 추가 이유다."],
        fill=GOLD_HEX,
        title_color=GOLD,
    )

    heading2(doc, "5-4. SK실트론 매각")
    bullet(doc, "지분 70.6%를 두산에 2.3조원 매각 공시")
    bullet(doc, "목적: 재무구조 개선, 투자재원, 초과이익 공유 예정")
    bullet(doc, "EV/EBITDA 8배 가정 시 지분가치 1.8조 → 2.3조 매각으로 NAV +0.5조")
    bullet(doc, "순자산 증가 + 주주환원 재원")

    # ============================================================
    heading1(doc, "6. Atlas  —  로봇이 아니라 Fleet × 가동률 × 데이터")
    # ============================================================

    callout(
        doc,
        "핵심",
        [
            "투자 포인트가 ‘로봇을 얼마나 잘 만드느냐’에서 ‘얼마나 많이, 오래, 실제 공장에서 일하게 하느냐’로 이동.",
            "경로: RMAC 학습 → 성능검증 → HMGMA·글로비스 파일럿 → 실제 생산. 2028 HMGMA, 2030 부품 조립 확대.",
            "현대차그룹 captive(제조 데이터 + 물류 데이터 + 한·미·중 훈련 거점)가 오히려 해자.",
            "수혜 가시성: 모비스 = 액추에이터 양산, 글로비스 = 서열(Sequencing) 물류 파일럿.",
        ],
        fill=BLUE_HEX,
        title_color=NAVY,
    )

    heading2(doc, "6-1. 단계가 바뀌었다")
    flow(doc, ["연구용 로봇", "RMAC 학습", "성능검증", "HMGMA·글로비스 파일럿", "실제 생산 투입"])
    bullet(doc, "2028년 HMGMA 투입")
    bullet(doc, "2030년 부품 조립공정 확대 목표")

    heading2(doc, "6-2. 평가 지표가 바뀐다")
    body(doc, "제조 로봇의 핵심은 수천 번을 안정적으로 반복하는 것.")
    flow(doc, ["신뢰성 ↑", "가동률 ↑", "데이터 ↑", "AI 성능 ↑", "추가 배치 ↑"])
    body(doc, "이 루프가 Physical AI Data Flywheel. 앞으로 볼 것은 스펙이 아니라 배치 대수·가동시간·활용률.")

    heading3(doc, "Captive가 약점이 아니라 해자")
    bullet(doc, "HMGMA = 제조 데이터")
    bullet(doc, "현대글로비스 = 물류 데이터")
    bullet(doc, "한국·미국·중국 로봇 훈련 거점 + 글로벌 생산 거점")
    bullet(doc, "학습 → 배치 → 데이터 → 재학습을 빠르게 반복 가능")

    heading2(doc, "6-3. 수혜주")
    make_table(
        doc,
        ["종목", "포인트", "왜 지금인가"],
        [
            [
                "현대모비스",
                "Atlas 1대당 바디 액추에이터 31개. 11월 시제품 납품.",
                "양산 확대 = 부품 물량. 로봇 부품 사업화 가시성.",
            ],
            [
                "현대글로비스",
                "부품 서열(Sequencing)은 반복·구조화가 높아 초기 상용화에 적합.",
                "파일럿 확인 시 현대차=제조, 글로비스=물류의 초기 생태계.",
            ],
        ],
        col_widths=[3.4, 7.3, 6.3],
    )

    callout(
        doc,
        "검증 이벤트",
        [
            "HMGMA·현대글로비스의 실제 배치가 Atlas 상용화의 가장 중요한 확인 지점.",
            "모비스는 로봇 부품 양산 수혜의 핵심 후보.",
        ],
        fill=GREEN_HEX,
        title_color=GREEN,
    )

    # ============================================================
    heading1(doc, "7. 당일 추가 ①  —  메모리 밸류 스케치")
    # ============================================================

    body(doc, "오늘은 NON-삼전닉스 본편. 오후 코멘트에 숫자가 있어 별첨으로만 정리. 8/14 종가 기준.")

    heading2(doc, "7-1. 글로벌 피어")
    make_table(
        doc,
        ["종목", "가격", "밸류"],
        [
            ["Micron", "971.66달러", "Fwd 12M PER 7.8배  ·  CY27 EPS 150달러 기준 PER 6.5배"],
            ["Sandisk", "1,644.11달러", "FY27 EPS 201달러 가정 PER 8.2배 (1Q 가이던스 45달러, 이후 QoQ +10/+5/+5)"],
            ["Kioxia", "오후 +10%대 (키옥시아/CXMT 각 +15%/+10%)", "7~9월 가이던스×4, 12M fwd PER 약 6배 중반"],
            ["CXMT", "공모가 대비 7배", "중국 메모리 센티멘트"],
        ],
        col_widths=[3.2, 5.5, 8.3],
    )

    heading2(doc, "7-2. 하이닉스 ADR vs 본주")
    make_table(
        doc,
        ["구분", "가격", "26Y PER", "27Y PER"],
        [
            ["SK하이닉스 ADR", "166.33달러 (235.9만원, 1,418.5원/$)", "6.8배", "5.4배"],
            ["SK하이닉스 본주", "164.5만원", "4.8배", "3.8배"],
            ["삼성전자", "27.45만원", "5.7배", "4.1배"],
        ],
        col_widths=[4.2, 7.4, 2.7, 2.7],
    )
    bullet(doc, "ADR은 마이크론 대비 −17%. 과거 할인 −20~−50%와 비교하면 할인이 이미 많이 줄었음")
    bullet(doc, "ADR 본주 대비 프리미엄 43%")
    bullet(doc, "정상 프리미엄 +20%(TSMC 약 15%)로 보면 본주 196만원이 정합")
    bullet(doc, "최근 실제 30%대 프리미엄(30~35%)을 적용하면 본주 175~181만원")
    bullet(doc, "ADR 프리미엄을 ‘인정한다’ 수준만 반영하면 하이닉스 170~180만, 같은 흐름이면 삼전 29만원대")

    heading3(doc, "컨센 이익")
    make_table(
        doc,
        ["", "26년 OP / EPS", "27년 OP / EPS"],
        [
            ["SK하이닉스", "266조 / 346,000원", "392조 / 437,000원"],
            ["삼성전자", "391조 / 47,900원", "549조 / 67,200원"],
        ],
        col_widths=[4.0, 6.5, 6.5],
    )

    heading3(doc, "다른 접근 — 27년 성장 0, 26년 PER 6~7배 (과거 사이클 4~8배)")
    bullet(doc, "하이닉스 208만 ~ 242만")
    bullet(doc, "삼성전자 28.7만 ~ 33.5만")
    body(doc, "보수적 성장 가정인데도 밴드가 이렇게 나오는 스케치. 본편 강의 숫자는 아님.", size=10.5, color=MUTED)

    # ============================================================
    heading1(doc, "8. 당일 추가 ②  —  Dark GPU, 금융이 붙을 때")
    # ============================================================

    body(
        doc,
        "출처 스케치: 월가CN / 중국 언론 요약. 트럼프 기술고문 David Sacks, 엔비디아 5,000억 달러 AI 금융 계획. "
        "https://wallstreetcn.com/articles/3779545",
        size=10,
        color=MUTED,
    )

    callout(
        doc,
        "한 줄",
        [
            "엔비디아 GPU 수요가 당장 꺾인다는 말이 아니다.",
            "AI 인프라에 금융이 붙으면서 ‘과잉투자 위험’이 커질 수 있다는 경고다.",
            "Sacks는 AI 수요 자체를 부정하지 않았다.",
        ],
        fill=BLUE_HEX,
        title_color=NAVY,
    )

    heading2(doc, "8-1. Dark GPU란")
    body(doc, "닷컴 버블 때 통신사가 인터넷 수요를 낙관해 광섬유를 깔고, 붕괴 후 쓰이지 않은 구간을 Dark Fiber라고 불렀다.")
    flow(doc, ["데이터센터 건설", "GPU 대량 구매", "AI 수요 예상", "실제 사용률 부족", "유휴 연산 = Dark GPU"])
    body(doc, "GPU가 고장 난다는 뜻이 아니다. 비싸게 샀는데 충분히 안 돌아가는 연산능력. Sacks가 지목한 AI 인프라의 가장 큰 잠재 위험.")

    heading2(doc, "8-2. 금융이 붙으면 속도가 빨라진다")
    flow(doc, ["AI 수요", "투자", "GPU 구매"])
    body(doc, "위 경로의 속도가 빨라진다. 실제 수요가 예상보다 느리면 반대 레버리지.")
    flow(
        doc,
        ["GPU 과잉", "가동률 하락", "임대료·가격 하락", "DC 수익성 악화", "금융비용 부담", "투자 축소"],
    )

    heading2(doc, "8-3. 역설 — 정치·물리가 Dark GPU를 막을 수 있다")
    body(doc, "Sacks: 데이터센터 건설의 정치적·물리적 장애물이 과잉공급을 막아주는 브레이크가 될 수 있다.")
    callout(
        doc,
        "안전한 등식",
        ["AI 수요 증가 속도  =  데이터센터 건설 속도"],
        fill=GOLD_HEX,
        title_color=GOLD,
    )
    body(doc, "GPU만 주문한다고 바로 안 돌아간다.")
    flow(doc, ["전력", "송전망", "부지", "인허가", "냉각", "데이터센터", "GPU", "네트워크"])
    body(doc, "정치적 반발과 인허가 지연이, 역설적으로 AI 공급과잉을 막는 브레이크가 될 수 있다는 논리.")

    heading2(doc, "8-4. 그래서 양면이 남는다")
    bullet(doc, "모든 것은 미래 어느 시점의 최종 수요. 그걸 정확히 맞히는 사람은 없다")
    bullet(doc, "그래서 AI 모델, 데이터센터, GPU 업체 전망에 기대는 것")
    bullet(doc, "투자는 해야 하는데 자금이 딸릴 것 같으니 서로 묶어서 가보자는 것이 금융 계획")
    bullet(doc, "이 우려·시나리오·병목이 말끔히 해소되기 전까지 양면적 시각은 계속된다")

    # ============================================================
    heading1(doc, "9. 녹화 큐시트")
    # ============================================================

    make_table(
        doc,
        ["블록", "종목·주제", "말할 한 줄", "시청자가 가져갈 질문"],
        [
            [
                "오프닝",
                "NON-삼전닉스",
                "오늘은 삼전·닉스가 아니라 그 돈이 흘러가는 주변.",
                "장비·지주·로봇 중 어디를 볼 것인가.",
            ],
            [
                "1",
                "테스",
                "실적보다 수주. BSD는 NAND에서 HBM으로 가는 옵션.",
                "퀄·양산 확대가 보이는가. 독점으로 포장하지 말 것.",
            ],
            [
                "2",
                "한미반도체",
                "2Q는 사상 최대. 이제 곱셈은 시장×점유율.",
                "한화세미텍, 마이크론, 삼성향. 신사업을 밸류에 얼마나 넣을 것인가.",
            ],
            [
                "3",
                "원익IPS",
                "2Q 부진은 이미 알려짐. 논쟁은 2027년.",
                "3Q OP 증가폭, 4Q 매출+OPM, 27년 수주. 세 개.",
            ],
            [
                "4",
                "SK",
                "할인율 41%가 좁혀질 이유가 늘었다. 에코플랜트 2.1조는 싸다.",
                "그래도 축은 스퀘어·이노베이션. 실트론은 NAV +0.5조.",
            ],
            [
                "5",
                "Atlas",
                "좋은 로봇이 아니라 공장에서 오래 일하는 로봇.",
                "HMGMA·글로비스 배치, 모비스 31개 액추에이터.",
            ],
            [
                "클로징",
                "공통",
                "수요·수주·금융이 붙을수록 양면이 남는다.",
                "선반영인가, 아직 실적으로 증명 중인가.",
            ],
        ],
        col_widths=[2.0, 3.2, 6.0, 5.8],
    )

    heading2(doc, "클로징 멘트 초안")
    body(
        doc,
        "소부장은 수주가 실적보다 먼저이고, 주가는 그 수주보다도 먼저 갈 수 있습니다. "
        "원익IPS의 BNK 코멘트가 그 문장입니다. 한미는 실적이 이미 강하니 점유율과 삼성향을 보고, "
        "테스는 BSD가 HBM으로 넘어가는지만 확인하면 됩니다.",
    )
    body(
        doc,
        "SK는 하이닉스·이노베이션이 본체이고, 에코플랜트와 실트론은 할인율을 좁히는 이유입니다. "
        "Atlas는 로봇 스펙이 아니라 공장 배치입니다. "
        "오후 코멘트의 Dark GPU는, 금융이 붙으면 속도가 빨라지고 수요가 느리면 반대로 접힌다는 양면입니다. "
        "병목이 해소되기 전까지 이 양면은 계속됩니다.",
    )

    heading2(doc, "숫자 메모 — 강의 중 실수하지 말 것")
    make_table(
        doc,
        ["실수하기 쉬운 곳", "코멘트 원문 기준"],
        [
            ["테스 신규수주 / 잔고", "2Q 신규 2,823억(QoQ 2배), 기말 잔고 2,000억+"],
            ["테스 TP", "SK 23만 유지, 삼성 20만 상향. 8.8~19만은 원익IPS"],
            ["한미 2Q", "매출 2,511억 / OP 1,303억 / OPM 51.9%"],
            ["BofA vs 상상인", "같은 28년 PER 47배여도 42만 vs 32만. EPS 가정 차이"],
            ["원익 잔고", "1Q 수주잔고 약 4,000억"],
            ["리노 파업", "7/23 시작, 8/17 약 4주. OPM 51.3% → 우려는 40%대 중반"],
            ["SK 연결 2Q", "매출 42.12조, OP 4.84조, +2,212.5% YoY"],
            ["에코플랜트 2Q", "매출 5.15조, OP 5,340억, 수주 7.8조, 잔고 26.8조, 지분 71.2%"],
            ["실트론", "70.6%, 2.3조, 가정 1.8조 대비 +0.5조"],
            ["모비스", "Atlas 1대당 바디 액추에이터 31개, 11월 시제품"],
            ["하이닉스 ADR", "166.33달러 = 235.9만원 @ 1,418.5원, 본주 대비 +43%"],
            ["엔비디아 금융", "5,000억 달러. Dark GPU ≠ GPU 고장"],
        ],
        col_widths=[4.8, 12.2],
    )

    p = doc.add_paragraph()
    set_para_spacing(p, before=16, after=4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "—  당일 Quick 코멘트 재구성  ·  원문 숫자 유지  ·  추정 추가 없음  —", size=9, color=MUTED)

    out = Path("/workspace/docs/8월 18일 NON-삼전닉스 (소부장, SK, 자동차).docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"saved {out} ({out.stat().st_size} bytes)")


if __name__ == "__main__":
    build()
