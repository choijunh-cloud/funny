"""Convert the simple markdown summary (headings + bullets + paragraphs) to .docx."""
import sys

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


def set_korean_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def convert(md_path: str, docx_path: str) -> None:
    doc = Document()
    set_korean_font(doc)

    with open(md_path, encoding="utf-8") as f:
        lines = f.read().splitlines()

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=3)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=0)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    convert(sys.argv[1], sys.argv[2])
