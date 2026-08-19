#!/usr/bin/env python3
"""8월 19일 주요 시장 코멘트 강의노트(.docx) 생성."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Mm, Pt, RGBColor

OUT_PATH = Path(
    "/workspace/lectures/8월 19일 주요 시장 코멘트 (매크로·하이닉스 환원·파운드리·소부장).docx"
)

KR_FONT = "맑은 고딕"
EN_FONT = "Calibri"

NAVY = RGBColor(0x0F, 0x20, 0x43)
NAVY2 = RGBColor(0x1E, 0x40, 0x7C)
GOLD = RGBColor(0xB8, 0x94, 0x3A)
GRAY = RGBColor(0x4B, 0x55, 0x63)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x16, 0x65, 0x34)
RED = RGBColor(0x99, 0x1B, 0x1B)
AMBER = RGBColor(0x7A, 0x5C, 0x12)

NAVY_HEX = "0F2043"
NAVY2_HEX = "1E407C"
GOLD_HEX = "B8943A"
LIGHT_HEX = "EEF2F8"
GREEN_HEX = "E8F5E9"
RED_HEX = "FDECEA"
AMBER_HEX = "FFF8E7"
BLUE_HEX = "E8F1FB"
ROW_HEX = "F7F9FC"
WHITE_HEX = "FFFFFF"


def set_run_font(run, size=11, bold=False, color=DARK, italic=False, font=KR_FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), KR_FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(
        parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill}"/>')
    )


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(
        parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'<w:top w:w="{top}" w:type="dxa"/>'
            f'<w:left w:w="{left}" w:type="dxa"/>'
            f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
            f'<w:right w:w="{right}" w:type="dxa"/>'
            f"</w:tcMar>"
        )
    )


def set_table_borders(table, color="D0D7E2", sz="4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f"</w:tblBorders>"
    )
    tbl_pr.append(borders)


def set_left_accent(cell, color=NAVY_HEX, sz="24"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(
        parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="nil"/>'
            f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="nil"/>'
            f'<w:right w:val="nil"/>'
            f"</w:tcBorders>"
        )
    )


def prevent_row_split(row):
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()
    tr_pr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))


def cell_text(cell, text, size=10, bold=False, color=DARK, align="left", font=KR_FONT):
    cell.text = ""
    align_enum = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    lines = str(text).split("\n")
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = align_enum
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1 if i < len(lines) - 1 else 0)
        p.paragraph_format.line_spacing = 1.12
        run = p.add_run(line)
        set_run_font(run, size=size, bold=bold, color=color, font=font)
    set_cell_margins(cell)


def add_runs(paragraph, parts, size=11, color=DARK):
    for part in parts:
        if isinstance(part, str):
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)
        else:
            text, bold, *rest = part
            c = rest[0] if rest else color
            run = paragraph.add_run(text)
            set_run_font(run, size=size, bold=bold, color=c)


class Notes:
    def __init__(self):
        self.doc = Document()
        self._setup()

    def _setup(self):
        sec = self.doc.sections[0]
        sec.page_width = Mm(210)
        sec.page_height = Mm(297)
        sec.left_margin = Mm(16)
        sec.right_margin = Mm(16)
        sec.top_margin = Mm(16)
        sec.bottom_margin = Mm(16)
        sec.header_distance = Mm(8)
        sec.footer_distance = Mm(8)

        normal = self.doc.styles["Normal"]
        normal.font.name = KR_FONT
        normal.font.size = Pt(11)
        normal.font.color.rgb = DARK
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), KR_FONT)
        pf = normal.paragraph_format
        pf.space_after = Pt(6)
        pf.space_before = Pt(0)
        pf.line_spacing = 1.18

        header = sec.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = hp.add_run("8/19 주요 시장 코멘트  ·  매크로 · 하이닉스 · 파운드리 · 소부장")
        set_run_font(r, size=8.5, color=GRAY)

        footer = sec.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run("영상 녹화용 정리  ·  숫자·시나리오는 공개 코멘트 기준  ·  ")
        set_run_font(r, size=8, color=GRAY)
        fld = parse_xml(
            f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE ">'
            f'<w:r><w:rPr><w:sz w:val="16"/><w:color w:val="4B5563"/>'
            f'<w:rFonts w:ascii="{KR_FONT}" w:hAnsi="{KR_FONT}" w:eastAsia="{KR_FONT}"/></w:rPr>'
            f"<w:t></w:t></w:r></w:fldSimple>"
        )
        fp._p.append(fld)

        core = self.doc.core_properties
        core.title = "8월 19일 주요 시장 코멘트 (매크로·하이닉스 환원·파운드리·소부장)"
        core.author = "준혁"
        core.subject = "매크로, SK하이닉스 주주환원, 삼성 파운드리, 메모리 논쟁, 이수페타시스, 기가비스"

    def p(self, text, size=11, bold=False, color=DARK, space_after=6, space_before=0, align="left"):
        para = self.doc.add_paragraph()
        para.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }[align]
        para.paragraph_format.space_after = Pt(space_after)
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.line_spacing = 1.18
        run = para.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
        return para

    def rich(self, parts, size=11, space_after=6, space_before=0):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_after = Pt(space_after)
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.line_spacing = 1.18
        add_runs(para, parts, size=size)
        return para

    def h1(self, text, num=None):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing = 1.1
        if num:
            run = para.add_run(f"{num}  ")
            set_run_font(run, size=16, bold=True, color=GOLD)
        run = para.add_run(text)
        set_run_font(run, size=16, bold=True, color=NAVY)
        pPr = para._p.get_or_add_pPr()
        pPr.append(
            parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="12" w:space="4" w:color="{NAVY_HEX}"/>'
                f"</w:pBdr>"
            )
        )
        return para

    def h2(self, text):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(text)
        set_run_font(run, size=13, bold=True, color=NAVY2)
        return para

    def h3(self, text):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(3)
        run = para.add_run(text)
        set_run_font(run, size=11.5, bold=True, color=NAVY)
        return para

    def bullet(self, text, level=0, bold_lead=None, size=11):
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.55 + level * 0.45)
        para.paragraph_format.first_line_indent = Cm(-0.35)
        para.paragraph_format.space_after = Pt(2.5)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.line_spacing = 1.15
        mark = "• " if level == 0 else "– "
        run = para.add_run(mark)
        set_run_font(run, size=size, color=NAVY2 if level == 0 else GRAY)
        if bold_lead:
            run = para.add_run(bold_lead)
            set_run_font(run, size=size, bold=True, color=DARK)
            run = para.add_run(text)
            set_run_font(run, size=size, color=DARK)
        else:
            run = para.add_run(text)
            set_run_font(run, size=size, color=DARK)
        return para

    def flow(self, items, size=11):
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing = 1.2
        for i, item in enumerate(items):
            if i:
                run = para.add_run("   →   ")
                set_run_font(run, size=size, bold=True, color=GOLD)
            run = para.add_run(item)
            set_run_font(run, size=size, bold=True, color=NAVY)
        return para

    def callout(self, title, body, kind="key"):
        palette = {
            "key": (NAVY_HEX, LIGHT_HEX, NAVY),
            "bull": ("166534", GREEN_HEX, GREEN),
            "bear": ("991B1B", RED_HEX, RED),
            "note": (GOLD_HEX, AMBER_HEX, AMBER),
            "blue": (NAVY2_HEX, BLUE_HEX, NAVY2),
        }
        accent, fill, title_color = palette[kind]
        table = self.doc.add_table(rows=1, cols=1)
        table.autofit = True
        table.allow_autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        shade_cell(cell, fill)
        set_left_accent(cell, accent, sz="28")
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        p1.paragraph_format.space_before = Pt(0)
        r = p1.add_run(title)
        set_run_font(r, size=10, bold=True, color=title_color)
        if isinstance(body, str):
            body = [body]
        for line in body:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(line)
            set_run_font(r, size=10.5, color=DARK)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def table(self, headers, rows, col_widths=None, first_col_bold=True):
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        set_table_borders(table, color="D5DCE6", sz="4")
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            shade_cell(cell, NAVY_HEX)
            cell_text(cell, h, size=9.5, bold=True, color=WHITE, align="center")
        prevent_row_split(table.rows[0])
        for r_i, row in enumerate(rows):
            for c_i, val in enumerate(row):
                cell = table.rows[r_i + 1].cells[c_i]
                if r_i % 2 == 1:
                    shade_cell(cell, ROW_HEX)
                else:
                    shade_cell(cell, WHITE_HEX)
                align = "left" if c_i == 0 else "center"
                bold = first_col_bold and c_i == 0
                cell_text(cell, str(val), size=9.5, bold=bold, color=DARK, align=align)
            prevent_row_split(table.rows[r_i + 1])
        if col_widths:
            for row in table.rows:
                for i, w in enumerate(col_widths):
                    row.cells[i].width = Cm(w)
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(8)
        spacer.paragraph_format.space_before = Pt(2)
        return table

    def spacer(self, pt=4):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(pt)
        p.paragraph_format.space_before = Pt(0)

    def save(self, path: Path):
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))


def build():
    n = Notes()

    # ── Cover ──────────────────────────────────────────────
    n.p("2026. 8. 19. 강의노트  ·  영상 녹화 / 오후–저녁 업로드", size=10.5, color=GRAY, align="center", space_after=4)
    n.p("주요 시장 코멘트", size=13, bold=True, color=GOLD, align="center", space_after=2)
    n.p("8월 19일  주요 시장 코멘트", size=22, bold=True, color=NAVY, align="center", space_after=2)
    n.p("(매크로 · 하이닉스 환원 · 파운드리 · 소부장)", size=15, bold=True, color=NAVY2, align="center", space_after=8)
    n.p(
        "유가·금리  ·  SK하이닉스 40조·FCF 50%+  ·  삼성 파운드리 +15%  ·  이수페타시스 · 기가비스",
        size=11,
        color=GRAY,
        align="center",
        space_after=10,
    )

    n.callout(
        "오늘 한 장으로 보면",
        [
            "단기 무게중심은 매크로다. 미·이란 불확실 → 유가 → 금리 → 고PER 압박. AI 수요가 하루아침에 꺾인 이야기가 아니다.",
            "펀더멘탈 축은 SK하이닉스 주주환원(40조 + FCF 50% 초과)과 삼성 파운드리 판가 인상. 소부장은 Multi-Lam·검사장비 수주.",
            "메모리 비관론(Wood/Thompson)은 ‘대체 유인’ 논리로 읽되, 메모리 vs 비메모리 헤게모니 싸움으로 몰아가진 않는다.",
        ],
        kind="key",
    )

    n.h2("강의 구성")
    n.table(
        ["파트", "대상", "오늘 확인할 것"],
        [
            ["1", "매크로", "10Y 4.7%↓ vs 5% · Brent ~90 vs 100+ · 엔캐리 여부"],
            ["2", "SK하이닉스", "40조 소각 + 25–27 FCF 50%+ · 3Q26 추가 규모"],
            ["3", "삼성 파운드리", "SF4 최대 +15% · 흑자전환·고객 파이프라인"],
            ["4", "메모리 논쟁", "가격↑→대체 유인 · Hormuz 비유 · 프레이밍 경계"],
            ["5", "이수페타시스", "Multi-Lam 비중 · Capa · 판가 +15%"],
            ["6", "기가비스", "AOI/AOR · 89.5억 수주 · 27년 성장"],
            ["보론", "주변 테마", "앤트로픽 · 유니트리 · 100조 美투자 · AI 비중"],
        ],
        col_widths=[2.2, 4.0, 11.4],
    )

    n.h2("오프닝 멘트 (녹화용)")
    n.p(
        "오늘은 매크로가 먼저입니다. 전쟁이 유가와 금리를 건드리면, AI CAPEX 스토리와 무관하게 밸류에이션이 먼저 맞습니다. "
        "그 위에서 SK하이닉스 주주환원, 삼성 파운드리 판가, 그리고 이수페타시스·기가비스를 보겠습니다. "
        "마지막에 Cathy Wood·Ben Thompson의 ‘왜 메모리주를 사지 않는가’ 논리를 다루되, 진영 싸움으로 읽지는 않겠습니다."
    )

    # ── 0. Cheat sheet ─────────────────────────────────────
    n.h1("오늘 숫자 한 장", num="0.")
    n.table(
        ["항목", "핵심 숫자", "한 줄 결론"],
        [
            ["매크로", "10Y 4.7%↓ vs 5%\nBrent ~$90 vs $100+\nUSD/JPY 157~159", "유가발 금리 쇼크 성격"],
            ["SK하이닉스", "40조 매입·소각 (시총 3.3%)\n25–27 FCF 50% 이상", "추가 환원 가시성 → 3Q26"],
            ["FCF 결론", "385조 → 192.5조\n−40조 = +152.5조 이상", "기간 누적 환원. 일시지급 아님"],
            ["삼성 파운드리", "첨단 최대 +15%\nSF4·SF5·8nm", "협상력↑ · 흑자전환 기대"],
            ["이수페타시스", "Multi-Lam 7→11%(잔고 20%+)\nOPM 20.3% · 판가 +15%", "Capa보다 이익 레버리지"],
            ["기가비스", "수주 89.5억 (매출 17.1%)\n컨센 TP ~19만원", "26년 비쌈 · 27년 성장·분할"],
        ],
        col_widths=[3.2, 7.0, 7.4],
    )

    # ── 1. Macro ───────────────────────────────────────────
    n.h1("매크로 — 전쟁발 할인율", num="1.")
    n.callout(
        "한 줄 결론",
        [
            "핵심 악재는 전쟁 자체보다 ‘금리’다.",
            "AI 수요가 갑자기 약해졌다(X) → 금리 상승으로 AI/반도체 고PER이 압박받았다(O)에 가깝다.",
            "메모리는 원래 사이클 특성을 반영한 낮은 배수였다는 점도 같이 본다.",
        ],
        kind="key",
    )

    n.h2("1) 인과 사슬")
    n.flow(
        [
            "미·이란 불확실",
            "호르무즈·유가↑",
            "인플레 우려",
            "미 국채금리↑",
            "고PER 압박",
        ],
        size=10.5,
    )
    n.p(
        "매크로의 역습이 펀더멘탈을 압도한 구간입니다. 펀더멘탈이 없어도 AI CAPEX가 계속돼도, "
        "전쟁발 매크로 vs 투자의지 vs 토큰 수요의 무게중심이 왼쪽으로 치우치면 위험자산이 먼저 흔들립니다."
    )

    n.h2("2) 모니터할 레벨")
    n.table(
        ["변수", "안정·완화", "부담·악화"],
        [
            ["브렌트유", "$90 전후 안정", "$100 이상 → 인플레·금리 악순환"],
            ["미 10년물", "4.7% 이하 안정", "5% 돌파·고착 → 위험자산 회피"],
            ["미 30년물", "장중 고점 후 하락 반전 주목\n(예: 5.34% → 5.285%)", "장기금리 재급등"],
            ["USD/JPY", "157~159 반등·현수지", "159→155→150 급락 시 엔캐리 청산 의심"],
            ["Fed", "기준금리 동결이 주식에 유리", "고용 둔화 인하 기대 vs 이란전 장기화"],
        ],
        col_widths=[3.0, 7.0, 7.6],
    )
    n.callout(
        "현실 인식",
        [
            "이란발 금리 이슈가 안정될 때까지는 채권시장 사람들의 주장에 힘이 실릴 가능성이 큽니다.",
            "그래도 토큰 수요 확대 → AI 투자 GO의 근본 궤도가, 웬만한 금리 상승만으로 바뀌긴 어렵다고 봅니다.",
        ],
        kind="note",
    )

    n.h2("3) 엔캐리 — 2024년 8월형인가")
    n.p("비교 기준은 2024년 7/31 → 8/5입니다.")
    n.table(
        ["지표", "당시 움직임", "의미"],
        [
            ["Nikkei", "39,102 → 31,458 (약 −19.5%)", "레버리지·포지션 청산 성격"],
            ["KOSPI", "2,771 → 2,442 (약 −11.9%)", "동반 유동성 쇼크"],
            ["USD/JPY", "152~153 → 142~145 (약 6% 엔화 강세)", "캐리 청산 경로"],
            ["익일", "Nikkei +10.23% / KOSPI +3.30%", "펀더멘털 하루 붕괴보다 유동성 쇼크 증거"],
        ],
        col_widths=[2.8, 7.2, 7.6],
    )
    n.callout(
        "현재 해석",
        [
            "USD/JPY가 157~159에서 오히려 반등하는 상태라면, 일본 장기금리 상승의 1차 충격은",
            "엔캐리 청산보다 글로벌 채권금리 상승·밸류에이션 압박 쪽으로 보는 편이 적절합니다.",
            "반대로 159→155→150으로 빠르게 내려가면 그때는 2024년 8월형 청산을 의심할 단계입니다.",
        ],
        kind="blue",
    )

    n.h2("4) FT 이란 보도 — 정확한 프레이밍")
    n.p(
        "Financial Times(2026-08-19): Iran eyes military targets in Europe if Donald Trump escalates war, insiders say."
    )
    n.bullet("트럼프가 추가 확전할 경우, 이란 내부에서 유럽 내 미군 시설 공격 옵션을 검토한다는 내용.")
    n.bullet("불가리아·키프로스 등 남동부 유럽, Bezmer 공군기지, 호르무즈 해저 인프라가 거론.")
    n.bullet("장거리 미사일 능력에는 한계가 있다는 평가도 함께 실림.")
    n.callout(
        "오독 금지",
        [
            "“이란이 유럽을 공격하기로 결정했다”가 아닙니다.",
            "정확히는 추가 확전 시를 대비한 내부 옵션 검토입니다.",
        ],
        kind="bear",
    )

    # ── 2. SK Hynix ────────────────────────────────────────
    n.h1("SK하이닉스 주주환원", num="2.")
    n.callout(
        "한 줄 결론",
        [
            "40조원 자사주 매입·소각은 시작이고, 본게임은 2025~2027 누적 FCF의 50% 초과 환원입니다.",
            "정확한 결론: 누적 FCF ≈385조 → 최소 192.5조 이상 → 이미 40조 → 추가 약 152.5조 이상.",
            "구체적 추가 규모·방식은 3Q26 실적발표에서 안내 예정입니다.",
        ],
        kind="key",
    )

    n.h2("1) 공시 요약 — 40조 매입·소각")
    n.table(
        ["항목", "내용"],
        [
            ["규모", "40조원 자기주식 취득 후 소각 (현 주가 기준 총주식수 대비 약 3.3%)"],
            ["기간", "2026년 8월 20일 ~ 11월 19일 (약 3개월, 62 영업일 전망)"],
            ["일 매입", "약 6,452억원/일  (6,452×62 ≈ 40.0조)"],
            ["목적", "내재가치 대비 저평가 판단 → 자본 재배치·주주가치 제고"],
            ["구조", "ADR 발행으로 희석된 지분율 → 매입·소각 → SK스퀘어 지분율을 ADR 이전 수준으로 복원"],
            ["공시", "DART rcpNo=20260819000254"],
        ],
        col_widths=[3.2, 14.4],
    )

    n.h2("2) 주주환원 정책 (2025~2027)")
    n.bullet("재무건전성 목표 달성을 전제로 하며, 목표 달성 경로는 계획대로 진행 중이라고 판단.")
    n.bullet("정책 기간 누적 Free Cash Flow의 ‘50% 이상’을 환원. (회사는 50% 초과 목표로 상향)")
    n.bullet("자기주식 취득·소각과 현금 배당 병행. 고정배당·특별배당 등 배당 확대 방안도 검토.")
    n.bullet("유의미한 FCF 창출에 따라 조기 환원. 향후에도 정책 기간 내 자사주 취득·소각을 지속할 계획.")
    n.bullet("추가 환원 규모와 방식은 이사회 결의 후 3분기 실적발표 시점에 안내 예정.")

    n.h2("3) 정확한 결론 — 숫자")
    n.callout(
        "메인 산식 (공개 코멘트 기준)",
        [
            "2025~2027 누적 FCF ≈ 385조원",
            "385조 × 50% = 192.5조원 이상을 프로그램 기간 동안 환원",
            "이미 발표한 40조 자사주 매입·소각 → 추가 약 152.5조원 이상 필요(50% 기준)",
            "회사는 50%가 아니라 ‘50% 초과’ + 특별배당 등도 열어 둠",
        ],
        kind="bull",
    )

    n.h2("4) 오해 방지 (필수)")
    n.callout(
        "192.5조는 2027년에 한 번에 주는 돈이 아니다",
        [
            "2025~2027년 프로그램 기간 동안 누적으로 환원한다는 의미입니다.",
            "이미 40조가 집행 구간에 들어가므로, 나머지는 기간 내 배당·추가 자사주 등으로 채워집니다.",
        ],
        kind="note",
    )
    n.callout(
        "표의 2028년 102.5조 ≠ 회사 정책",
        [
            "FCF 모델에 50%를 적용해 본 참고치일 뿐입니다.",
            "2028년 주주환원은 별도 정책이 새로 나와야 합니다.",
            "오늘 발표만으로 2028년에도 FCF의 50%를 환원한다고 볼 수 없습니다.",
        ],
        kind="bear",
    )

    n.h2("5) 내부 FCF 래더 — 참고 모델")
    n.p(
        "업데이트 버전의 연간 FCF 추산입니다. 메인 결론(385조)과 병존하니, 강의에서는 ‘참고 모델’로만 씁니다."
    )
    n.table(
        ["구분", "Y1", "Y2", "Y3", "3년 합"],
        [
            ["기존 계산", "179조", "242조", "237조", "658조"],
            ["보수 (WC·기타 20~30조 차감)", "150조", "210조", "205조", "565조"],
        ],
        col_widths=[5.6, 2.8, 2.8, 2.8, 3.6],
    )
    n.p(
        "보수 래더라도 3년 누적 565조의 50%면 약 282조 수준입니다. "
        "메인으로 쓰는 385→192.5는 더 보수적으로 잡아 둔 ‘정확한 결론’ 프레임입니다."
    )

    n.h2("6) Peer buyback — 과대해석 금지")
    n.table(
        ["사례", "내용", "읽기"],
        [
            ["키옥시아", "8/3–8/10 약 8,000억엔 매입\n7/31 46,500 → 8/10 48,010 (+3.2%)\n8/18 49,950 누적 +7.4%", "시황 나쁠 때는 소폭. 그래도 매입 비중 이상은 상승"],
            ["샌디스크", "자사주 매입한도 확대", "한도 확대 뉴스 자체"],
            ["하이닉스 맥락", "8/5~8/7 주가 −15.1%\n이후 Investor Day(8/13) 재평가 + Buyback", "$140억급 추가 승인만으로 +20%를 설명하기 어려움"],
        ],
        col_widths=[3.0, 8.0, 6.6],
    )
    n.callout(
        "하이닉스 주가 해석",
        [
            "8/13 Investor Day 이후 장기 성장률·마진 전망 재평가가 본격 반등의 축.",
            "Buyback은 그 위에 EPS 상승 레버리지를 추가한 요인으로 보는 것이 적절합니다.",
            "직관적으로 규모 확대 반영에 따라 +5~9% 구간을 가늠해 볼 수 있다는 코멘트가 있었습니다.",
        ],
        kind="blue",
    )

    n.h2("7) 확인 포인트")
    n.bullet("3Q26: 추가 환원 규모·방식(추가 자사주 / 현금·특별배당)이 숫자로 나오는가.")
    n.bullet("누적 FCF가 385조 프레임에 실제로 수렴하는가. (CAPEX·운전자본)")
    n.bullet("‘50% 초과’가 형식적 문구가 아니라 실행으로 이어지는가.")
    n.bullet("회사 안팎의 자금 수요(투자·파이)와 환원 속도의 균형.")

    # ── 3. Samsung foundry ─────────────────────────────────
    n.h1("삼성전자 파운드리 가격 인상", num="3.")
    n.callout(
        "한 줄 결론",
        [
            "AI 반도체 수요 급증 + TSMC 포화로 삼성 파운드리 협상력이 올라갔고, 신규 주문 기준 최대 15% 인상.",
            "원문 톤: 삼성전자에 긍정적 소식.",
        ],
        kind="bull",
    )

    n.h2("1) 공정별 인상")
    n.table(
        ["공정", "인상폭", "비고"],
        [
            ["4나노 (SF4)", "중국·미국 고객 10~15%\n대만 고객 5~10%", "평택 SF4: 퀄컴 + HBM 베이스다이로 풀가동"],
            ["5나노 (SF5)", "웨이퍼 기준 10~15%", "첨단 공정"],
            ["8나노", "레거시 약 10%", "레거시도 동반 인상"],
        ],
        col_widths=[3.4, 6.0, 8.2],
    )
    n.p("출처: Reuters, 2026-08-19 — Samsung hikes chipmaking prices by up to 15%.")

    n.h2("2) 왜 지금인가")
    n.bullet("TSMC 첨단 용량 포화(점유율 70%+) → 고객이 삼성으로 분산 → 가격 협상력 강화.")
    n.bullet("미 장비 규제로 중국 팹리스의 해외 파운드리 의존 심화.")
    n.bullet("평택 SF4 라인 작년 말부터 완전 가동.")

    n.h2("3) 사업 전망")
    n.bullet("2022년 이후 적자 파운드리 → 단가·수율·가동률로 이르면 내년 흑자 전환 기대.")
    n.bullet("올해 첨단 공정 매출 비중 절반 이상. AI·HPC 비중 30% 이상(2025년 말 15~20% → 확대).")
    n.bullet("고객: 테슬라·애플·브로드컴 계약, 엔비디아 신규 AI 추론칩 수주, 구글 4나노 협의 중.")

    # ── 4. Memory debate ───────────────────────────────────
    n.h1("왜 메모리주를 사지 않는가 — 논리와 경계", num="4.")
    n.callout(
        "한 줄 결론",
        [
            "Cathy Wood·Ben Thompson의 요지: 메모리 가격 급등 자체가 장기적으로는 ‘대체 유인’을 키운다.",
            "핵심은 “메모리 고점 = 수요 고점”이 아니라, 비쌀수록 AI 업계가 메모리 의존도를 낮출 경제적 유인이 커진다는 점.",
            "다만 이 시각을 미국 큰손의 헤게모니로 ‘메모리 vs 비메모리’ 싸움까지 몰아가는 프레이밍은 부담스럽다는 톤을 유지합니다.",
        ],
        kind="key",
    )

    n.h2("1) 캐시 우드 — 가격 급등은 부정 신호일 수 있다")
    n.flow(["HBM 가격 급등", "AI 원가 상승", "HBM 의존도 축소", "대체 기술 투자 가속"], size=10.5)
    n.bullet("단기간 수배 상승을 호재로만 보지 않음.")
    n.bullet("가격이 높을수록 AI 업체는 ① SRAM 확대 ② 메모리 압축 ③ 모델 경량화 ④ ASIC 개발을 추진.")
    n.bullet("Cerebras·Groq 등 추론칩 업체는 엔지니어링 단계에서 HBM 의존을 낮추거나 우회하려는 움직임.")

    n.h2("2) 벤 톰슨 — Hormuz 비유")
    n.p(
        "HBM 공급자를 이란·호르무즈 해협에 비유합니다. 특정 자원을 통제해 가격을 올리면 단기 이익은 가능하지만, "
        "상대는 장기적으로 다른 공급망과 기술을 찾아 의존도를 낮춥니다."
    )
    n.callout(
        "함의",
        [
            "메모리 가격 상승이 영원히 메모리 업체의 가격결정력을 보장하지 않는다.",
            "현재 수요 폭증은 실적·주가 호재이지만, 과도한 가격 지속은 대체·효율화 투자를 자극할 수 있다.",
        ],
        kind="note",
    )

    n.h2("3) 프레이밍 경계")
    n.callout(
        "작성자 톤",
        [
            "이런 시각이 미국 큰손·영향력 있는 사람들을 통해 퍼진다면,",
            "일종의 헤게모니 싸움(메모리 vs 비메모리)으로까지 엮어 그 방향으로 모는 것은 부담스럽습니다.",
            "강의에서는 ‘대체 유인’ 논리로만 정리하고, 섹터 전쟁 서사는 채택하지 않습니다.",
        ],
        kind="bear",
    )
    n.p("출처 요약: 월가견문(wallstreetcn) — 为什么不买存储股？这是两位大佬的回答")

    # ── 5. Isu Petasys ─────────────────────────────────────
    n.h1("이수페타시스", num="5.")
    n.callout(
        "한 줄 결론",
        [
            "이제 Capa 증설만 보는 회사가 아니라, AI용 Multi-Lam 비중 상승에 따른 이익 레버리지를 봐야 한다.",
            "주가 조정 = 실적 상향 국면의 매수 기회라는 논리가 가능. 상대적으로 편안한 인식 가능성도 언급.",
        ],
        kind="key",
    )

    n.h2("1) 2Q26 Review")
    n.table(
        ["항목", "숫자", "비고"],
        [
            ["매출", "3,799억원 · YoY +57.4%", "컨센 +4.9% 상회"],
            ["영업이익", "771억원 · YoY +83.3%", "컨센 +2.7% 상회"],
            ["OPM", "20.3%", "본사 Capa + AI 고수익 + Multi-Lam 믹스"],
        ],
        col_widths=[3.2, 6.4, 8.0],
    )

    n.h2("2) Multi-Lam이란")
    n.p(
        "여러 겹의 회로층을 쌓은 고다층 PCB입니다. AI 서버·가속기·고속 스위치는 더 복잡하고 많은 신호를 처리해야 해서 "
        "고다층이 필요합니다. Multi-Lam 비중↑ → 단가↑ · 난이도↑ · 진입장벽↑ · 수익성↑."
    )
    n.flow(["1Q 비중 7%", "2Q 비중 11%", "수주잔고 20%+", "4Q G사 전환·M사 ASIC"], size=10.5)

    n.h2("3) Capa · 판가 · 이익")
    n.table(
        ["시점", "월 매출 Capa"],
        [
            ["현재", "약 1,200억원"],
            ["2027년 2Q", "1,500억원"],
            ["2028년 하반기", "1,800억원"],
        ],
        col_widths=[6.0, 11.6],
    )
    n.bullet("원재료 반영 판가 협상 → 평균 약 +15% 효과, 하반기부터.")
    n.bullet("① Capa 확대 + ② Multi-Lam 비중 + ③ 수율 개선 + ④ 판가 인상 → 양적·질적 성장 동시.")
    n.bullet("2027 영업이익 컨센 대비 +10% 전후 상향여지 — 시장이 Multi-Lam 레버리지를 과소평가했을 가능성.")
    n.bullet("글로벌 Peer 하락으로 Target Multiple 30.9배 → 26.0배, 그러나 이익상향으로 TP 하향폭은 제한적일 듯.")

    # ── 6. Gigavis ─────────────────────────────────────────
    n.h1("기가비스", num="6.")
    n.callout(
        "한 줄 결론",
        [
            "고사양 FC-BGA 기판의 ‘눈(AOI)’과 ‘레이저 수리공(AOR)’을 만드는 Top-tier 검사·수리 장비.",
            "연결고리: AI → FC-BGA 고사양화·증설 → 기가비스 수주 → 매출·영업 레버리지.",
        ],
        kind="key",
    )

    n.h2("1) 사업 한 장")
    n.bullet("2004년 설립, 2023년 코스닥 상장. FC-BGA(ABF) 기판용 AOI·AOR.")
    n.bullet("고객: 이비덴, 신코, 유니마이크론, 삼성전기 등 글로벌 기판업체.")
    n.bullet("미세 회로 결함 검출 + 레이저 수리. FC-BGA 검사·수리 장비 글로벌 Top-tier.")

    n.h2("2) 단일판매·공급계약")
    n.table(
        ["항목", "내용"],
        [
            ["상대", "일본 반도체 기판 제조회사"],
            ["내용", "반도체 기판 검사 및 수리장비 외"],
            ["금액", "89.5억원 (매출대비 17.1%)"],
            ["기간", "2026-08-18 ~ 2027-09-30"],
            ["지역", "해외"],
        ],
        col_widths=[3.2, 14.4],
    )

    n.h2("3) 실적 레버리지 · 접근")
    n.table(
        ["연도", "매출", "영업이익"],
        [
            ["2025 (집계)", "약 847억원", "약 121억원"],
            ["2026 (메리츠)", "1,835억원", "721억원"],
        ],
        col_widths=[4.0, 6.4, 7.2],
    )
    n.bullet("증권사 컨센 TP 약 19만원선. 적정 안전마진 고려 시 분할접근 후보.")
    n.bullet("26년 실적 기준으로는 비싸고, 27년 성장 잠재력이 초점.")

    # ── 7. Boron ───────────────────────────────────────────
    n.h1("보론", num="7.")

    n.h2("1) 앤트로픽")
    n.p("올 매출 최대 1,200억 달러 전망, 기업가치 2조 달러 예측도 제기. (국내 언론 요약)")

    n.h2("2) 유니트리 IPO")
    n.bullet("상하이 커촹반 상장. 공모가 컨센(104위안) 대비 +45% 확정.")
    n.bullet("Hyperliquid 무기한 선물 ~100달러 거래 → 기업가치 ~405억 달러로, 상장 첫날 4배+ 폭등 베팅 해석.")
    n.bullet("커촹반 첫날 3~5배 관행을 감안하면 시총 급등 시나리오도 나오지만, 휴머노이드 상용화 초기라 장기 흥행은 신중론.")
    n.bullet("1분기 순이익은 R&D 등으로 YoY −47.69%.")

    n.h2("3) 100조 美 메가투자")
    n.p("삼성·SK·현대차그룹, 100조 규모 미국 메가투자 속도. (국내 언론)")

    n.h2("4) AI 비중 · 선호 순서")
    n.bullet("글로벌 투자자가 AI 비중을 임계점 전에 줄이기는 어렵다고 봄.")
    n.bullet("선호/순서 감각: 빅테크 → 파운드리 → 메모리 / 미국·일본·한국.")
    n.bullet("국내기관: 대형주 exposure와 변압기·소부장 믹스 비율이 고민의 중심.")

    # ── 8. Close ───────────────────────────────────────────
    n.h1("클로징", num="8.")
    n.callout(
        "오늘 강의에서 가져갈 세 문장",
        [
            "1) 단기: 유가·10년물·엔화. 4.7% 이하 안정이면 숨이 트이고, 5% 고착이면 밸류 압박이 계속된다.",
            "2) 하이닉스: 40조는 시작. 385→192.5→추가 152.5조+ 프레임과 3Q26 추가 환원을 본다. 2028 숫자는 정책이 아니다.",
            "3) 소부장: 이수페타시스는 Multi-Lam 이익 레버리지, 기가비스는 FC-BGA 검사·수리 수주. 삼전 파운드리 +15%는 협상력 회복 신호.",
        ],
        kind="key",
    )

    n.h2("클로징 멘트 (녹화용)")
    n.p(
        "모든 것은 돌고 돌아 전쟁으로 연결되고, 전쟁이 유가와 금리를 건드리면 AI 스토리와 무관하게 할인율이 먼저 움직입니다. "
        "그래도 토큰 수요가 살아있는 한 CAPEX의 근본 궤도가 쉽게 꺾이진 않는다고 봅니다. "
        "그 위에서 오늘은 하이닉스 환원 실행, 삼성 파운드리 판가, 그리고 Multi-Lam·검사장비의 이익 레버리지를 확인했습니다. "
        "메모리 비관론은 대체 유인으로만 읽고, 섹터 헤게모니 싸움으로는 몰아가지 않겠습니다."
    )

    n.spacer(8)
    n.p(
        "— 8월 19일 주요 시장 코멘트 강의노트. 원문 퀵코멘트(06:20~00:14)를 강의 순서로 재구성.",
        size=9.5,
        color=GRAY,
        align="right",
    )

    n.save(OUT_PATH)
    print(f"Wrote {OUT_PATH} ({OUT_PATH.stat().st_size} bytes)")


if __name__ == "__main__":
    build()
