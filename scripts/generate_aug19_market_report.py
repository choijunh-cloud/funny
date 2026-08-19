#!/usr/bin/env python3
"""2026년 8월 19일 퀵 코멘트를 시각화한 시장전략 보고서(.docx) 생성."""

from __future__ import annotations

from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from generate_aug18_lecture import (
    AMBER,
    DARK,
    GOLD,
    GRAY,
    GREEN,
    NAVY,
    NAVY2,
    RED,
    Notes,
    set_run_font,
)

ROOT = Path("/workspace")
OUT_PATH = ROOT / "lectures/8월 19일 시장전략 보고서 (매크로·반도체·주주환원).docx"
CHART_DIR = Path("/tmp/aug19_market_report_charts")

FONT = "WenQuanYi Micro Hei"
BLUE = "#1E407C"
NAVY_HEX = "#0F2043"
GOLD_HEX = "#B8943A"
RED_HEX = "#B42318"
GREEN_HEX = "#16803A"
GRAY_HEX = "#667085"
LIGHT = "#EEF2F8"


def setup_charts() -> None:
    CHART_DIR.mkdir(parents=True, exist_ok=True)
    plt.rcParams.update(
        {
            "font.family": FONT,
            "axes.unicode_minus": False,
            "axes.titleweight": "bold",
            "axes.titlesize": 15,
            "axes.labelsize": 10,
            "xtick.labelsize": 9,
            "ytick.labelsize": 9,
            "figure.facecolor": "white",
            "axes.facecolor": "white",
            "axes.edgecolor": "#D0D5DD",
            "axes.spines.top": False,
            "axes.spines.right": False,
        }
    )


def savefig(fig: plt.Figure, name: str) -> Path:
    path = CHART_DIR / name
    fig.savefig(path, dpi=190, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def add_chart(n: Notes, path: Path, width=17.0, caption: str | None = None) -> None:
    p = n.doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Cm(width))
    if caption:
        cp = n.doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(7)
        r = cp.add_run(caption)
        set_run_font(r, size=8.5, color=GRAY)


def page_break(n: Notes) -> None:
    n.doc.add_page_break()


def chart_macro_thresholds() -> Path:
    fig, ax = plt.subplots(figsize=(9.3, 4.2))
    ax.set_xlim(0, 100)
    ax.set_ylim(-0.7, 3.7)
    rows = [
        ("브렌트유", 54, "$90 전후 안정", "$100+ 악순환 위험"),
        ("미 10년물", 47, "4.7% 이하 완화", "5.0%+ 위험회피"),
        ("USD/JPY", 58, "157~159: 금리 충격", "150 초반↓: 캐리 청산"),
        ("달러-원", 42, "1,400 하회", "1,340~1,360 조건부"),
    ]
    for i, (label, marker, good, bad) in enumerate(rows[::-1]):
        ax.barh(i, 56, left=0, color="#DDF4E4", height=0.56)
        ax.barh(i, 22, left=56, color="#FFF1CC", height=0.56)
        ax.barh(i, 22, left=78, color="#FDE2E0", height=0.56)
        ax.scatter(marker, i, s=120, color=NAVY_HEX, zorder=3, edgecolor="white", linewidth=1.5)
        ax.text(-1.5, i, label, va="center", ha="right", fontweight="bold", color=NAVY_HEX)
        ax.text(4, i - 0.37, good, fontsize=8, color=GREEN_HEX)
        ax.text(79, i - 0.37, bad, fontsize=8, color=RED_HEX)
    ax.set_yticks([])
    ax.set_xticks([])
    ax.set_title("단기 시장 대시보드: 숫자보다 ‘구간 전환’을 본다", pad=14, color=NAVY_HEX)
    ax.text(0, 3.48, "완화/우호", color=GREEN_HEX, fontweight="bold")
    ax.text(58, 3.48, "경계", color="#9A6700", fontweight="bold")
    ax.text(80, 3.48, "위험", color=RED_HEX, fontweight="bold")
    return savefig(fig, "01_macro_thresholds.png")


def chart_carry_shock() -> Path:
    labels = ["Nikkei", "KOSPI", "USD/JPY"]
    changes = [-19.5, -11.9, -6.5]
    colors = [RED_HEX, "#D0473E", GOLD_HEX]
    fig, ax = plt.subplots(figsize=(8.8, 4.2))
    bars = ax.barh(labels[::-1], changes[::-1], color=colors[::-1], height=0.58)
    ax.axvline(0, color="#98A2B3", lw=1)
    ax.set_xlim(-22, 1)
    ax.set_title("2024년 7/31→8/5: 주가 하락과 엔화 강세가 동시에 발생", color=NAVY_HEX)
    ax.set_xlabel("기간 등락률(%, 제공 코멘트 기준)")
    ax.grid(axis="x", color="#EAECF0", lw=0.8)
    for b, v in zip(bars, changes[::-1]):
        ax.text(v - 0.4, b.get_y() + b.get_height() / 2, f"{v:.1f}%", va="center", ha="right", color="white", fontweight="bold")
    ax.text(
        0.01,
        -0.28,
        "8/6 반등: Nikkei +10.23% · KOSPI +3.30% → 펀더멘털보다 레버리지/유동성 쇼크의 성격",
        transform=ax.transAxes,
        fontsize=9,
        color=GRAY_HEX,
    )
    return savefig(fig, "02_carry_shock.png")


def chart_fx_eps() -> Path:
    fx_change = (1420 / 1520 - 1) * 100
    names = ["삼성전자", "SK하이닉스"]
    eps = [fx_change * 0.4, fx_change * 0.9]
    fig, ax = plt.subplots(figsize=(8.8, 4.1))
    bars = ax.bar(names, eps, color=[BLUE, RED_HEX], width=0.55)
    ax.axhline(0, color="#98A2B3", lw=1)
    ax.set_ylim(-7, 0.5)
    ax.set_ylabel("EPS 변화 추정(%)")
    ax.set_title("달러-원 1,520→1,420(-6.6%) 가정의 EPS 민감도", color=NAVY_HEX)
    ax.grid(axis="y", color="#EAECF0", lw=0.8)
    for b, v in zip(bars, eps):
        ax.text(b.get_x() + b.get_width() / 2, v + 0.15, f"{v:.1f}%", ha="center", va="bottom", color="white", fontweight="bold")
    ax.text(
        0.5,
        -0.23,
        "민감도 가정: 원/달러 +1% → 삼성전자 EPS +0.4%, SK하이닉스 +0.9%",
        transform=ax.transAxes,
        ha="center",
        fontsize=9,
        color=GRAY_HEX,
    )
    return savefig(fig, "03_fx_eps.png")


def chart_shareholder_return() -> Path:
    labels = ["발행주식 감소", "기계적 EPS 증가", "매입 규모/시총"]
    values = [3.3, 3.4, 3.3]
    fig, ax = plt.subplots(figsize=(8.8, 4.0))
    bars = ax.barh(labels[::-1], values[::-1], color=[GOLD_HEX, BLUE, NAVY_HEX][::-1], height=0.55)
    ax.set_xlim(0, 4)
    ax.set_xlabel("%")
    ax.set_title("SK하이닉스 40조원 자사주 매입·소각의 1차 효과", color=NAVY_HEX)
    ax.grid(axis="x", color="#EAECF0", lw=0.8)
    for b, v in zip(bars, values[::-1]):
        ax.text(v + 0.06, b.get_y() + b.get_height() / 2, f"{v:.1f}%", va="center", fontweight="bold", color=NAVY_HEX)
    return savefig(fig, "04_shareholder_return.png")


def chart_fcf_reconciliation() -> Path:
    labels = ["누적 FCF 385조\n시나리오", "연도별 FCF\n150+210+205조"]
    fcf = np.array([385, 565])
    return_50 = fcf * 0.5
    extra_after_40 = return_50 - 40
    x = np.arange(2)
    fig, ax = plt.subplots(figsize=(9.2, 4.6))
    ax.bar(x, fcf, width=0.56, color="#DCE6F4", label="누적 FCF")
    ax.bar(x, return_50, width=0.56, color=BLUE, label="50% 환원 기준")
    ax.scatter(x, extra_after_40, s=90, color=GOLD_HEX, zorder=3, label="40조 제외 추가분")
    for i in range(2):
        ax.text(i, fcf[i] + 10, f"{fcf[i]:.0f}", ha="center", color=NAVY_HEX, fontweight="bold")
        ax.text(i, return_50[i] - 18, f"{return_50[i]:.1f}", ha="center", color="white", fontweight="bold")
        ax.text(i + 0.04, extra_after_40[i] + 10, f"추가 {extra_after_40[i]:.1f}", color="#7A5C12", fontweight="bold")
    ax.set_xticks(x, labels)
    ax.set_ylabel("조원")
    ax.set_title("FCF 코멘트의 두 산식은 서로 다른 시나리오다", color=NAVY_HEX)
    ax.legend(frameon=False, ncol=3, loc="upper left")
    ax.grid(axis="y", color="#EAECF0", lw=0.8)
    return savefig(fig, "05_fcf_reconciliation.png")


def chart_memory_balance() -> Path:
    cases = ["AI 연산 +50%\n효율 +20%", "AI 연산 +20%\n효율 +30%"]
    compute = [50, 20]
    efficiency = [20, 30]
    gap = np.array(compute) - np.array(efficiency)
    x = np.arange(2)
    fig, ax = plt.subplots(figsize=(9.0, 4.4))
    w = 0.25
    ax.bar(x - w / 2, compute, width=w, label="AI 연산 증가", color=BLUE)
    ax.bar(x + w / 2, efficiency, width=w, label="메모리 효율 개선", color=GOLD_HEX)
    for i, g in enumerate(gap):
        c = GREEN_HEX if g > 0 else RED_HEX
        ax.text(i, max(compute[i], efficiency[i]) + 4, f"방향성 갭 {g:+.0f}%p", ha="center", color=c, fontweight="bold")
    ax.set_xticks(x, cases)
    ax.set_ylim(0, 65)
    ax.set_ylabel("%")
    ax.set_title("HBM 수요의 핵심: 연산 증가율 − 효율 개선률", color=NAVY_HEX)
    ax.legend(frameon=False, ncol=2, loc="upper right")
    ax.grid(axis="y", color="#EAECF0", lw=0.8)
    return savefig(fig, "06_memory_balance.png")


def chart_market_moves() -> Path:
    names = ["Marvell", "Broadcom", "Moderna", "Merck", "SOX"]
    moves = [9.9, -4.6, 77.0, 12.9, -2.12]
    colors = [GREEN_HEX if v > 0 else RED_HEX for v in moves]
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(9.4, 4.2), gridspec_kw={"width_ratios": [1, 1.25]})
    for ax, idx, title in [(ax1, [0, 1, 4], "AI/반도체"), (ax2, [2, 3], "바이오")]:
        vals = [moves[i] for i in idx]
        nms = [names[i] for i in idx]
        bars = ax.bar(nms, vals, color=[colors[i] for i in idx], width=0.6)
        ax.axhline(0, color="#98A2B3", lw=1)
        ax.set_title(title, color=NAVY_HEX)
        ax.grid(axis="y", color="#EAECF0", lw=0.8)
        for b, v in zip(bars, vals):
            ax.text(b.get_x() + b.get_width() / 2, v + (1.5 if v >= 0 else -1.5), f"{v:+.1f}%", ha="center", va="bottom" if v >= 0 else "top", fontweight="bold")
    fig.suptitle("같은 날의 가격 신호: ASIC 재편 + 헬스케어 로테이션", color=NAVY_HEX, fontweight="bold")
    return savefig(fig, "07_market_moves.png")


def chart_foundry_hikes() -> Path:
    names = ["4nm\n미·중", "4nm\n대만", "5nm", "8nm"]
    low = [10, 5, 10, 10]
    high = [15, 10, 15, 10]
    midpoint = [(a + b) / 2 for a, b in zip(low, high)]
    err = [[m - a for m, a in zip(midpoint, low)], [b - m for m, b in zip(midpoint, high)]]
    fig, ax = plt.subplots(figsize=(8.8, 4.2))
    ax.bar(names, midpoint, yerr=err, capsize=6, color=[NAVY_HEX, BLUE, GOLD_HEX, "#7F8C8D"], width=0.58)
    ax.set_ylim(0, 18)
    ax.set_ylabel("신규 주문 가격 인상률(%)")
    ax.set_title("삼성 파운드리 공정별 가격 인상 보도 범위", color=NAVY_HEX)
    ax.grid(axis="y", color="#EAECF0", lw=0.8)
    for i, (a, b) in enumerate(zip(low, high)):
        label = f"{a}~{b}%" if a != b else f"약 {a}%"
        ax.text(i, b + 0.7, label, ha="center", fontweight="bold", color=NAVY_HEX)
    return savefig(fig, "08_foundry_hikes.png")


def chart_domestic_growth() -> Path:
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(9.4, 4.3))
    ax1.bar(["1Q26", "2Q26", "수주잔고"], [7, 11, 20], color=[BLUE, NAVY_HEX, GOLD_HEX], width=0.58)
    ax1.set_title("이수페타시스 Multi-Lam 비중", color=NAVY_HEX)
    ax1.set_ylabel("%")
    ax1.set_ylim(0, 24)
    ax1.grid(axis="y", color="#EAECF0", lw=0.8)
    for i, v in enumerate([7, 11, 20]):
        ax1.text(i, v + 0.7, "20%+" if i == 2 else f"{v}%", ha="center", fontweight="bold")
    ax2.bar(["현재", "27년 2Q", "28년 하반기"], [1200, 1500, 1800], color=[BLUE, NAVY_HEX, GOLD_HEX], width=0.58)
    ax2.set_title("월 매출 Capa 로드맵", color=NAVY_HEX)
    ax2.set_ylabel("억원/월")
    ax2.set_ylim(0, 2050)
    ax2.grid(axis="y", color="#EAECF0", lw=0.8)
    for i, v in enumerate([1200, 1500, 1800]):
        ax2.text(i, v + 55, f"{v:,}", ha="center", fontweight="bold")
    return savefig(fig, "09_domestic_growth.png")


def chart_valuation() -> Path:
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(9.6, 4.4))
    names = ["SK하이닉스\n26E", "SK하이닉스\n27E", "삼성전자\n26E", "삼성전자\n27E", "Micron\nCY27", "Sandisk\nFY27"]
    per = [4.3, 3.4, 5.2, 3.7, 6.25, 7.8]
    ax1.bar(names, per, color=[NAVY_HEX, BLUE, "#4356A4", "#6574C4", GOLD_HEX, "#8D6E36"])
    ax1.set_ylim(0, 9)
    ax1.set_title("메모리 피어 PER", color=NAVY_HEX)
    ax1.set_ylabel("배")
    ax1.tick_params(axis="x", rotation=28)
    ax1.grid(axis="y", color="#EAECF0", lw=0.8)
    for i, v in enumerate(per):
        ax1.text(i, v + 0.2, f"{v:.1f}", ha="center", fontsize=8, fontweight="bold")
    ax2.bar(["시장 프레임", "Unitree\n2026E"], [60, 155], color=[GOLD_HEX, RED_HEX], width=0.58)
    ax2.set_title("Unitree PSR 비교", color=NAVY_HEX)
    ax2.set_ylabel("배")
    ax2.set_ylim(0, 180)
    ax2.grid(axis="y", color="#EAECF0", lw=0.8)
    for i, v in enumerate([60, 155]):
        ax2.text(i, v + 5, f"{v}배", ha="center", fontweight="bold")
    fig.suptitle("싸 보이는 메모리 vs 기대가 앞선 휴머노이드", color=NAVY_HEX, fontweight="bold")
    return savefig(fig, "10_valuation.png")


def make_charts() -> dict[str, Path]:
    setup_charts()
    return {
        "macro": chart_macro_thresholds(),
        "carry": chart_carry_shock(),
        "fx": chart_fx_eps(),
        "return": chart_shareholder_return(),
        "fcf": chart_fcf_reconciliation(),
        "memory": chart_memory_balance(),
        "moves": chart_market_moves(),
        "foundry": chart_foundry_hikes(),
        "domestic": chart_domestic_growth(),
        "valuation": chart_valuation(),
    }


def build() -> None:
    charts = make_charts()
    n = Notes()

    # 문서 메타데이터/머리말 교체
    n.doc.core_properties.title = "8월 19일 시장전략 보고서"
    n.doc.core_properties.author = "준혁"
    n.doc.core_properties.subject = "매크로·반도체·주주환원·섹터 로테이션"
    header = n.doc.sections[0].header.paragraphs[0]
    header.clear()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("2026. 8. 19.  ·  MARKET STRATEGY  ·  QUICK COMMENT SYNTHESIS")
    set_run_font(r, size=8.5, color=GRAY)

    n.p("2026. 8. 19.  |  장 마감 이후 업데이트 포함", size=10, color=GRAY, align="center", space_after=5)
    n.p("시장전략 보고서", size=24, bold=True, color=NAVY, align="center", space_after=3)
    n.p("매크로의 역습과 반도체의 방어선", size=17, bold=True, color=NAVY2, align="center", space_after=5)
    n.p("유가 · 미 국채금리 · 엔화 · 원화  |  HBM · ASIC · 파운드리  |  SK하이닉스 주주환원", size=10.5, color=GRAY, align="center", space_after=12)
    n.callout(
        "EXECUTIVE VIEW",
        [
            "이번 조정의 1차 성격은 AI 수요 훼손보다 ‘유가발 장기금리 쇼크 + 포지션 조정’에 가깝습니다.",
            "단기 방어선은 미 10년물 4.7%와 브렌트유 $90, 위험선은 각각 5.0%와 $100입니다. USD/JPY가 150 초반으로 빠르게 내려가면 엔캐리 청산 경계가 한 단계 올라갑니다.",
            "반도체 내부에서는 SK하이닉스의 40조원 자사주 매입·소각, 삼성 파운드리 가격 인상, Multi-Lam/검사장비 수요가 펀더멘털 방어선입니다.",
            "다만 원화 강세는 수출주 EPS의 역풍입니다. 달러-원 1,520→1,420 가정 시 제공 민감도로 삼성전자 약 -2.6%, SK하이닉스 약 -5.9%입니다.",
        ],
        kind="key",
    )
    n.h2("오늘의 결론 — 투자 판단을 가르는 4개 문장")
    n.table(
        ["축", "판단", "다음 확인"],
        [
            ["매크로", "전쟁 → 유가 → 금리의 연결고리가 주가를 압도", "10Y 4.7% / 5.0%, Brent $90 / $100"],
            ["수급", "금리 하락만으로 기술주가 오르지 않음: 바이오로 일부 로테이션", "SOX 상대수익률, NVIDIA 8/26 실적"],
            ["반도체", "AI CAPEX 궤도는 유지되나 HBM 효율화·ASIC 재편 논쟁 확대", "연산 증가율 vs 메모리 효율 개선률"],
            ["한국", "주주환원은 하방을 보강하지만 원화 강세가 EPS를 상쇄", "달러-원 1,400/1,360, 외국인 수급"],
        ],
        col_widths=[2.5, 8.2, 6.8],
    )

    page_break(n)
    n.h1("시장 국면: 매크로가 펀더멘털을 덮은 날", num="1.")
    n.flow(["미·이란 확전 우려", "호르무즈 리스크", "유가 상승", "인플레이션 우려", "장기금리 상승", "AI·반도체 할인율 압박"])
    add_chart(n, charts["macro"], caption="※ 구간은 제공 코멘트의 투자 프레임을 시각화한 것이며 실시간 시세가 아닙니다.")
    n.h2("왜 반도체가 더 크게 흔들렸나")
    n.bullet("고PER 성장주는 장기 현금흐름의 듀레이션이 길어 할인율 상승에 민감합니다.")
    n.bullet("그러나 메모리는 이미 사이클 특성을 반영한 낮은 PER이어서 ‘금리 하나만으로’ 낙폭을 설명하기에는 부족합니다.")
    n.bullet("최근 2주 급반등 이후 차익실현, 외국인 현·선물 수급, AI 롱 포지션 피로가 동시에 작용한 것으로 보는 편이 타당합니다.")
    n.callout("판단", ["현재까지는 AI CAPEX 훼손보다 유가·금리·포지션의 복합 쇼크. 10년물 4.7% 아래 안착 여부가 가장 빠른 확인 지표입니다."], kind="blue")

    n.h2("미 재무부 바이백: 단기 방어, 구조적 해결은 아님")
    n.table(
        ["효과", "단기", "한계"],
        [
            ["장기채 수급", "장기채 수요 보강 → 금리 급등 속도 완화", "재정적자와 발행 증가를 없애지는 못함"],
            ["위험자산", "할인율 진정 → 성장주에 우호적", "섹터 로테이션이 효과를 상쇄할 수 있음"],
            ["정책 신호", "금리 급등을 좌시하지 않겠다는 신호", "Fed의 매파적 반응과 별개"],
        ],
        col_widths=[3.0, 7.1, 7.4],
    )

    page_break(n)
    n.h1("엔캐리: ‘일본 금리’보다 엔화 속도가 중요", num="2.")
    add_chart(n, charts["carry"], caption="제공 코멘트의 2024년 7/31→8/5 수치를 사용한 비교.")
    n.h2("현재와 2024년 8월형 쇼크를 구분하는 체크포인트")
    n.table(
        ["구분", "정상화 흡수", "캐리 청산 경계"],
        [
            ["USD/JPY", "157~159에서 안정·반등", "159→155→150 초반으로 급락"],
            ["미국 금리", "안정 또는 완만한 움직임", "미 금리 하락과 엔화 급등 동시 발생"],
            ["일본 국채", "입찰 견조, 변동성 제한", "입찰 부진·장기금리 급등"],
            ["주식", "밸류에이션 조정", "Nikkei·KOSPI 동반 투매"],
        ],
        col_widths=[3.0, 7.2, 7.3],
    )
    n.callout(
        "해석",
        [
            "USD/JPY가 높은 수준을 유지하면 1차 충격은 ‘글로벌 금리 상승/밸류에이션 압박’입니다.",
            "엔화가 짧은 기간 강하게 절상되고 주식이 동반 급락하면 레버리지 청산 가능성을 본격적으로 의심해야 합니다.",
        ],
        kind="note",
    )

    page_break(n)
    n.h1("달러-원: 국내 달러 공급이 먼저 움직였다", num="3.")
    n.h2("1,400원 아래로 내려온 경로")
    n.flow(["8월 말 법인세·설비투자 원화 수요", "수출기업 달러 매도", "환헤지 비중 상승", "고환율 추가 매도", "달러-원 하락 가속"])
    n.p("이번 원화 강세는 달러인덱스와 엔화만으로 설명하기 어렵습니다. 핵심은 국내 달러 공급이 강해 달러-원이 먼저 하락했다는 수급 해석입니다.")
    n.table(
        ["시나리오", "조건", "달러-원 함의"],
        [
            ["기본", "국내 달러 공급 지속", "1,400원 하회 유지 가능"],
            ["추가 하락", "Fed 인상 기대 되돌림 → DXY 99→96~97", "약 1,360원 하단 가능"],
            ["강한 원화", "약달러 + 국내 공급 동시", "1,340원대 조건부"],
            ["제동", "외국인 주식 매도·유가 상승·1,350원대 달러 수요", "추가 하락 제한"],
        ],
        col_widths=[3.0, 8.3, 6.2],
    )
    add_chart(n, charts["fx"], caption="단순 선형 민감도. 실제 EPS는 환헤지·통화별 매출/비용·가격 전가에 따라 달라질 수 있습니다.")
    n.callout(
        "원문 오기 교정",
        [
            "원문에는 ‘삼성전자 환율 민감도’ 설명 뒤에도 SK하이닉스가 반복 표기되어 있습니다.",
            "문맥과 -5.9% 산식을 기준으로 삼성전자 +0.4%, SK하이닉스 +0.9%로 정리했습니다.",
        ],
        kind="note",
    )

    page_break(n)
    n.h1("SK하이닉스: 40조 소각과 FCF 50% 이상", num="4.")
    n.callout(
        "핵심",
        [
            "40조원 자사주를 8/20~11/19 매입 후 전량 소각. 총주식의 약 3.3%, 동일 순이익 가정 시 EPS 약 +3.4%.",
            "2025~2027 누적 FCF의 ‘50% 범위 내’가 아니라 ‘50% 이상’을 환원. 자사주와 현금배당을 병행하고 특별배당도 검토.",
            "구체적 추가 환원 규모와 방식은 3Q26 실적 발표 시 이사회 결의를 거쳐 안내 예정.",
        ],
        kind="key",
    )
    add_chart(n, charts["return"])
    n.h2("FCF 숫자: 반드시 두 시나리오를 분리해야 한다")
    add_chart(n, charts["fcf"], caption="단순 산술 검증. 실제 FCF·배당가능이익·이사회 결의에 따라 달라집니다.")
    n.table(
        ["입력", "산식", "50% 기준", "40조 제외 추가분"],
        [
            ["누적 FCF 385조", "385 × 50%", "192.5조", "152.5조"],
            ["연도별 150/210/205조", "합계 565 × 50%", "282.5조", "242.5조"],
        ],
        col_widths=[4.1, 5.3, 3.8, 4.3],
    )
    n.callout(
        "산술 불일치 교정",
        [
            "‘150 + 210 + 205 = 누적 385조’는 성립하지 않습니다. 합계는 565조입니다.",
            "따라서 192.5조 환원 결론은 누적 FCF 385조 시나리오에만 해당합니다.",
            "2028년 102.5조는 회사 정책이 아니라 별도 FCF 모델에 50%를 적용한 참고치입니다. 2028 정책은 새로 발표돼야 합니다.",
        ],
        kind="bear",
    )

    # FCF 보정 박스 다음의 여백을 활용해 HBM 파트를 이어 배치한다.
    n.h1("HBM 논쟁: 수요 절벽이 아니라 효율화의 경제성", num="5.")
    n.p("“HBM 가격이 올랐으니 수요가 곧 꺾인다”는 단정은 약합니다. 더 정확한 명제는 “가격이 높을수록 HBM 사용량을 줄이는 기술 개발의 경제적 유인이 커진다”입니다.")
    add_chart(n, charts["memory"], caption="방향성 비교이며 수요량을 직접 계산한 값이 아닙니다.")
    n.h2("대체가 아니라 메모리 계층의 재배치")
    n.table(
        ["기술/구조", "현실성", "의미"],
        [
            ["Cerebras 온칩 SRAM", "높음", "HBM 없이 대규모 온칩 메모리 활용"],
            ["Groq SRAM 중심 추론", "높음", "결정론적·저지연 추론에 최적화"],
            ["KV Cache 압축", "높음", "필요 HBM 용량과 대역폭 절감"],
            ["CPU DRAM·SSD 오프로딩", "진행 중", "비용·지연을 절충한 계층화"],
            ["HBF/스토리지 계층", "진행 중", "용량 중심 워크로드의 대안"],
        ],
        col_widths=[4.5, 2.7, 10.0],
    )
    n.callout(
        "2026~28 vs 2028 이후",
        [
            "단기: HBM 가격결정력은 메모리 업체의 초과이익을 확대.",
            "장기: 가격이 지나치면 SRAM 확대·압축·모델 경량화·ASIC 전환을 촉진.",
            "따라서 ‘현재 메모리주가 틀렸다’가 아니라 ‘현재의 가격결정력이 영구적이라고 가정하지 말라’는 경고로 읽는 편이 타당합니다.",
        ],
        kind="blue",
    )

    page_break(n)
    n.h1("AI 내부 로테이션: Broadcom 독점 견제와 바이오 피벗", num="6.")
    add_chart(n, charts["moves"], caption="등락률은 제공 코멘트 수치를 사용. Moderna 수치는 코멘트 내 +77%/+117%/+177%가 혼재해 8/19 종목표의 +77%를 채택.")
    n.h2("Google–Marvell: 칩 설계에서 TPU 주변 생태계로")
    n.table(
        ["항목", "내용", "투자 함의"],
        [
            ["적용 영역", "AI 추론 가속기·Storage Controller·NIC·Memory Interface·Near-memory", "커스텀 ASIC 가시성 확대"],
            ["Warrant", "최대 5,897만주, 행사가 $206.58", "장기 협력의 경제적 이해관계"],
            ["Vesting", "Google 관련 Custom Products 매출 $5억마다 tranche", "매출 증가와 보상 연동"],
            ["기간", "FY27 Q3~FY33", "단기 이벤트보다 장기 계약 성격"],
            ["경쟁 구도", "Broadcom의 TPU 독점적 지위 견제 신호", "ASIC 내 점유율 재배분"],
        ],
        col_widths=[3.1, 8.1, 6.2],
    )
    n.h2("금리 하락에도 기술주가 못 오른 이유")
    n.flow(["재무부 바이백", "장기금리 진정", "기술주 할인율 개선", "그러나 AI 차익실현", "헬스케어 유입", "기술주 약세/바이오 강세"])
    n.p("금리 하나만으로 시장을 설명할 수 없는 국면입니다. NVIDIA 8/26 실적 전후 AI 심리가 재강화되는지, 헬스케어로의 성장주 로테이션이 이어지는지를 함께 봐야 합니다.")

    page_break(n)
    n.h1("국내 반도체: 가격·믹스·검사 난이도가 이익을 만든다", num="7.")
    n.h2("삼성 파운드리 — 생산 여력 부족이 가격 협상력으로")
    add_chart(n, charts["foundry"], caption="Reuters 링크가 포함된 제공 코멘트 기준.")
    n.bullet("TSMC 첨단공정 포화와 중국 팹리스의 해외 파운드리 의존이 신규 주문의 가격 협상력을 높이는 구도.")
    n.bullet("평택 SF4는 퀄컴 물량과 차세대 HBM 베이스다이 생산으로 높은 가동률을 유지한다는 보도.")
    n.bullet("가격 인상 + 수율 개선 + 가동률 상승이 결합되면 파운드리 적자 축소/흑자 전환 기대를 강화.")
    n.h2("이수페타시스 — Capa보다 Multi-Lam 믹스가 중요")
    add_chart(n, charts["domestic"], caption="수주잔고 비중은 ‘20%+’를 시각화하기 위해 20으로 표시.")
    n.flow(["AI 서버·가속기", "신호 복잡도 증가", "고다층 Multi-Lam", "ASP·난이도 상승", "수율 안정", "이익 레버리지"])
    n.h2("기가비스 — FC-BGA의 눈(AOI)과 레이저 수리공(AOR)")
    n.table(
        ["촉매", "현재 확인", "리스크"],
        [
            ["AI → FC-BGA 고사양화", "미세회로 검사·수리 중요도 상승", "기판 CAPEX 지연"],
            ["글로벌 고객 증설", "일본 고객 89.5억 계약(매출의 17.1%)", "장비 수주 변동성"],
            ["영업 레버리지", "26E 매출 1,785억·OP 721억 제시", "높은 기대치와 고객 집중"],
        ],
        col_widths=[4.3, 7.1, 5.8],
    )

    # 기가비스 표 뒤에 비반도체 확산 파트를 이어 불필요한 공백 페이지를 막는다.
    n.h1("비반도체 확산: 전력·ESS·방산", num="8.")
    n.table(
        ["테마", "기업/이벤트", "핵심 연결고리"],
        [
            ["전력 인프라", "LS·LS ELECTRIC·LS전선", "배전설비·케이블·특수권선 + 구리가격 상승"],
            ["북미 ESS", "LG에너지솔루션", "EV 라인 전환, Lansing 35GWh+, Tesla Megapack LFP"],
            ["배터리 방산", "LG에너지솔루션", "드론·무인무기체계용 배터리 공급 협상"],
            ["미 자주포", "한화에어로스페이스", "MTC 시제기 단독 선정 → 최대 18문, 양산 진입 옵션"],
            ["바이오", "알테오젠·mRNA 밸류체인", "키트루다SC + 개인맞춤형 암백신 병용 확장"],
        ],
        col_widths=[3.0, 5.1, 9.1],
    )
    n.h2("LS: 자회사 이익 체력의 구조적 상승")
    n.table(
        ["자회사", "2Q26 영업이익", "동력"],
        [
            ["LS ELECTRIC", "1,785억 (+64% YoY)", "배전설비·변압기·배전반 수주"],
            ["LS전선", "1,413억 (+71%)", "고부가 케이블·구리가격"],
            ["LS MnM", "1,757억 (흑전)", "황산 수익성"],
            ["LS I&D", "741억 (+166%)", "Essex 특수권선"],
        ],
        col_widths=[4.1, 5.3, 7.8],
    )
    n.callout("포트폴리오 시사점", ["국내 기관의 고민은 AI 비중을 ‘0’으로 줄이는 것이 아니라 대형 반도체 노출과 전력·소부장·방산의 믹스 비율을 조정하는 문제에 가깝습니다."], kind="key")

    page_break(n)
    n.h1("밸류에이션: 메모리의 저PER와 Unitree의 초고PSR", num="9.")
    add_chart(n, charts["valuation"], caption="모든 배수는 제공 코멘트의 추정치·가격 기준. 서로 다른 회계연도와 추정치이므로 방향성 비교용.")
    n.h2("메모리: 낮은 PER은 기회인 동시에 피크 우려의 가격")
    n.table(
        ["종목", "가격/기준", "26E PER", "27E PER", "보수적 27E"],
        [
            ["SK하이닉스", "본주 150만원", "4.3배", "3.4배", "5.1배"],
            ["삼성전자", "24.75만원", "5.2배", "3.7배", "5.6배"],
            ["Micron", "$937.11", "F12M 7.5배", "CY27 6.25배", "—"],
            ["Sandisk", "$1,568.37", "—", "FY27 7.8배", "—"],
        ],
        col_widths=[3.2, 4.2, 3.2, 3.2, 3.2],
    )
    n.h2("Unitree: 산업 성장과 주가의 속도를 분리")
    n.bullet("휴머노이드 시장 CAGR 31% 전망과 중국의 제조·공급망·원가 우위는 강점.")
    n.bullet("그러나 2026E 매출 약 22억 위안 대비 종가 시총 3,418억 위안이면 PSR 약 155배.")
    n.bullet("‘시장 성장의 2배인 60%+ 성장 시 PSR 60배까지 허용’이라는 공격적 프레임보다도 약 2.6배 높음.")
    n.bullet("다음 관문은 연구·교육·시연이 아니라 산업 현장의 가동률, 반복 신뢰성, 고객 ROI입니다.")

    page_break(n)
    n.h1("실전 체크리스트", num="10.")
    n.table(
        ["우선순위", "변수", "완화 신호", "위험 신호", "영향"],
        [
            ["1", "미 10년물", "4.7% 이하 안착", "5.0% 돌파·고착", "AI/반도체 멀티플"],
            ["2", "브렌트유", "$90 전후 안정", "$100 이상", "인플레·금리 악순환"],
            ["3", "USD/JPY", "157~159 안정", "150 초반 급락", "엔캐리 청산"],
            ["4", "달러-원", "완만한 하락", "급격한 1,340대 진입", "수출주 EPS·외인 수급"],
            ["5", "AI 수요", "토큰·Hyperscaler CAPEX 지속", "OpenAI 성장 둔화·손실 확대", "GPU/HBM/네트워크"],
            ["6", "주주환원", "3Q 추가 환원 구체화", "FCF·배당가능이익 제약", "SK하이닉스 하방"],
            ["7", "HBM", "연산 증가 > 효율 개선", "효율 개선이 연산 증가 추월", "메모리 가격결정력"],
        ],
        col_widths=[1.5, 3.0, 4.2, 4.3, 4.3],
    )
    n.h2("시나리오 맵")
    n.table(
        ["시나리오", "매크로", "시장 반응", "선호"],
        [
            ["Bull", "유가 안정 + 10Y <4.7 + 엔화 안정", "반도체 리레이팅·주주환원 반영", "대형 메모리 + AI 소부장"],
            ["Base", "금리 고점권 횡보 + AI CAPEX 유지", "변동성 속 실적·수주 차별화", "대형주 + 전력/장비 바벨"],
            ["Bear", "유가 >$100 + 10Y >5.0 + 엔화 급등", "멀티플 축소·캐리 청산", "현금흐름·방어주·비중 축소"],
        ],
        col_widths=[2.5, 6.2, 5.2, 3.5],
    )
    n.callout(
        "최종 결론",
        [
            "매크로가 시장의 속도를 결정하지만 AI 투자의 방향을 아직 뒤집지는 않았습니다.",
            "주주환원은 한국 메모리의 하방을 강화하지만 원화 강세와 장기금리가 상단을 제약합니다.",
            "따라서 단기에는 유가·10년물·엔화를, 중기에는 토큰 수요·CAPEX·HBM 효율화·추가 환원을 순서대로 확인합니다.",
        ],
        kind="key",
    )

    n.h2("자료 기준과 주의")
    n.p("본 보고서는 사용자가 제공한 2026년 8월 19일 퀵 코멘트를 재구성한 자료입니다. Reuters·DART·FT·국내 기사 링크가 포함된 항목은 원 코멘트의 출처 표기를 유지했으며, 모든 수치를 독립적으로 재검증한 실시간 시세 보고서는 아닙니다.", size=9.5, color=GRAY)
    n.p("수치가 충돌하는 항목은 본문에서 채택 기준을 밝혔습니다. 특히 FCF 385조/565조, Moderna 등락률, 환율 민감도 회사명은 원문 내 불일치를 교정하거나 병기했습니다.", size=9.5, color=GRAY)
    n.p("투자 권유가 아니며, 실제 의사결정 전 공시·회사자료·실시간 시장 데이터를 다시 확인해야 합니다.", size=9.5, color=RED, bold=True)

    n.save(OUT_PATH)
    print(f"Wrote {OUT_PATH} ({OUT_PATH.stat().st_size:,} bytes)")
    print(f"Embedded charts: {len(charts)}")


if __name__ == "__main__":
    build()
