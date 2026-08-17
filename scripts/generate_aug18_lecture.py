#!/usr/bin/env python3
"""8월 18일 NON-삼전닉스 강의노트(.docx) 생성."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Mm, Pt, RGBColor

OUT_PATH = Path("/workspace/lectures/8월 18일 NON-삼전닉스 (소부장, SK, 자동차).docx")

# 한국 Windows에서 기본으로 열리는 서체. 없는 환경은 Calibri로 폴백.
KR_FONT = "맑은 고딕"
EN_FONT = "Calibri"

NAVY = RGBColor(0x0F, 0x20, 0x43)
NAVY2 = RGBColor(0x1E, 0x40, 0x7C)
GOLD = RGBColor(0xB8, 0x94, 0x3A)
GRAY = RGBColor(0x4B, 0x55, 0x63)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x16, 0x65, 0x34)
RED = RGBColor(0x99, 0x1B, 0x1B)
AMBER = RGBColor(0x7A, 0x5C, 0x12)

NAVY_HEX = "0F2043"
NAVY2_HEX = "1E407C"
GOLD_HEX = "B8943A"
LIGHT_HEX = "EEF2F8"
GREEN_HEX = "E8F5E9"
RED_HEX = "FDECEA"
AMBER_HEX = "FFF8E7"
BLUE_HEX = "E8F1FB"
ROW_HEX = "F7F9FC"
WHITE_HEX = "FFFFFF"


def set_run_font(run, size=11, bold=False, color=DARK, italic=False, font=KR_FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), KR_FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(
        parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill}"/>')
    )


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(
        parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'<w:top w:w="{top}" w:type="dxa"/>'
            f'<w:left w:w="{left}" w:type="dxa"/>'
            f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
            f'<w:right w:w="{right}" w:type="dxa"/>'
            f"</w:tcMar>"
        )
    )


def set_table_borders(table, color="D0D7E2", sz="4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f"</w:tblBorders>"
    )
    tbl_pr.append(borders)


def set_left_accent(cell, color=NAVY_HEX, sz="24"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(
        parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="nil"/>'
            f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="nil"/>'
            f'<w:right w:val="nil"/>'
            f"</w:tcBorders>"
        )
    )


def prevent_row_split(row):
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()
    tr_pr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))


def cell_text(cell, text, size=10, bold=False, color=DARK, align="left", font=KR_FONT):
    cell.text = ""
    align_enum = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    lines = str(text).split("\n")
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = align_enum
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1 if i < len(lines) - 1 else 0)
        p.paragraph_format.line_spacing = 1.12
        run = p.add_run(line)
        set_run_font(run, size=size, bold=bold, color=color, font=font)
    set_cell_margins(cell)


def add_runs(paragraph, parts, size=11, color=DARK):
    """parts: str or (text, bold, color?) tuples."""
    for part in parts:
        if isinstance(part, str):
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)
        else:
            text, bold, *rest = part
            c = rest[0] if rest else color
            run = paragraph.add_run(text)
            set_run_font(run, size=size, bold=bold, color=c)


class Notes:
    def __init__(self):
        self.doc = Document()
        self._setup()

    def _setup(self):
        sec = self.doc.sections[0]
        sec.page_width = Mm(210)
        sec.page_height = Mm(297)
        sec.left_margin = Mm(16)
        sec.right_margin = Mm(16)
        sec.top_margin = Mm(16)
        sec.bottom_margin = Mm(16)
        sec.header_distance = Mm(8)
        sec.footer_distance = Mm(8)

        normal = self.doc.styles["Normal"]
        normal.font.name = KR_FONT
        normal.font.size = Pt(11)
        normal.font.color.rgb = DARK
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), KR_FONT)
        pf = normal.paragraph_format
        pf.space_after = Pt(6)
        pf.space_before = Pt(0)
        pf.line_spacing = 1.18

        header = sec.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = hp.add_run("8/18 NON-삼전닉스  ·  소부장 · SK · Atlas  ·  강의노트")
        set_run_font(r, size=8.5, color=GRAY)

        footer = sec.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run("영상 녹화용 정리  ·  숫자·시나리오는 공개 코멘트 기준  ·  ")
        set_run_font(r, size=8, color=GRAY)
        fld = parse_xml(
            f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE ">'
            f'<w:r><w:rPr><w:sz w:val="16"/><w:color w:val="4B5563"/>'
            f'<w:rFonts w:ascii="{KR_FONT}" w:hAnsi="{KR_FONT}" w:eastAsia="{KR_FONT}"/></w:rPr>'
            f"<w:t></w:t></w:r></w:fldSimple>"
        )
        fp._p.append(fld)

        core = self.doc.core_properties
        core.title = "8월 18일 NON-삼전닉스 (소부장, SK, 자동차)"
        core.author = "준혁"
        core.subject = "소부장(테스·한미반도체·원익IPS), SK 지주사/에코플랜트, Atlas"

    def p(self, text, size=11, bold=False, color=DARK, space_after=6, space_before=0, align="left"):
        para = self.doc.add_paragraph()
        para.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }[align]
        para.paragraph_format.space_after = Pt(space_after)
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.line_spacing = 1.18
        run = para.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
        return para

    def rich(self, parts, size=11, space_after=6, space_before=0):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_after = Pt(space_after)
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.line_spacing = 1.18
        add_runs(para, parts, size=size)
        return para

    def h1(self, text, num=None):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing = 1.1
        if num:
            run = para.add_run(f"{num}  ")
            set_run_font(run, size=16, bold=True, color=GOLD)
        run = para.add_run(text)
        set_run_font(run, size=16, bold=True, color=NAVY)
        # underline bar
        pPr = para._p.get_or_add_pPr()
        pPr.append(
            parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="12" w:space="4" w:color="{NAVY_HEX}"/>'
                f"</w:pBdr>"
            )
        )
        return para

    def h2(self, text):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(text)
        set_run_font(run, size=13, bold=True, color=NAVY2)
        return para

    def h3(self, text):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(3)
        run = para.add_run(text)
        set_run_font(run, size=11.5, bold=True, color=NAVY)
        return para

    def bullet(self, text, level=0, bold_lead=None, size=11):
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.55 + level * 0.45)
        para.paragraph_format.first_line_indent = Cm(-0.35)
        para.paragraph_format.space_after = Pt(2.5)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.line_spacing = 1.15
        mark = "• " if level == 0 else "– "
        run = para.add_run(mark)
        set_run_font(run, size=size, color=NAVY2 if level == 0 else GRAY)
        if bold_lead:
            run = para.add_run(bold_lead)
            set_run_font(run, size=size, bold=True, color=DARK)
            run = para.add_run(text)
            set_run_font(run, size=size, color=DARK)
        else:
            run = para.add_run(text)
            set_run_font(run, size=size, color=DARK)
        return para

    def flow(self, items, size=11):
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing = 1.2
        for i, item in enumerate(items):
            if i:
                run = para.add_run("   →   ")
                set_run_font(run, size=size, bold=True, color=GOLD)
            run = para.add_run(item)
            set_run_font(run, size=size, bold=True, color=NAVY)
        return para

    def callout(self, title, body, kind="key"):
        palette = {
            "key": (NAVY_HEX, LIGHT_HEX, NAVY),
            "bull": ("166534", GREEN_HEX, GREEN),
            "bear": ("991B1B", RED_HEX, RED),
            "note": (GOLD_HEX, AMBER_HEX, AMBER),
            "blue": (NAVY2_HEX, BLUE_HEX, NAVY2),
        }
        accent, fill, title_color = palette[kind]
        table = self.doc.add_table(rows=1, cols=1)
        table.autofit = True
        table.allow_autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        shade_cell(cell, fill)
        set_left_accent(cell, accent, sz="28")
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        p1.paragraph_format.space_before = Pt(0)
        r = p1.add_run(title)
        set_run_font(r, size=10, bold=True, color=title_color)
        if isinstance(body, str):
            body = [body]
        for line in body:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(line)
            set_run_font(r, size=10.5, color=DARK)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def table(self, headers, rows, col_widths=None, first_col_bold=True):
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        set_table_borders(table, color="D5DCE6", sz="4")
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            shade_cell(cell, NAVY_HEX)
            cell_text(cell, h, size=9.5, bold=True, color=WHITE, align="center")
        prevent_row_split(table.rows[0])
        for r_i, row in enumerate(rows):
            for c_i, val in enumerate(row):
                cell = table.rows[r_i + 1].cells[c_i]
                if r_i % 2 == 1:
                    shade_cell(cell, ROW_HEX)
                else:
                    shade_cell(cell, WHITE_HEX)
                align = "left" if c_i == 0 else "center"
                bold = first_col_bold and c_i == 0
                cell_text(cell, str(val), size=9.5, bold=bold, color=DARK, align=align)
            prevent_row_split(table.rows[r_i + 1])
        if col_widths:
            for row in table.rows:
                for i, w in enumerate(col_widths):
                    row.cells[i].width = Cm(w)
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(8)
        spacer.paragraph_format.space_before = Pt(2)
        return table

    def spacer(self, pt=4):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(pt)
        p.paragraph_format.space_before = Pt(0)

    def save(self, path: Path):
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))


def build():
    n = Notes()

    # ── Cover ──────────────────────────────────────────────
    n.p("2026. 8. 18. 강의노트  ·  영상 녹화 / 오후–저녁 업로드", size=10.5, color=GRAY, align="center", space_after=4)
    n.p("NON-삼전닉스", size=13, bold=True, color=GOLD, align="center", space_after=2)
    n.p("8월 18일  NON-삼전닉스", size=22, bold=True, color=NAVY, align="center", space_after=2)
    n.p("(소부장, SK, 자동차)", size=16, bold=True, color=NAVY2, align="center", space_after=8)
    n.p("테스 · 한미반도체 · 원익IPS  |  SK 지주사 SOTP · SK에코플랜트  |  Atlas", size=11, color=GRAY, align="center", space_after=10)

    n.callout(
        "오늘 한 장으로 보면",
        [
            "공통 분모는 AI CAPEX입니다. 다만 수요를 정확히 예측할 수 없으니 모델·데이터센터·GPU 전망에 기대고, 자금이 빠듯하니 서로 묶어서 가는 구간입니다.",
            "그래서 장비는 수주, SK는 할인율, Atlas는 실제 공장 배치를 확인하면 됩니다. 우려·시나리오·병목이 말끔히 해소되기 전까지 양면적 시각은 계속됩니다.",
        ],
        kind="key",
    )

    n.h2("강의 구성")
    n.table(
        ["파트", "대상", "오늘 확인할 것"],
        [
            ["1", "시장 프레임", "Dark GPU — 수요 부정이 아니라 과잉투자·금융 레버리지 경고"],
            ["2", "테스", "수주 2배 · DRAM+NAND 동시 수혜 · BSD/Tetra 성장 옵션"],
            ["3", "한미반도체", "HBM4 TCB · 한화세미텍 변수 · 고객 다변화"],
            ["4", "원익IPS", "2Q 부진 ≠ 수주 훼손 · 2027 성장 지속 여부"],
            ["5", "SK / 에코플랜트", "SOTP 할인율 · 에코플랜트 저평가 · 실트론 매각"],
            ["6", "Atlas", "연구용 → 공장 투입 · 모비스/글로비스 수혜"],
            ["보론", "리노공업 · 메모리", "파업·마진 / ADR·PER 비교 프레임"],
        ],
        col_widths=[2.2, 4.4, 11.0],
    )

    n.h2("오프닝 멘트 (녹화용)")
    n.p("오늘은 삼전닉스가 아니라 NON-삼전닉스입니다. 소부장 세 종목 — 테스, 한미반도체, 원익IPS — 그리고 SK 지주사, 특히 SK에코플랜트, 마지막으로 Atlas 포인트입니다. 테스·한미반도체 목표주가 시나리오의 세부 산식은 강의에서 풀겠습니다.")

    # ── Cheat sheet ────────────────────────────────────────
    n.h1("오늘 숫자 한 장", num="0.")
    n.table(
        ["종목", "핵심 숫자", "한 줄 결론"],
        [
            ["테스", "2Q 신규수주 2,823억(QoQ 2배)\n잔고 2,000억+  /  TP 20~23만", "DRAM+NAND 동시. BSD가 성장 옵션"],
            ["한미반도체", "2Q 매출 2,511억 · OPM 51.9%\nBofA 42만 vs 상상인 38만", "핵심은 TCB 성장률 × 점유율"],
            ["원익IPS", "1Q 잔고 ~4,000억\nTP 평균 14.8만 → 안전마진 10.4~11.8만", "하반기 레벨업 vs 2027 정점론"],
            ["SK", "NAV 81.7조 × 할인 41%\n대신 TP 88만 유지", "에코플랜트 재평가 시 96~106만"],
            ["SK에코플랜트", "2Q OP 5,340억 · 잔고 26.8조\n표 가치 2.1조", "반도체·AI DC를 동시에 먹는 회사"],
            ["Atlas", "2028 HMGMA → 2030 조립 확대\n대당 액추에이터 31개", "성능보다 Fleet × 가동률 × 데이터"],
            ["리노공업", "2Q OPM 51.3% · 파업 약 4주\n삼성 TP 10만(하향)", "매출보다 마진 훼손이 핵심"],
        ],
        col_widths=[3.4, 7.2, 7.0],
        first_col_bold=True,
    )

    # ── 1. Dark GPU ────────────────────────────────────────
    n.h1("AI 인프라 금융과 Dark GPU", num="1.")
    n.callout(
        "한 줄 결론",
        [
            "핵심은 “엔비디아 GPU 수요가 당장 꺾인다”가 아닙니다.",
            "AI 인프라에 금융이 붙으면서 ‘과잉투자 위험’이 커질 수 있다는 경고입니다.",
        ],
        kind="key",
    )

    n.h2("1) 중국 언론이 전한 Sacks 발언 요약")
    n.p("월가견문(wallstreetcn) 보도: 엔비디아 5,000억 달러 AI 금융 계획에 숨은 우려, 트럼프 고문 ‘다크 GPU’ 과잉 위험 경고.")
    n.bullet("트럼프 대통령 기술 고문 David Sacks는 팟캐스트에서, 엔비디아 5,000억 달러 AI 금융 계획이 직면한 가장 큰 위험을 연산능력 과잉으로 지목.")
    n.bullet("인터넷 버블 이후 ‘다크 광섬유(dark fiber)’ 위기에 비유. GPU가 대량 유휴 상태에 놓이고 가격이 붕괴하면 AI 인프라 투자망 전체에 충격.")
    n.bullet("동시에 데이터센터 건설에 대한 정치적 반발이 오히려 과잉 공급을 막는 역할을 할 수도 있다고 지적.")

    n.h2("2) Dark GPU가 정확히 무엇인가")
    n.p("과거 닷컴버블 때 통신업체들이 인터넷 수요를 지나치게 낙관해 광섬유를 대량 설치했습니다. 버블 붕괴 후 상당 부분이 사용되지 않았고, 이를 Dark Fiber(미사용 광섬유)라고 불렀습니다.")
    n.flow(["데이터센터 건설", "GPU 대량 구매", "AI 수요 예상", "실제 사용률 부족"])
    n.p("이렇게 되면 가동되지 않는 GPU = Dark GPU가 됩니다. GPU 자체가 고장 난다는 의미가 아니라, 비싸게 투자했지만 충분히 사용되지 않는 연산능력을 뜻합니다. Sacks는 이를 AI 인프라의 가장 큰 잠재적 위험으로 지목했습니다.")

    n.h2("3) 금융이 붙으면 속도가 빨라진다 — 그리고 반대로도")
    n.callout(
        "순방향 레버리지",
        ["AI 수요  →  투자  →  GPU 구매의 속도가 빨라집니다."],
        kind="bull",
    )
    n.callout(
        "역방향 레버리지 (수요가 예상보다 느릴 때)",
        [
            "GPU 과잉 → 가동률 하락 → GPU 임대료/가격 하락",
            "→ 데이터센터 수익성 악화 → 금융비용 부담 → 투자 축소",
        ],
        kind="bear",
    )

    n.h2("4) Sacks는 AI 수요를 부정한 것이 아니다")
    n.p("의외로 Sacks는 AI 수요 자체를 부정하지 않았습니다. 오히려 현재 상황에서 데이터센터 건설의 정치적·물리적 장애물이 Dark GPU를 막아주는 역할을 할 수 있다고 봅니다.")
    n.callout(
        "그의 안전 조건",
        ["AI 수요 증가 속도  =  데이터센터 건설 속도  가 되는 것이 오히려 안전하다."],
        kind="note",
    )
    n.p("데이터센터는 GPU만 주문한다고 바로 가동할 수 없습니다.")
    n.flow(["전력", "송전망", "부지", "인허가", "냉각", "데이터센터", "GPU", "네트워크"])
    n.p("그래서 정치적 반발과 인허가 지연이 역설적으로 AI 공급과잉을 방지하는 브레이크가 될 수 있다는 논리입니다.")

    n.h2("5) 그래서 양면적 시각이 계속된다")
    n.bullet("모든 것은 미래 어느 시점의 최종수요입니다. 그걸 정확히 아무도 예측하지 못합니다.")
    n.bullet("그나마 AI 모델, 데이터센터, GPU 업체 전망에 기대는 것입니다.")
    n.bullet("그런데 투자는 GO GO 하는데/해야 하는데, 자금이 딸릴 것 같으니 서로 묶어서 가보자는 것입니다.")
    n.bullet("앞으로도 이런 우려(시나리오)와 병목이 말끔히 해소되기 전까지, 양면적 시각은 계속될 것입니다.")
    n.callout(
        "오늘 종목과의 연결",
        [
            "테스·한미·원익IPS = CAPEX가 장비 수주로 떨어지는 속도",
            "SK에코플랜트 = 팹 건설 + AI DC + 메모리 유통을 동시에 먹는 회사",
            "Dark GPU 리스크가 현실화되면 가장 먼저 흔들리는 것은 ‘수주 가시성’과 ‘할인율’입니다.",
        ],
        kind="blue",
    )

    # ── 2. TES ─────────────────────────────────────────────
    n.h1("테스", num="2.")
    n.callout(
        "한 줄 결론",
        [
            "2Q 매출은 서프라이즈, 이익은 일회성 비용. 그러나 신규 수주가 전분기 대비 2배로 늘었고 잔고가 쌓이고 있다.",
            "기존 ACL/ARC(NAND)에 BSD(HBM/DRAM)와 Tetra(DRAM 선단)가 붙으면 2027년까지 실적 레벨업 옵션이 열린다.",
            "독점 장비라고 단정하면 안 된다. 핵심은 삼성·하이닉스 퀄리피케이션과 양산 확대다.",
        ],
        kind="key",
    )

    n.h2("1) 2Q26 리뷰 — 증권사 톤")
    n.table(
        ["증권사", "제목/톤", "TP", "포인트"],
        [
            ["SK증권", "2Q26 Review, TP 유지", "23만원 유지", "신규 수주 2,823억, QoQ 2배. 잔고 고려 시 하반기 컨센 상회 여지"],
            ["삼성증권", "안팎으로 좋다, TP 상향", "20만원 상향", "R&D 일시 증가로 OP 부합. BSD 기여 확대. 27년 PER 30배"],
        ],
        col_widths=[2.8, 4.4, 3.2, 7.2],
    )
    n.p("삼성증권 밸류: 27년 PER 30배. 글로벌 전공정 장비 상위 5개사 평균 30배. 과거 테스는 제한적 포트폴리오·높은 NAND 의존으로 동종 대비 할인 거래.")
    n.p("※ 테스·한미반도체 TP 시나리오의 세부 산식은 강의에서 다룹니다.", size=10.5, color=AMBER, bold=True)

    n.h2("2) 실적 흐름 — 수주가 매출로 바뀌는 경로")
    n.flow(["2Q 말 수주잔고 2,000억+", "대부분 연내 매출 인식", "3Q26부터 상반기 수주 매출화", "하반기 추가 수주 반영", "연간 컨센 상향 가능"])
    n.p("2Q는 매출 서프라이즈, 이익은 일회성 비용 영향이었습니다. 쌓이는 수주 잔고 흐름이 주가의 중요한 트리거입니다.")

    n.h2("3) 고객·공정 — DRAM + NAND 동시 수혜")
    n.bullet("삼성전자: P4 → PH4, NAND → 시안, V9 전환 투자")
    n.bullet("SK하이닉스: M15X DRAM 투자 지속")
    n.p("즉, DRAM + NAND 동시 수혜가 테스의 강점입니다. 기존 ACL/ARC가 주로 NAND 투자와 연결됐다면, BSD는 DRAM/HBM으로 고객·공정이 확장될 수 있는 장비입니다.")

    n.h2("4) BSD = Backside Deposition (훨씬 중요)")
    n.p("웨이퍼 뒷면에 막을 입혀 휘어짐(Warpage)을 잡아주는 장비입니다.")
    n.bullet("HBM·고단화 NAND처럼 웨이퍼가 얇아질수록 휘어짐 문제가 커집니다.")
    n.bullet("웨이퍼 뒷면에 SiO/SiN 등을 증착해 앞면과 반대 방향의 응력을 만들어 평탄하게 합니다.")
    n.bullet("따라서 HBM 고단화 → BSD 필요성 증가라는 논리.")
    n.bullet("NAND·DRAM·Foundry 모두 적용 가능.")
    n.bullet("테스는 2025년부터 국내 메모리 고객 NAND 공정에 공급.")
    n.bullet("현재 HBM용 DRAM 및 다른 메모리 고객의 NAND/Foundry용으로 퀄리피케이션 진행 중.")
    n.callout(
        "“테스만 공급?” → 현재는 그렇게 단정하면 안 됩니다",
        [
            "BSD가 차별화된 신규 장비인 것은 맞지만, 글로벌 시장에서 테스만 만들 수 있는 독점 장비라고 보기는 어렵습니다.",
            "중요한 것은 장비 자체의 독점성보다 삼성전자·SK하이닉스에서 실제 퀄리피케이션을 통과하고 양산 적용을 확대할 수 있느냐입니다.",
            "공개자료상 테스는 이미 국내 메모리 고객 NAND에 공급했고, DRAM/HBM으로 확대 평가 중이라는 점이 핵심입니다.",
        ],
        kind="note",
    )

    n.h2("5) Tetra = 고처리량 PECVD")
    n.bullet("Quad chamber 구조. 한 번에 웨이퍼 4장을 처리 → Throughput 향상.")
    n.bullet("SiCN 박막 증착 용도.")
    n.bullet("현재 국내 메모리 고객 DRAM용 퀄리피케이션 진행.")
    n.bullet("기존 장비보다 ASP가 높은 신규 장비로 평가.")
    n.callout(
        "Tetra도 테스 독점이라고 볼 근거는 없습니다",
        [
            "PECVD라는 큰 범주에서는 Applied Materials, Lam Research 등 글로벌 업체와 경쟁하고, 국내에도 여러 증착장비 업체가 있습니다.",
            "Tetra의 의미는 테스가 기존 ACL/ARC 중심에서 DRAM 선단 공정으로 장비 포트폴리오를 확장한다는 것입니다.",
        ],
        kind="note",
    )

    n.h2("6) 확인 포인트")
    n.bullet("3Q부터 상반기 수주가 실제로 매출화되는가.")
    n.bullet("BSD의 DRAM/HBM 퀄 → 양산 전환 여부.")
    n.bullet("Tetra DRAM 퀄 통과와 ASP 믹스 개선.")
    n.bullet("NAND(시안·V9)와 DRAM(M15X·PH4) 동시 수주가 유지되는가.")

    # ── 3. Hanmi ───────────────────────────────────────────
    n.h1("한미반도체", num="3.")
    n.callout(
        "한 줄 결론",
        [
            "2Q26은 확실히 강하다. HBM4 양산 본격화 → TC 본더 수요. OPM 51.9% 사상 최대.",
            "시장의 질문은 실적이 아니라 TCB 시장 성장률 × 한미반도체 점유율이다. 변수는 한화세미텍, 마이크론 발주, 삼성전자향 수주다.",
            "그래도 HBM 세대가 올라갈수록 장비 난이도가 올라간다. 12단→16단에서 본딩 정밀도와 생산성이 핵심이다.",
        ],
        kind="key",
    )

    n.h2("1) 2Q26 실적")
    n.table(
        ["항목", "숫자", "의미"],
        [
            ["매출", "2,511억원  ·  +39.5% YoY", "HBM4 양산 본격화"],
            ["영업이익", "1,303억원  ·  +51.0% YoY", "레버리지 확인"],
            ["영업이익률", "51.9%", "사상 최대 분기 실적"],
        ],
        col_widths=[3.4, 6.2, 8.0],
    )
    n.p("특히 메모리 업체들이 HBM4E를 올해 말~내년 초 준비하면서 12단·16단용 차세대 TC 본더 수요가 이어질 가능성이 높습니다.")

    n.h2("2) 증권사 시각 — BofA vs 상상인")
    n.table(
        ["출처", "포지션 / TP", "산식·논리"],
        [
            ["BofA 2월", "Neutral → Buy, TP 30만", "전환의 시작"],
            ["BofA 4월", "30만 → 42만", "2028년 EPS × 47배. 과거 평균 PER ~60배보다 보수적"],
            ["BofA 이후 보도", "50만까지 제시로 보도", "하반기 매출 약 5,200억, 실적 회복 강하게 전망"],
            ["상상인 7/10", "TP 38만", "27~28년 평균 EPS 6,057원 × PER 62.4배 (직전연도 56.7배 +10%)"],
            ["같은 47배 비교", "상상인 이익 기준 32만", "BofA 42만 vs 상상인 32만 — 차이 = 이익 추정치"],
        ],
        col_widths=[3.6, 5.4, 8.6],
    )
    n.p("BofA의 주장: “2028년 성장성을 보되, 밸류에이션은 오히려 보수적으로 적용했다.” 과거 평균 PER 약 60배 → BofA 적용 47배.")
    n.p("※ 한미반도체 TP 시나리오의 세부 전개는 강의에서 다룹니다.", size=10.5, color=AMBER, bold=True)

    n.h3("BofA 핵심 논리 다섯 가지")
    n.bullet("TCB 시장 지배력")
    n.bullet("SK하이닉스 → Micron → 중국 OSAT 등 고객 다변화")
    n.bullet("삼성전자 공급 가능성")
    n.bullet("용인 SK하이닉스 + 삼성 P5 신규 팹")
    n.bullet("2028년부터 신규 메가팹의 TCB 수요 본격화")

    n.h2("3) 성장 옵션 — HBM 하나에만 의존하지 않는다면")
    n.p("“HBM 하나에만 의존하지 않는다”가 증명된다면, 기존 한미반도체 = HBM → TC 본더 구조에서 시장이 넓어집니다.")
    n.table(
        ["신규 성장축", "내용", "의미"],
        [
            ["Logic TC 본더", "CoWoS 외주화 확대. OSAT가 로직용 TCB 필요", "한미반도체 신규 시장"],
            ["HBF TC 본더", "HBF = High Bandwidth Flash, NAND 기반 AI 메모리", "초기 시장이지만 장기 옵션"],
            ["MSVP", "대면적 기판을 자르는 장비. PLP(Panel Level Packaging) 확대", "MSVP 매출 빠른 증가 가능"],
        ],
        col_widths=[3.6, 8.0, 6.0],
    )
    n.p("결국 질문은 이것입니다. Logic·HBF·PLP까지 성공한다는 ‘성장 옵션’을 얼마나 반영할 것인가?")

    n.h2("4) 변수는 한화세미텍, 마이크론, 삼성향")
    n.p("한미반도체의 리스크 = 한화세미텍입니다.")
    n.rich(
        [
            ("과거:  ", True),
            "SK하이닉스 HBM  →  한미반도체 TCB",
        ]
    )
    n.rich(
        [
            ("앞으로:  ", True),
            "한화세미텍과 경쟁 / 공급망 다변화가 진행될 가능성",
        ]
    )
    n.callout(
        "시장이 앞으로 확인하려는 것",
        ["TCB 시장 성장률  ×  한미반도체 점유율"],
        kind="blue",
    )

    n.h2("5) 그래도 한미가 유리한 이유 — 그리고 기본 가정")
    n.flow(["HBM3E", "HBM4", "HBM4E", "12단 → 16단 이상"])
    n.p("적층수가 증가할수록 본딩 정밀도와 생산성이 중요해집니다.")
    n.h3("기본 가정")
    n.bullet("HBM4/HBM4E 수요는 강함")
    n.bullet("그러나 한화세미텍 점유율 상승")
    n.bullet("한미반도체 TCB 점유율 55~60%")
    n.bullet("Micron/중국 고객 증가")
    n.bullet("삼성 공급은 일부 장비부터")
    n.bullet("영업이익률은 45~50% 유지")

    n.h2("6) 확인 포인트")
    n.bullet("한화세미텍 점유율 상승 속도 vs 한미 55~60% 방어.")
    n.bullet("마이크론 발주, 중국 OSAT, 삼성전자향 일부 장비 공급의 가시화.")
    n.bullet("HBM4E 12단·16단 차세대 TCB 매출 인식 시점.")
    n.bullet("Logic TCB / MSVP가 ‘옵션’에서 ‘실적’으로 넘어가는지.")

    # ── 4. Wonik IPS ───────────────────────────────────────
    n.h1("원익IPS", num="4.")
    n.callout(
        "한 줄 결론",
        [
            "주류: 2Q 실적은 부진했지만 수주잔고 증가 → 하반기 이익 레벨업.",
            "BNK만 차별적: 전방 수요 둔화 → 주가 정점 통과 가능성. 실적 개선을 부정하는 것이 아니라, 이미 주가에 상당 부분 반영됐다는 경고.",
            "논쟁의 핵심은 실적 개선 여부가 아니라 2027년까지 성장이 지속되느냐다.",
        ],
        kind="key",
    )

    n.h2("1) 2Q26 이후 증권사 시각")
    n.bullet("8/7 전후 대부분 증권사: 목표주가를 하향하더라도 긍정적 톤 유지.")
    n.bullet("핵심 근거: 수주잔고 증가 + 3Q~4Q 매출 인식 확대.")
    n.bullet("BNK만 차별적 시각: 전방 수요 둔화 → 주가 정점 통과 가능성.")
    n.p("현재 논쟁의 핵심 = 실적 개선 여부가 아니라 2027년까지 성장 지속 여부.")

    n.h2("2) 왜 주류는 긍정적인가")
    n.p("수주가 아직 꺾이지 않았기 때문입니다.")
    n.flow(["1Q 수주잔고 약 4,000억", "2Q 신규 수주도 견조", "3Q~4Q 수주잔고 매출화", "하반기 매출·OP 증가"])
    n.callout("핵심", ["2Q 실적 부진  ≠  수주 사이클 훼손"], kind="blue")

    n.h2("3) 하반기가 중요하다")
    n.table(
        ["시점", "내용", "봐야 할 것"],
        [
            ["1H26", "수주 증가, 실적은 기대보다 부진", "이미 확인"],
            ["3Q26", "수주잔고 본격 매출화, 이익 증가폭 확대 기대", "OP 증가폭이 커지는가"],
            ["4Q26", "매출 증가 + 영업레버리지, 실적 레벨업 여부 확인", "매출↑와 OPM↑가 동시에 나오는가"],
        ],
        col_widths=[2.6, 8.4, 6.6],
    )

    n.h2("4) 2027년이 진짜 승부처")
    n.callout(
        "Bull Case",
        [
            "DRAM + NAND + Foundry + CXMT → 투자 지속 → 신규 수주 증가",
            "→ 2027년 매출·이익 증가 → 현재 Valuation 부담 완화",
        ],
        kind="bull",
    )
    n.callout(
        "Bear Case (BNK)",
        [
            "CAPEX 증가율 둔화 → 신규 수주 둔화 → 2027년 이익 성장 제한",
            "→ 현재 주가 선반영 부담",
        ],
        kind="bear",
    )

    n.h2("5) BNK의 경고가 중요한 이유")
    n.p("“실적 개선을 부정하는 것이 아니라, 실적 개선이 이미 주가에 상당 부분 반영됐다.”")
    n.rich([("주류 경로:  ", True), "수주 증가  →  실적 증가  →  주가 상승"])
    n.rich(
        [
            ("BNK 경로:  ", True, RED),
            "수주 증가  →  실적 증가 기대  →  주가 선반영  →  이후 CAPEX 둔화",
        ]
    )
    n.callout(
        "장비주의 타이밍",
        ["실적 정점  ≠  주가 정점.  장비주는 실적보다 주가가 먼저 움직일 수 있습니다."],
        kind="note",
    )

    n.h2("6) 앞으로 딱 3개만 확인")
    n.table(
        ["#", "확인 항목", "YES이면", "NO이면"],
        [
            ["①", "3Q26 영업이익 증가폭 확대?", "주류 전망 우세", "하반기 레벨업 지연"],
            ["②", "4Q26 매출 증가 + OPM 개선 동시 발생?", "실적 레벨업 확인", "레버리지 실패"],
            ["③", "2027년 신규 수주가 다시 증가?", "성장 지속", "BNK 정점론 현실화"],
        ],
        col_widths=[1.4, 7.2, 4.6, 4.4],
    )
    n.p("3개 모두 YES → 주류 전망 우세.  수주 둔화 + OPM 개선 실패 → BNK 정점론 현실화.")

    n.h2("7) 목표주가 시나리오")
    n.p("8/8일 리포트 목표주가: 8.8만, 14만, 15.3만, 16.7만, 19만.")
    n.table(
        ["방법", "가정", "산출"],
        [
            ["증권사 TP 평균 + 안전마진", "평균 14.8만 × 안전마진 20~30%", "10.4만 ~ 11.8만원"],
            ["이익 믹스 × PER (BNK 제외)", "26년 50% + 27년 50% 이익 추정 × PER 20~25배", "8.4~10.5 / 9.6~12 / 9.8~12.2 / 11.9~14.9만"],
        ],
        col_widths=[4.6, 7.2, 5.8],
    )
    n.p("즉, 안전마진을 두고 보면 10만원 전후가 공통 구간입니다. 상단은 수주가 2027까지 이어진다는 확신이 붙을 때 열립니다.")

    # ── 5. SK ──────────────────────────────────────────────
    n.h1("SK 지주사 · SK에코플랜트", num="5.")
    n.callout(
        "한 줄 결론",
        [
            "대신: SK할인율 좁혀질 이유가 늘었다. TP 88만원 유지. NAV 81.7조 × 할인율 41%.",
            "에코플랜트는 반도체 팹 + AI DC + 메모리 유통을 동시에 먹는다. 표의 2.1조는 2Q OP 연환산·수주잔고 대비 상당히 낮다.",
            "에코플랜트만 재평가해도 TP 96~106만 산출 가능. 기본은 여전히 SK스퀘어(하이닉스)와 SK이노베이션 가치다.",
        ],
        kind="key",
    )

    n.h2("1) 대신 — SK할인율 좁혀질 이유가 늘었다, TP 88만원 유지")
    n.table(
        ["항목", "2Q26"],
        [
            ["매출액", "42조 1,247억원"],
            ["영업이익", "4조 8,412억원  (+2,212.5% YoY)"],
            ["코멘트", "포트폴리오 리밸런싱을 마친 자회사 이익 체력 확인"],
        ],
        col_widths=[4.4, 13.2],
    )
    n.p("SK에코플랜트가 분기 최대 매출을 달성하며 AI DC 관련 실적 반영이 본격화됐습니다.")

    n.h3("SOTP (대신)")
    n.table(
        ["구성", "금액"],
        [
            ["자체사업 및 로열티 영업가치", "13.0조원"],
            ["상장자회사 지분가치", "73.4조원"],
            ["비상장 자회사 지분가치", "3.8조원"],
            ["별도 순차입금 (차감)", "△ 8.5조원"],
            ["NAV", "81.7조원"],
            ["Target 할인율", "41%"],
            ["함의 (81.6조 × 59%)", "약 48.2조원  →  TP 88만 유지"],
        ],
        col_widths=[8.0, 9.6],
    )

    n.h2("2) SK에코플랜트 — 무엇이 다른 회사인가")
    n.p("반도체·AI 인프라의 종합 솔루션 회사입니다. 반도체 공장 건설 + AI 데이터센터 + 반도체 소재/가스 + 메모리 유통/재활용.")
    n.callout(
        "한 줄 정의",
        ["SK하이닉스의 반도체 투자와 AI 데이터센터 투자를 동시에 먹는 회사"],
        kind="blue",
    )

    n.h3("4대 사업 구조 (26년 상반기 매출 비중)")
    n.table(
        ["사업", "내용", "비중"],
        [
            ["Asset Lifecycle", "메모리 모듈 유통 · IT자산 재활용", "45%"],
            ["Hi-Tech", "반도체 Fab · AI 데이터센터 건설", "33%"],
            ["Solution", "기존 건설 · 인프라", "18%"],
            ["Gas & Material", "반도체용 가스 · 소재", "4%"],
        ],
        col_widths=[4.4, 9.6, 3.6],
    )
    n.p("반도체 관련 매출 ≈ 80%+.")

    n.h2("3) 2Q26 실적 — 왜 급증했나")
    n.table(
        ["항목", "2Q26", "변화"],
        [
            ["매출", "5.15조원", "+68.4% YoY / 분기 최대"],
            ["영업이익", "5,340억원", "+265.7% YoY"],
            ["EBITDA Margin", "11.8%", "+5.1%p YoY"],
            ["수주잔고", "26.8조원", "향후 매출 가시성"],
            ["2Q 신규 수주", "7.8조원", "Hi-Tech 대형 프로젝트 확대"],
        ],
        col_widths=[4.0, 5.0, 8.6],
    )
    n.h3("실적 상승의 3가지 엔진")
    n.bullet("반도체 CAPEX ↑  —  용인 반도체 클러스터 + 청주 M15X → Hi-Tech EPC 매출 증가")
    n.bullet("AI CAPEX ↑  —  울산 AI 데이터센터 등 → AI DC 공정 본격화")
    n.bullet("DRAM·NAND 가격 ↑  —  Essencore 메모리 유통 매출 증가")

    n.h2("4) 표의 2.1조는 너무 낮다")
    n.p("현재 밸류에이션 표: 지분율 71.2%, 장부가 1.271조원, Value 2.106조원.")
    n.callout(
        "숫자 체크",
        [
            "1.271조 × 71.2% = 0.905조 → 표의 2.106조와 맞지 않습니다.",
            "해석: 2.106조를 SK 지분가치(이미 지분율 반영)로 보는 편이 자연스럽습니다. 장부가 1.271조는 별도 장부가일 가능성.",
            "어느 쪽이든 2Q 영업이익 5,340억 × 4 = 연환산 2.14조, 수주잔고 26.8조입니다.",
            "과거의 건설회사와 달리 반도체 EPC + AI DC + 메모리 유통으로 이익 체력이 구조적으로 올라왔다면, 2.1조의 기업가치를 부여하는 것은 상당히 낮습니다.",
        ],
        kind="note",
    )

    n.h3("시나리오 A — 보수")
    n.flow(["영업이익 1.6~2조", "이자·세금·비영업·일회성", "순이익 1.0조 가정", "PER 10배 → 10조", "SK 71.2% → 7.1조"])
    n.p("현재 표 2.1조와 비교하면 약 5조원 차이.")

    n.h3("시나리오 B — 성장·수주 감안")
    n.flow(["순이익 1.5조", "PER 10배 → 15조", "SK 71.2% → 10.7조"])
    n.p("현재 표 2.1조와 비교하면 약 8.6조원 차이. (이노베이션의 가치기여와 비슷한 규모로 볼 수 있음)")

    n.h2("5) 할인율만 안 바꿔도 TP가 올라간다")
    n.p("현재 대신 목표가 산식: 81.6조 × 59% = 48.2조.")
    n.p("SK에코플랜트 재평가로 NAV가 (시나리오 A/B 평균 차이 약 8조를 더해) 89.5조가 되면:")
    n.table(
        ["가정", "산식", "결과", "vs 현재"],
        [
            ["할인율 41% 유지", "89.5조 × 59%", "52.8조", "현재 TP 대비 ~10%"],
            ["할인율 41% → 35%", "89.5조 × 65%", "58.2조", "현재 TP 대비 ~21%"],
        ],
        col_widths=[4.6, 4.6, 3.8, 4.6],
    )
    n.callout(
        "함의",
        [
            "따라서 TP는 96~106만원도 산출 가능합니다.",
            "현재 증권사 목표주가 최고치는 100만원 (6/24 흥국) 이후 97만 1개, 88만 1개, 85만 2개, 73.8만 1개.",
            "안전마진을 충분히 고려해도 긍정적 시각이 가능합니다. 단, 기본은 SK스퀘어(즉 SK하이닉스), SK이노베이션 가치입니다.",
        ],
        kind="key",
    )

    n.h2("6) SK실트론 매각 — NAV +0.5조, 옵션은 남긴다")
    n.p("SK실트론 지분 70.6%를 두산에 2.3조원 매각 공시. 목적은 재무구조 개선 및 투자재원 확보, 초과이익공유 예정. 순자산가치 증가 요인이자 주주환원 재원 확보.")
    n.bullet("SK실트론 가치를 EV/EBITDA 8배로 적용해 지분가치를 1.8조원으로 가정")
    n.bullet("하지만 2.3조원에 매각 → SK의 NAV는 0.5조원 증가")
    n.p("재무구조를 개선하면서 미래 성장가치 옵션을 확보하는 거래로 읽으면 됩니다.")

    n.h2("7) 확인 포인트")
    n.bullet("에코플랜트 Hi-Tech 수주잔고 26.8조가 실제로 매출·마진으로 이어지는가.")
    n.bullet("표의 2.1조가 리포트에서 상향 재평가되는가. (이게 할인율 축소의 재료)")
    n.bullet("실트론 매각 대금의 사용처 — 차입 축소 vs 성장 투자 vs 주주환원.")
    n.bullet("그래도 NAV의 중심은 SK스퀘어(하이닉스)와 이노베이션. 에코플랜트는 ‘할인율 좁힐 이유’다.")

    # ── 6. Atlas ───────────────────────────────────────────
    n.h1("Atlas / Physical AI", num="6.")
    n.callout(
        "한 줄 결론",
        [
            "Atlas의 투자 포인트가 ‘로봇을 얼마나 잘 만드느냐’에서 ‘얼마나 많이, 오래, 실제 공장에서 일하게 하느냐’로 이동했다.",
            "HMGMA·현대글로비스의 실제 배치가 상용화의 가장 중요한 검증 이벤트다. 현대모비스는 로봇 부품 양산 수혜의 핵심 후보.",
        ],
        kind="key",
    )

    n.h2("1) 연구용 로봇 → 공장 투입")
    n.p("상용화 경로가 구체화되고 있습니다.")
    n.flow(["RMAC 학습", "성능검증", "HMGMA·글로비스 파일럿", "실제 생산 투입"])
    n.bullet("2028년 HMGMA 투입")
    n.bullet("2030년 부품 조립공정 확대 목표")

    n.h2("2) 핵심 경쟁력 = Fleet × 가동률 × 데이터")
    n.p("로봇 성능보다, 수천 번 안정적으로 반복 작업하는 것이 제조용 로봇의 핵심입니다.")
    n.p("향후 평가지표: 로봇 성능 → 실제 배치 대수 · 가동시간 · 활용률")
    n.flow(["신뢰성 ↑", "가동률 ↑", "데이터 ↑", "AI 성능 ↑", "추가 배치 ↑"])
    n.callout("Physical AI Data Flywheel", ["학습 → 배치 → 데이터 축적 → 재학습을 빠르게 반복할 수 있는 쪽이 이깁니다."], kind="blue")

    n.h2("3) 현대차그룹 Captive 환경 = 오히려 강력한 경쟁우위")
    n.bullet("HMGMA = 제조 데이터")
    n.bullet("현대글로비스 = 물류 데이터")
    n.bullet("한국 · 미국 · 중국 로봇훈련 거점 + 글로벌 생산거점")
    n.p("학습 → 배치 → 데이터 축적 → 재학습을 빠르게 반복할 수 있습니다.")

    n.h2("4) 수혜주 관점")
    n.h3("① 현대모비스")
    n.bullet("Atlas 1대당 바디 액추에이터 31개")
    n.bullet("11월 시제품 납품 예정")
    n.bullet("Atlas 양산 확대 → 액추에이터 물량 증가")
    n.bullet("로봇 부품 사업의 사업화 가시성 상승")
    n.h3("② 현대글로비스")
    n.bullet("부품 서열(Sequencing)은 반복성·구조화가 높아 초기 상용화에 적합")
    n.bullet("실제 파일럿/배치 확인 시 → 현대차=제조 + 현대글로비스=물류의 Physical AI 초기 생태계 구축")

    n.callout(
        "최종 결론",
        [
            "HMGMA·현대글로비스의 실제 배치 = Atlas 상용화의 가장 중요한 검증 이벤트.",
            "현대모비스는 로봇 부품 양산 수혜의 핵심 후보.",
        ],
        kind="key",
    )

    n.h2("5) 확인 포인트")
    n.bullet("11월 모비스 시제품 납품이 일정대로 나가는가.")
    n.bullet("HMGMA·글로비스 파일럿의 실제 대수·가동시간.")
    n.bullet("2028 투입 → 2030 조립공정 확대 로드맵이 숫자로 구체화되는가.")

    # ── 7. Leeno ───────────────────────────────────────────
    n.h1("보론  ·  리노공업 (파업)", num="7.")
    n.callout(
        "한 줄 결론",
        [
            "파업 직전까지 실적·수익성은 강한 성장 국면. 2Q OPM 51.3%, 컨센 상회.",
            "창사 이래 첫 무기한 총파업(7/23~8/17 약 4주)이 최대 변수. 핵심은 매출 감소보다 마진 훼손.",
            "파업만 마무리되면 peers → 리노공업으로 순환매가 가능한 시기가 올 수 있다.",
        ],
        kind="key",
    )

    n.h2("1) 삼성증권 2Q26 리뷰")
    n.p("시장의 우려를 재차 실적으로 반박. 물량/판가 모두 증가, 영업이익률 50% 초과, 컨센서스 상회. 판가 확대 지속 vs 추후 비용 구조 변화 리스크도 상쇄될 것. 목표주가를 100,000원 (27년 PER 32배)으로 하향.")

    n.h2("2) 2Q26 실적 — 매우 양호")
    n.table(
        ["항목", "숫자"],
        [
            ["매출", "약 1,432억원  ·  YoY +30%"],
            ["영업이익", "약 735억원  ·  YoY +37%"],
            ["영업이익률", "51.3%"],
            ["상반기 영업이익", "1,208억원  ·  +36.7%"],
            ["믹스", "핀·소켓 물량 + ASP 동반 상승. AI/HPC용 고성능 반도체 테스트 수요 견조"],
        ],
        col_widths=[4.4, 13.2],
    )
    n.p("핵심: 파업 직전까지는 실적·수익성 모두 강한 성장 국면.")

    n.h2("3) 최대 변수 — 창사 이래 첫 무기한 총파업")
    n.p("7/23부터 무기한 총파업 → 8/17 현재 약 4주.")
    n.table(
        ["구분", "내용"],
        [
            ["노조 요구", "연 800% 고정상여  ·  영업이익의 15% 성과급"],
            ["회사 측 우려", "인건비가 매출액의 약 30% 수준까지 상승 가능"],
            ["", "생산 차질 → 3Q 매출 감소 가능성"],
            ["", "장기화 시 고객의 경쟁사 전환 가능성"],
        ],
        col_widths=[3.6, 14.0],
        first_col_bold=True,
    )
    n.callout(
        "핵심은 ‘매출 감소’보다 마진 훼손",
        [
            "현재 OPM 50%+  →  임금체계가 구조적으로 바뀌면  →  40%대 중반까지 하락 가능성을 시장이 우려.",
        ],
        kind="bear",
    )
    n.p("파업만 마무리/종료가 잘 되기만 하면 다시 peers → 리노공업으로 순환매도 가능한 시기가 올 수 있습니다.")

    # ── 8. Memory ──────────────────────────────────────────
    n.h1("보론  ·  메모리 밸류에이션 비교 프레임", num="8.")
    n.callout(
        "왜 보론인가",
        [
            "오늘은 NON-삼전닉스 강의입니다. 다만 당일 코멘트에 삼전닉스·마이크론·샌디스크·키옥시아·CXMT 비교가 있어, 녹화 중 질문 대비로 숫자를 한곳에 모아 둡니다.",
            "SK 지주사 이야기의 기본이 SK스퀘어(하이닉스)이기도 합니다.",
        ],
        kind="note",
    )

    n.h2("1) 해외 피어 — 8/14 종가 기준")
    n.table(
        ["종목", "가격", "밸류에이션"],
        [
            ["마이크론", "971.66달러", "Forward 12개월 PER 7.8배  /  CY27 EPS(컨센) 150달러 → PER 6.5배"],
            ["Sandisk", "1,644.11달러", "FY27.1Q 가이던스 EPS 45달러, QoQ +10%/+5%/+5% 가정 → FY27 EPS 201달러, PER 8.2배"],
            ["키옥시아", "7~9월 가이던스 ×4", "12개월 선행 PER ~6배 중반. 메모리 업황 개선 기대에 오후장 +10%대"],
            ["CXMT", "공모가 7배 상승", "키옥시아 +15% / CXMT +10% 가량 동반 강세"],
        ],
        col_widths=[3.2, 4.4, 10.0],
    )

    n.h2("2) SK하이닉스 ADR vs 본주")
    n.table(
        ["항목", "숫자"],
        [
            ["SK하이닉스 ADR", "166.33달러  (235.9만원 = 1,418.5원/달러 기준)"],
            ["ADR 26년 / 27년 PER", "6.8배 / 5.4배"],
            ["마이크론 대비", "−17%  (과거 −20~−50%)"],
            ["본주 대비 프리미엄", "43%"],
            ["정상 프리미엄 +20% 가정 (TSMC 15% 수준)", "본주는 현시점 196만원은 되어야"],
            ["최근 실제 30% 이상 프리미엄 감안, 30~35% 적용", "본주 175~181만원"],
        ],
        col_widths=[8.4, 9.2],
    )

    n.h2("3) 본주 밸류 — 8/14 종가 흐름")
    n.table(
        ["종목", "가격", "26년 PER", "27년 PER", "26년 OP / EPS", "27년 OP / EPS"],
        [
            ["SK하이닉스", "164.5만원", "4.8배", "3.8배", "266조 / 346K", "392조 / 437K"],
            ["삼성전자", "27.45만원", "5.7배", "4.1배", "391조 / 47.9K", "549조 / 67.2K"],
        ],
        col_widths=[2.8, 2.6, 2.4, 2.4, 3.8, 3.8],
    )

    n.h2("4) 두 가지 접근")
    n.h3("접근 A — ADR 프리미엄")
    n.bullet("최근 ADR 프리미엄을 인정한다면 감안 정도만: 하이닉스 170~180만원, 정상적 프리미엄이면 196만원.")
    n.bullet("같은 흐름이라면 삼전 29만원대.")

    n.h3("접근 B — 밸류에이션/성장률 (보수)")
    n.p("26년 실적 대비 27년 성장이 없다고 단순화. 26년 PER 6~7배 (과거 사이클 주식 4~8배) 적용.")
    n.table(
        ["종목", "PER 6배", "PER 7배"],
        [
            ["SK하이닉스 (26년 EPS 346K)", "208만원", "242만원"],
            ["삼성전자 (26년 EPS 47.9K)", "28.7만원", "33.5만원"],
        ],
        col_widths=[6.4, 5.6, 5.6],
    )
    n.p("산출 가능 구간: 하이닉스 208만~242만, 삼전 28.7만~33.5만원.")

    # ── Close ──────────────────────────────────────────────
    n.h1("클로징", num="9.")
    n.callout(
        "오늘 강의에서 가져갈 세 문장",
        [
            "1) 소부장: 2Q 숫자보다 수주가 매출로 바뀌는 하반기, 그리고 2027 수주가 다시 증가하는지가 승부처다. 테스는 BSD 퀄, 한미는 점유율, 원익은 3개 체크리스트.",
            "2) SK: 할인율이 좁혀질 이유가 늘었다. 에코플랜트 2.1조는 낮고, 재평가만으로도 96~106만이 나온다. 그래도 기본은 하이닉스와 이노베이션.",
            "3) Atlas: 로봇 성능 논쟁은 끝나가고, 공장에서 몇 대가 몇 시간 일하느냐가 시작이다. 모비스는 부품, 글로비스는 서열·물류 데이터.",
        ],
        kind="key",
    )
    n.h2("클로징 멘트 (녹화용)")
    n.p("모든 게 미래 어느 시점의 최종수요입니다. 그걸 정확히 아무도 예측하지 못하니, AI 모델·데이터센터·GPU 전망에 기대는 것이고, 자금이 딸리니 서로 묶어서 가는 것입니다. 금융이 붙으면 속도가 빨라지지만 Dark GPU라는 역방향 레버리지도 있습니다. 우려와 병목이 해소되기 전까지 양면적 시각은 계속됩니다. 그래서 오늘은 그 양면 위에서, 삼전닉스가 아니라 수주·할인율·배치를 봤습니다.")

    n.spacer(8)
    n.p("— 8월 18일 NON-삼전닉스 강의노트. 원문 퀵코멘트(10:25~17:05)를 강의 순서로 재구성.", size=9.5, color=GRAY, align="right")

    n.save(OUT_PATH)
    print(f"Wrote {OUT_PATH} ({OUT_PATH.stat().st_size} bytes)")


if __name__ == "__main__":
    build()
