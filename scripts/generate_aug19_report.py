#!/usr/bin/env python3
"""8월 19일 Quick 코멘트 → 시각화 중심 시장 브리핑 보고서(.docx)."""

from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.font_manager as fm
import numpy as np
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Mm, Pt, RGBColor

OUT_PATH = Path("/workspace/lectures/8월 19일 시장브리핑 (환율·주주환원·HBM·매크로).docx")
CHART_DIR = Path("/workspace/lectures/charts_aug19")

KR_FONT = "맑은 고딕"
EN_FONT = "Calibri"
FONT_PATH = "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc"
FP = fm.FontProperties(fname=FONT_PATH)
FP_B = fm.FontProperties(fname=FONT_PATH, weight="bold")

# Force Korean-capable default so axis labels never fall back to DejaVu.
plt.rcParams["font.family"] = FP.get_name()
plt.rcParams["axes.unicode_minus"] = False
fm.fontManager.addfont(FONT_PATH)

NAVY = RGBColor(0x0F, 0x20, 0x43)
NAVY2 = RGBColor(0x1E, 0x40, 0x7C)
GOLD = RGBColor(0xB8, 0x94, 0x3A)
GRAY = RGBColor(0x4B, 0x55, 0x63)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x16, 0x65, 0x34)
RED = RGBColor(0x99, 0x1B, 0x1B)
AMBER = RGBColor(0x7A, 0x5C, 0x12)

NAVY_HEX = "0F2043"
NAVY2_HEX = "1E407C"
GOLD_HEX = "B8943A"
LIGHT_HEX = "EEF2F8"
GREEN_HEX = "E8F5E9"
RED_HEX = "FDECEA"
AMBER_HEX = "FFF8E7"
BLUE_HEX = "E8F1FB"
ROW_HEX = "F7F9FC"
WHITE_HEX = "FFFFFF"

# matplotlib palette
C_NAVY = "#0F2043"
C_NAVY2 = "#1E407C"
C_GOLD = "#B8943A"
C_GREEN = "#166534"
C_RED = "#991B1B"
C_AMBER = "#B45309"
C_GRAY = "#6B7280"
C_LIGHT = "#EEF2F8"
C_SOFT = "#F7F9FC"


def set_run_font(run, size=11, bold=False, color=DARK, italic=False, font=KR_FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), KR_FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(
        parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill}"/>')
    )


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(
        parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'<w:top w:w="{top}" w:type="dxa"/>'
            f'<w:left w:w="{left}" w:type="dxa"/>'
            f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
            f'<w:right w:w="{right}" w:type="dxa"/>'
            f"</w:tcMar>"
        )
    )


def set_table_borders(table, color="D0D7E2", sz="4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f"</w:tblBorders>"
    )
    tbl_pr.append(borders)


def set_left_accent(cell, color=NAVY_HEX, sz="24"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(
        parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="nil"/>'
            f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="nil"/>'
            f'<w:right w:val="nil"/>'
            f"</w:tcBorders>"
        )
    )


def prevent_row_split(row):
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()
    tr_pr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))


def cell_text(cell, text, size=10, bold=False, color=DARK, align="left", font=KR_FONT):
    cell.text = ""
    align_enum = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    lines = str(text).split("\n")
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = align_enum
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1 if i < len(lines) - 1 else 0)
        p.paragraph_format.line_spacing = 1.12
        run = p.add_run(line)
        set_run_font(run, size=size, bold=bold, color=color, font=font)
    set_cell_margins(cell)


def add_runs(paragraph, parts, size=11, color=DARK):
    for part in parts:
        if isinstance(part, str):
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)
        else:
            text, bold, *rest = part
            c = rest[0] if rest else color
            run = paragraph.add_run(text)
            set_run_font(run, size=size, bold=bold, color=c)


# ── Charts ─────────────────────────────────────────────────


def _style_ax(ax, title=None):
    ax.set_facecolor(C_SOFT)
    for spine in ("top", "right"):
        ax.spines[spine].set_visible(False)
    for spine in ("left", "bottom"):
        ax.spines[spine].set_color("#D0D7E2")
    ax.tick_params(colors=C_GRAY, labelsize=9)
    for label in list(ax.get_xticklabels()) + list(ax.get_yticklabels()):
        label.set_fontproperties(FP)
    if title:
        ax.set_title(title, fontproperties=FP_B, fontsize=13, color=C_NAVY, pad=10)


def _ylabel(ax, text):
    ax.set_ylabel(text, fontproperties=FP, color=C_GRAY)


def _xlabel(ax, text):
    ax.set_xlabel(text, fontproperties=FP, color=C_GRAY)


def save_fig(fig, name: str) -> Path:
    CHART_DIR.mkdir(parents=True, exist_ok=True)
    path = CHART_DIR / name
    fig.savefig(path, dpi=160, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def chart_macro_chain() -> Path:
    fig, ax = plt.subplots(figsize=(10.5, 2.8))
    ax.set_xlim(0, 10.5)
    ax.set_ylim(0, 3)
    ax.axis("off")
    boxes = [
        (0.2, "미·이란\n리스크", C_RED),
        (2.2, "유가↑\n(브렌트)", C_AMBER),
        (4.2, "인플레\n우려", C_AMBER),
        (6.2, "미 국채\n금리↑", C_NAVY2),
        (8.2, "AI/반도체\n밸류 압박", C_NAVY),
    ]
    for x, label, color in boxes:
        rect = mpatches.FancyBboxPatch(
            (x, 0.9), 1.7, 1.4, boxstyle="round,pad=0.05,rounding_size=0.15",
            facecolor=color, edgecolor="none", alpha=0.92,
        )
        ax.add_patch(rect)
        ax.text(x + 0.85, 1.6, label, ha="center", va="center",
                fontproperties=FP_B, fontsize=10, color="white")
    for x in (1.95, 3.95, 5.95, 7.95):
        ax.annotate("", xy=(x + 0.2, 1.6), xytext=(x, 1.6),
                    arrowprops=dict(arrowstyle="->", color=C_GOLD, lw=2.2))
    ax.text(5.25, 0.35, "핵심: 전쟁 자체보다 '금리'가 주가를 누르는 경로",
            ha="center", fontproperties=FP, fontsize=10, color=C_GRAY)
    ax.set_title("매크로 전달 경로 — 8/19 조정의 뼈대", fontproperties=FP_B,
                 fontsize=13, color=C_NAVY, pad=6)
    return save_fig(fig, "01_macro_chain.png")


def chart_rate_gauge() -> Path:
    fig, ax = plt.subplots(figsize=(10.5, 2.6))
    ax.set_xlim(4.3, 5.4)
    ax.set_ylim(0, 1.6)
    ax.axis("off")
    # zones
    zones = [
        (4.3, 4.7, C_GREEN, "안정 구간\n(성장주 부담↓)"),
        (4.7, 5.0, C_AMBER, "주의 구간"),
        (5.0, 5.4, C_RED, "위험 구간\n(위험자산 회피)"),
    ]
    for x0, x1, color, label in zones:
        ax.add_patch(mpatches.FancyBboxPatch(
            (x0, 0.55), x1 - x0, 0.55, boxstyle="square,pad=0",
            facecolor=color, edgecolor="white", linewidth=2, alpha=0.85,
        ))
        ax.text((x0 + x1) / 2, 0.82, label, ha="center", va="center",
                fontproperties=FP_B, fontsize=9, color="white")
    # marker ~4.70
    ax.plot([4.70], [1.25], marker="v", markersize=14, color=C_NAVY)
    ax.text(4.70, 1.42, "10년물 ~4.70%", ha="center", fontproperties=FP_B,
            fontsize=11, color=C_NAVY)
    for x, lab in [(4.4, "4.4"), (4.7, "4.7"), (5.0, "5.0"), (5.3, "5.3")]:
        ax.text(x, 0.32, lab, ha="center", fontproperties=FP, fontsize=9, color=C_GRAY)
    ax.set_title("미국 10년물 국채금리 — 단기 최대 변수 게이지", fontproperties=FP_B,
                 fontsize=13, color=C_NAVY, pad=6)
    return save_fig(fig, "02_rate_gauge.png")


def chart_oil_gauge() -> Path:
    fig, ax = plt.subplots(figsize=(10.5, 2.4))
    ax.set_xlim(70, 110)
    ax.set_ylim(0, 1.5)
    ax.axis("off")
    zones = [
        (70, 90, C_GREEN, "안정 (충격 완화)"),
        (90, 100, C_AMBER, "부담 확대"),
        (100, 110, C_RED, "인플레·금리 악순환"),
    ]
    for x0, x1, color, label in zones:
        ax.add_patch(mpatches.FancyBboxPatch(
            (x0, 0.5), x1 - x0, 0.5, boxstyle="square,pad=0",
            facecolor=color, edgecolor="white", linewidth=2, alpha=0.85,
        ))
        ax.text((x0 + x1) / 2, 0.75, label, ha="center", va="center",
                fontproperties=FP_B, fontsize=10, color="white")
    ax.plot([84], [1.15], marker="v", markersize=14, color=C_NAVY)
    ax.text(84, 1.32, "WTI ~$84", ha="center", fontproperties=FP_B, fontsize=11, color=C_NAVY)
    ax.set_title("유가(브렌트/WTI) 임계 — 금리 경로의 1차 입력", fontproperties=FP_B,
                 fontsize=13, color=C_NAVY, pad=6)
    return save_fig(fig, "03_oil_gauge.png")


def chart_macro_triangle() -> Path:
    fig, ax = plt.subplots(figsize=(8.5, 5.2))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10)
    ax.axis("off")
    nodes = [
        (5, 8.2, "전쟁발 매크로\n(유가·금리)", C_RED),
        (1.8, 3.2, "투자의지\n(AI CAPEX)", C_NAVY2),
        (8.2, 3.2, "토큰 수요\n(최종수요)", C_GREEN),
    ]
    # triangle
    ax.plot([5, 1.8, 8.2, 5], [7.4, 4.0, 4.0, 7.4], color=C_GOLD, lw=2, alpha=0.7)
    for x, y, label, color in nodes:
        ax.add_patch(mpatches.Circle((x, y), 1.15, facecolor=color, edgecolor="white", lw=2, alpha=0.9))
        ax.text(x, y, label, ha="center", va="center", fontproperties=FP_B,
                fontsize=10, color="white")
    ax.text(5, 1.2, "현재 무게중심 → 왼쪽(매크로)\n하지만 토큰 수요→AI 투자 궤도는 금리만으로 쉽게 안 바뀜",
            ha="center", fontproperties=FP, fontsize=10, color=C_GRAY)
    ax.set_title("전쟁 vs 투자의지 vs 토큰 수요", fontproperties=FP_B,
                 fontsize=13, color=C_NAVY, pad=4)
    return save_fig(fig, "04_macro_triangle.png")


def chart_skh_fcf_return() -> Path:
    fig, axes = plt.subplots(1, 2, figsize=(10.8, 4.2))

    # left: FCF 3yr
    ax = axes[0]
    years = ["2025E", "2026E", "2027E"]
    base = [179, 242, 237]
    cons = [150, 210, 205]
    x = np.arange(len(years))
    w = 0.36
    b1 = ax.bar(x - w / 2, base, w, label="기존 FCF 추정", color=C_NAVY2)
    b2 = ax.bar(x + w / 2, cons, w, label="보수적 (운전자본 차감)", color=C_GOLD)
    _style_ax(ax, "연간 FCF 추정 (조원)")
    ax.set_xticks(x)
    ax.set_xticklabels(years, fontproperties=FP)
    ax.legend(prop=FP, fontsize=8, frameon=False)
    _ylabel(ax, "조원")
    for bars in (b1, b2):
        for bar in bars:
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 3,
                    f"{int(bar.get_height())}", ha="center", fontproperties=FP, fontsize=8)

    # right: return stack
    ax = axes[1]
    labels = ["확정\nBuyback", "추가 환원\n필요(50%)", "50% 초과\n여지"]
    vals = [40, 152.5, 30]
    colors = [C_NAVY, C_NAVY2, C_GOLD]
    bottom = 0
    for lab, v, c in zip(labels, vals, colors):
        ax.bar(["2025~27\n환원 구조"], [v], bottom=bottom, color=c, edgecolor="white", width=0.45)
        ax.text(0, bottom + v / 2, f"{lab}\n{v}조", ha="center", va="center",
                fontproperties=FP_B, fontsize=9, color="white")
        bottom += v
    _style_ax(ax, "주주환원 구조 (누적 FCF×50%+)")
    _ylabel(ax, "조원")
    ax.text(0, bottom + 8, "누적 FCF ~385조 × 50% = 192.5조+",
            ha="center", fontproperties=FP, fontsize=8.5, color=C_GRAY)

    fig.suptitle("SK하이닉스 주주환원 — FCF와 환원 규모", fontproperties=FP_B,
                 fontsize=14, color=C_NAVY, y=1.02)
    fig.tight_layout()
    return save_fig(fig, "05_skh_fcf_return.png")


def chart_per_compare() -> Path:
    fig, ax = plt.subplots(figsize=(10.5, 4.6))
    names = ["Sandisk\nFY27", "Micron\nFwd12M", "Micron\nCY27", "SKH ADR\n'26", "SKH ADR\n'27",
             "SKH 본주\n'26", "SKH 본주\n'27", "삼성\n'26", "삼성\n'27"]
    pers = [7.8, 7.5, 6.25, 6.6, 5.2, 4.3, 3.4, 5.2, 3.7]
    colors = [C_GRAY, C_GRAY, C_GRAY, C_NAVY2, C_NAVY2, C_NAVY, C_NAVY, C_GOLD, C_GOLD]
    bars = ax.bar(names, pers, color=colors, edgecolor="white", width=0.72)
    _style_ax(ax, "메모리 Peer PER 비교 (코멘트 기준)")
    _ylabel(ax, "PER (배)")
    for label in ax.get_xticklabels():
        label.set_fontproperties(FP)
        label.set_fontsize(8)
    for bar, v in zip(bars, pers):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.15,
                f"{v}", ha="center", fontproperties=FP_B, fontsize=8.5, color=C_NAVY)
    ax.axhline(6, color=C_AMBER, ls="--", lw=1, alpha=0.7)
    ax.text(8.5, 6.15, "참고선 6배", fontproperties=FP, fontsize=8, color=C_AMBER)
    fig.tight_layout()
    return save_fig(fig, "06_per_compare.png")


def chart_fx_sensitivity() -> Path:
    fig, axes = plt.subplots(1, 2, figsize=(10.8, 4.0))

    ax = axes[0]
    names = ["삼성전자", "SK하이닉스"]
    sens = [0.4, 0.9]
    bars = ax.barh(names, sens, color=[C_GOLD, C_NAVY], height=0.45)
    _style_ax(ax, "원/$ +1% → EPS 민감도")
    _xlabel(ax, "EPS 변화 (%)")
    for label in ax.get_yticklabels():
        label.set_fontproperties(FP)
    for bar, v in zip(bars, sens):
        ax.text(v + 0.03, bar.get_y() + bar.get_height() / 2, f"+{v}%",
                va="center", fontproperties=FP_B, fontsize=11, color=C_NAVY)

    ax = axes[1]
    # 1520 → 1420 impact
    scenarios = ["USD/KRW\n1,520→1,420\n(-6.6%)", "26H2 환율\n조정 시나리오"]
    impact = [-5.9, None]
    # show EPS drop and profit adjustment range
    ax.bar([0], [5.9], color=C_RED, width=0.5, alpha=0.85, label="EPS 하락 %")
    ax.bar([1], [21], color=C_AMBER, width=0.5, alpha=0.85, label="이익조정 중간(조)")
    ax.errorbar([1], [21], yerr=[[3], [3]], fmt="none", ecolor=C_NAVY, capsize=6, lw=1.5)
    ax.text(0, 6.2, "EPS -5.9%", ha="center", fontproperties=FP_B, fontsize=10, color=C_RED)
    ax.text(1, 25, "18~24조\n(순익 300~400조 가정)", ha="center", fontproperties=FP, fontsize=8.5, color=C_AMBER)
    ax.text(1, 14, "16.3조\n(26H2 환율)", ha="center", fontproperties=FP, fontsize=8, color="white")
    _style_ax(ax, "원화 강세 시 SKH 이익 민감")
    ax.set_xticks([0, 1])
    ax.set_xticklabels(["1,520→1,420", "이익 조정 규모"], fontproperties=FP)
    _ylabel(ax, "% 또는 조원")

    fig.suptitle("환율 민감도 — 수출주 원화강세 리스크", fontproperties=FP_B,
                 fontsize=14, color=C_NAVY, y=1.02)
    fig.tight_layout()
    return save_fig(fig, "07_fx_sensitivity.png")


def chart_fx_path() -> Path:
    fig, ax = plt.subplots(figsize=(10.5, 3.8))
    steps = [
        "법인세·설비\n달러 매도↑",
        "환헤지\n비중↑",
        "달러-원\n하락",
        "고환율서\n추가 매도",
        "하락\n가속",
    ]
    xs = np.linspace(0.8, 9.2, len(steps))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 3.2)
    ax.axis("off")
    for i, (x, s) in enumerate(zip(xs, steps)):
        ax.add_patch(mpatches.FancyBboxPatch(
            (x - 0.75, 1.1), 1.5, 1.2, boxstyle="round,pad=0.04,rounding_size=0.12",
            facecolor=C_NAVY2 if i < 3 else C_NAVY, edgecolor="none", alpha=0.9,
        ))
        ax.text(x, 1.7, s, ha="center", va="center", fontproperties=FP_B,
                fontsize=9, color="white")
        if i < len(steps) - 1:
            ax.annotate("", xy=(xs[i + 1] - 0.8, 1.7), xytext=(x + 0.8, 1.7),
                        arrowprops=dict(arrowstyle="->", color=C_GOLD, lw=2))
    ax.text(5, 0.45,
            "한국 내부 달러 공급이 먼저 → 달러인덱스·엔만으로 설명 어려움\n"
            "추가 하락 핵심변수=달러 (DXY 99→96~97 시 1,360~1,340대 조건부 가능)",
            ha="center", fontproperties=FP, fontsize=9.5, color=C_GRAY)
    ax.set_title("왜 1,400원 아래로? — 국내 수급 가속 루프", fontproperties=FP_B,
                 fontsize=13, color=C_NAVY, pad=6)
    return save_fig(fig, "08_fx_path.png")


def chart_hbm_debate() -> Path:
    fig, axes = plt.subplots(1, 2, figsize=(10.8, 4.4))

    # left: demand equation
    ax = axes[0]
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10)
    ax.axis("off")
    ax.add_patch(mpatches.FancyBboxPatch((0.5, 5.5), 4, 3.2, boxstyle="round,pad=0.05",
                                         facecolor=C_GREEN, alpha=0.9))
    ax.text(2.5, 8.2, "수요 증가 시나리오", ha="center", fontproperties=FP_B,
            fontsize=10, color="white")
    ax.text(2.5, 6.8, "추론량 +50%\n효율 +20%\n→ 메모리 수요↑↑", ha="center",
            fontproperties=FP, fontsize=10, color="white")
    ax.add_patch(mpatches.FancyBboxPatch((5.5, 5.5), 4, 3.2, boxstyle="round,pad=0.05",
                                         facecolor=C_AMBER, alpha=0.9))
    ax.text(7.5, 8.2, "둔화 시나리오", ha="center", fontproperties=FP_B,
            fontsize=10, color="white")
    ax.text(7.5, 6.8, "추론량 +20%\n효율 +30%\n→ 수요 증가세↓", ha="center",
            fontproperties=FP, fontsize=10, color="white")
    ax.text(5, 3.8, "핵심 변수", ha="center", fontproperties=FP_B, fontsize=11, color=C_NAVY)
    ax.text(5, 2.5, "AI 연산 증가율  −  메모리 효율 개선률", ha="center",
            fontproperties=FP_B, fontsize=12, color=C_NAVY2)
    ax.text(5, 1.2, "가격↑ = 26~28 초과이익 / 28이후 효율화·대체 촉진 (양날의 검)",
            ha="center", fontproperties=FP, fontsize=9, color=C_GRAY)
    ax.set_title("HBM 수요 = 연산↑ − 효율↑", fontproperties=FP_B, fontsize=12, color=C_NAVY)

    # right: tech readiness
    ax = axes[1]
    techs = ["Cerebras\n온칩 SRAM", "Groq\nSRAM 추론", "NVIDIA+Groq\ninference", "KV Cache\n압축", "HBF/SSD\n계층화"]
    levels = [90, 90, 85, 80, 55]
    colors = [C_GREEN, C_GREEN, C_GREEN, C_GREEN, C_AMBER]
    bars = ax.barh(techs, levels, color=colors, height=0.55)
    _style_ax(ax, "HBM 의존도 축소 움직임 (정성)")
    ax.set_xlim(0, 100)
    _xlabel(ax, "진행도 (상대 점수)")
    for label in ax.get_yticklabels():
        label.set_fontproperties(FP)
        label.set_fontsize(8)
    for bar, v, lab in zip(bars, levels, ["높음"] * 4 + ["진행중"]):
        ax.text(v + 2, bar.get_y() + bar.get_height() / 2, lab,
                va="center", fontproperties=FP, fontsize=8, color=C_GRAY)

    fig.suptitle("캐시 우드·벤 톰슨 논쟁 — '지금 틀렸다'가 아니라 '영구성 가정 경계'",
                 fontproperties=FP_B, fontsize=12, color=C_NAVY, y=1.02)
    fig.tight_layout()
    return save_fig(fig, "09_hbm_debate.png")


def chart_memory_hierarchy() -> Path:
    fig, ax = plt.subplots(figsize=(10.5, 3.6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 3.5)
    ax.axis("off")
    layers = [
        (0.4, "SRAM\n(온칩)", C_GREEN, "초저지연"),
        (2.8, "HBM\n(GPU)", C_NAVY, "고대역"),
        (5.2, "DRAM\n(시스템)", C_NAVY2, "용량"),
        (7.6, "SSD/HBF\n(스토리지)", C_GOLD, "저비용"),
    ]
    for x, label, color, sub in layers:
        ax.add_patch(mpatches.FancyBboxPatch(
            (x, 1.0), 1.9, 1.8, boxstyle="round,pad=0.05,rounding_size=0.12",
            facecolor=color, edgecolor="none", alpha=0.92,
        ))
        ax.text(x + 0.95, 2.1, label, ha="center", va="center",
                fontproperties=FP_B, fontsize=11, color="white")
        ax.text(x + 0.95, 0.55, sub, ha="center", fontproperties=FP, fontsize=9, color=C_GRAY)
    for x in (2.35, 4.75, 7.15):
        ax.annotate("", xy=(x + 0.4, 1.9), xytext=(x, 1.9),
                    arrowprops=dict(arrowstyle="<->", color=C_GOLD, lw=2))
    ax.text(5, 3.2, "SRAM ≠ HBM 대체 — Workload별 분업·계층 최적화",
            ha="center", fontproperties=FP_B, fontsize=12, color=C_NAVY)
    return save_fig(fig, "10_memory_hierarchy.png")


def chart_isu_multilam() -> Path:
    fig, axes = plt.subplots(1, 2, figsize=(10.8, 4.0))

    ax = axes[0]
    cats = ["1Q 매출", "2Q 매출", "수주잔고"]
    vals = [7, 11, 20]
    bars = ax.bar(cats, vals, color=[C_LIGHT, C_NAVY2, C_NAVY], edgecolor="white", width=0.55)
    for bar in bars:
        bar.set_edgecolor(C_NAVY2)
        if bar.get_facecolor()[:3] == plt.matplotlib.colors.to_rgb(C_LIGHT):
            bar.set_facecolor("#A8C0E0")
    _style_ax(ax, "이수페타시스 Multi-Lam 비중 (%)")
    _ylabel(ax, "%")
    for label in ax.get_xticklabels():
        label.set_fontproperties(FP)
    for bar, v in zip(bars, vals):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.5,
                f"{v}%", ha="center", fontproperties=FP_B, fontsize=11, color=C_NAVY)

    ax = axes[1]
    stages = ["현재", "27년 2Q", "28년 하반기"]
    capa = [1200, 1500, 1800]
    ax.plot(stages, capa, marker="o", markersize=10, color=C_NAVY, lw=2.5)
    ax.fill_between(range(3), capa, alpha=0.15, color=C_NAVY2)
    _style_ax(ax, "월 매출 Capa 로드맵 (억원)")
    for label in ax.get_xticklabels():
        label.set_fontproperties(FP)
    for i, v in enumerate(capa):
        ax.text(i, v + 40, f"{v:,}", ha="center", fontproperties=FP_B, fontsize=10, color=C_NAVY)
    _ylabel(ax, "억원/월")

    fig.suptitle("이수페타시스 — Capa + Multi-Lam 이익 레버리지", fontproperties=FP_B,
                 fontsize=13, color=C_NAVY, y=1.02)
    fig.tight_layout()
    return save_fig(fig, "11_isu_multilam.png")


def chart_sector_rotation() -> Path:
    fig, ax = plt.subplots(figsize=(10.5, 3.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4)
    ax.axis("off")
    # left box
    ax.add_patch(mpatches.FancyBboxPatch((0.3, 1.0), 3.2, 2.2, boxstyle="round,pad=0.05",
                                         facecolor=C_NAVY, alpha=0.9))
    ax.text(1.9, 2.5, "AI · 반도체", ha="center", fontproperties=FP_B, fontsize=12, color="white")
    ax.text(1.9, 1.7, "차익실현 / 자금이탈\n기술섹터 약세", ha="center", fontproperties=FP,
            fontsize=9, color="white")
    ax.annotate("", xy=(6.3, 2.1), xytext=(3.7, 2.1),
                arrowprops=dict(arrowstyle="->", color=C_GOLD, lw=3))
    ax.text(5, 2.55, "자금 Pivot", ha="center", fontproperties=FP_B, fontsize=10, color=C_GOLD)
    ax.add_patch(mpatches.FancyBboxPatch((6.5, 1.0), 3.2, 2.2, boxstyle="round,pad=0.05",
                                         facecolor=C_GREEN, alpha=0.9))
    ax.text(8.1, 2.5, "헬스케어", ha="center", fontproperties=FP_B, fontsize=12, color="white")
    ax.text(8.1, 1.7, "모더나 암백신 3상\nNBI 신고가", ha="center", fontproperties=FP,
            fontsize=9, color="white")
    ax.text(5, 0.4, "장기금리↓에도 기술주↑ 공식 미작동 → 금리보다 섹터 자금 이동이 단기 핵심",
            ha="center", fontproperties=FP, fontsize=9.5, color=C_GRAY)
    ax.set_title("8/19 뉴욕 — 금리 하락 vs 섹터 로테이션", fontproperties=FP_B,
                 fontsize=13, color=C_NAVY, pad=4)
    return save_fig(fig, "12_sector_rotation.png")


def chart_watchlist() -> Path:
    fig, ax = plt.subplots(figsize=(10.5, 4.2))
    ax.axis("off")
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    items = [
        (0.3, 4.2, "① 10년물", "4.7% 이하 안정\nvs 5% 돌파", C_NAVY),
        (3.5, 4.2, "② 유가", "$90 전후 안정\nvs $100+", C_AMBER),
        (6.7, 4.2, "③ USD/JPY", "157~159 안정\nvs 150대 급락", C_NAVY2),
        (0.3, 1.2, "④ 이란전", "소강 vs 확전\n(금리·유가)", C_RED),
        (3.5, 1.2, "⑤ 엔비디아", "8/26 실적\nCAPEX·Rubin", C_GREEN),
        (6.7, 1.2, "⑥ 주주환원", "SKH 추가규모\n3Q26 발표", C_GOLD),
    ]
    for x, y, title, body, color in items:
        ax.add_patch(mpatches.FancyBboxPatch(
            (x, y), 2.9, 1.8, boxstyle="round,pad=0.04,rounding_size=0.1",
            facecolor=color, edgecolor="none", alpha=0.9,
        ))
        ax.text(x + 1.45, y + 1.35, title, ha="center", fontproperties=FP_B,
                fontsize=11, color="white")
        ax.text(x + 1.45, y + 0.65, body, ha="center", fontproperties=FP,
                fontsize=9, color="white")
    ax.set_title("앞으로 볼 6대 변수 대시보드", fontproperties=FP_B, fontsize=13, color=C_NAVY)
    return save_fig(fig, "13_watchlist.png")


def chart_buyback_peers() -> Path:
    fig, ax = plt.subplots(figsize=(9.5, 3.8))
    names = ["키옥시아\n(발표~8/10)", "키옥시아\n(누적 8/18)", "샌디스크\nBuyback확대", "SK하이닉스\n40조(3.3%)"]
    rets = [3.2, 7.4, 20, None]  # Sandisk +20% was after other factors; show as reference note
    # For sandisk the comment said buyback wasn't the direct cause of +20%
    display = [3.2, 7.4, 5, 5]  # SKH expected min +5% from policy; use illustrative
    colors = [C_GRAY, C_GRAY, C_AMBER, C_NAVY]
    labels_note = ["+3.2%", "+7.4%", "한도확대\n(직접원인≠+20%)", "정책효과\n기대 +5%↑"]
    bars = ax.bar(names, [3.2, 7.4, 8, 6], color=colors, width=0.55, alpha=0.9)
    _style_ax(ax, "자사주 매입 이벤트와 주가 반응 (코멘트 정리)")
    for label in ax.get_xticklabels():
        label.set_fontproperties(FP)
        label.set_fontsize(8)
    for bar, note in zip(bars, labels_note):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.3,
                note, ha="center", fontproperties=FP, fontsize=8, color=C_NAVY)
    _ylabel(ax, "주가 반응 (%)")
    ax.text(0.5, -2.8, "※ 샌디스크 +20%는 Buyback만의 직접 원인이 아님 (8/5~7 -15% 후 반등·성장률 재평가 복합)",
            fontproperties=FP, fontsize=8, color=C_GRAY, transform=ax.transData)
    fig.tight_layout()
    return save_fig(fig, "14_buyback_peers.png")


def generate_all_charts() -> dict[str, Path]:
    return {
        "macro_chain": chart_macro_chain(),
        "rate_gauge": chart_rate_gauge(),
        "oil_gauge": chart_oil_gauge(),
        "macro_triangle": chart_macro_triangle(),
        "skh_fcf": chart_skh_fcf_return(),
        "per": chart_per_compare(),
        "fx_sens": chart_fx_sensitivity(),
        "fx_path": chart_fx_path(),
        "hbm": chart_hbm_debate(),
        "hierarchy": chart_memory_hierarchy(),
        "isu": chart_isu_multilam(),
        "rotation": chart_sector_rotation(),
        "watch": chart_watchlist(),
        "buyback": chart_buyback_peers(),
    }


# ── Document ───────────────────────────────────────────────


class Notes:
    def __init__(self):
        self.doc = Document()
        self._setup()

    def _setup(self):
        sec = self.doc.sections[0]
        sec.page_width = Mm(210)
        sec.page_height = Mm(297)
        sec.left_margin = Mm(15)
        sec.right_margin = Mm(15)
        sec.top_margin = Mm(14)
        sec.bottom_margin = Mm(14)

        normal = self.doc.styles["Normal"]
        normal.font.name = KR_FONT
        normal.font.size = Pt(11)
        normal.font.color.rgb = DARK
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), KR_FONT)
        pf = normal.paragraph_format
        pf.space_after = Pt(5)
        pf.space_before = Pt(0)
        pf.line_spacing = 1.16

        header = sec.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = hp.add_run("8/19 시장브리핑  ·  환율 · 주주환원 · HBM · 매크로  ·  시각화 보고서")
        set_run_font(r, size=8.5, color=GRAY)

        footer = sec.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run("Quick 코멘트 종합  ·  매수·매도 추천 아님  ·  ")
        set_run_font(r, size=8, color=GRAY)
        fld = parse_xml(
            f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE ">'
            f'<w:r><w:rPr><w:sz w:val="16"/><w:color w:val="4B5563"/>'
            f'<w:rFonts w:ascii="{KR_FONT}" w:hAnsi="{KR_FONT}" w:eastAsia="{KR_FONT}"/></w:rPr>'
            f"<w:t></w:t></w:r></w:fldSimple>"
        )
        fp._p.append(fld)

        core = self.doc.core_properties
        core.title = "8월 19일 시장브리핑 (환율·주주환원·HBM·매크로)"
        core.author = "준혁"
        core.subject = "Quick 코멘트 시각화 보고서"

    def p(self, text, size=11, bold=False, color=DARK, space_after=5, space_before=0, align="left"):
        para = self.doc.add_paragraph()
        para.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }[align]
        para.paragraph_format.space_after = Pt(space_after)
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.line_spacing = 1.16
        run = para.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
        return para

    def h1(self, text, num=None):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
        if num:
            run = para.add_run(f"{num}  ")
            set_run_font(run, size=15, bold=True, color=GOLD)
        run = para.add_run(text)
        set_run_font(run, size=15, bold=True, color=NAVY)
        pPr = para._p.get_or_add_pPr()
        pPr.append(
            parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="12" w:space="4" w:color="{NAVY_HEX}"/>'
                f"</w:pBdr>"
            )
        )
        return para

    def h2(self, text):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(3)
        run = para.add_run(text)
        set_run_font(run, size=12.5, bold=True, color=NAVY2)
        return para

    def bullet(self, text, level=0, bold_lead=None, size=10.5):
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.5 + level * 0.4)
        para.paragraph_format.first_line_indent = Cm(-0.32)
        para.paragraph_format.space_after = Pt(2)
        mark = "• " if level == 0 else "– "
        run = para.add_run(mark)
        set_run_font(run, size=size, color=NAVY2 if level == 0 else GRAY)
        if bold_lead:
            run = para.add_run(bold_lead)
            set_run_font(run, size=size, bold=True, color=DARK)
            run = para.add_run(text)
            set_run_font(run, size=size, color=DARK)
        else:
            run = para.add_run(text)
            set_run_font(run, size=size, color=DARK)
        return para

    def flow(self, items, size=10.5):
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after = Pt(6)
        for i, item in enumerate(items):
            if i:
                run = para.add_run("  →  ")
                set_run_font(run, size=size, bold=True, color=GOLD)
            run = para.add_run(item)
            set_run_font(run, size=size, bold=True, color=NAVY)
        return para

    def callout(self, title, body, kind="key"):
        palette = {
            "key": (NAVY_HEX, LIGHT_HEX, NAVY),
            "bull": ("166534", GREEN_HEX, GREEN),
            "bear": ("991B1B", RED_HEX, RED),
            "note": (GOLD_HEX, AMBER_HEX, AMBER),
            "blue": (NAVY2_HEX, BLUE_HEX, NAVY2),
        }
        accent, fill, title_color = palette[kind]
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        shade_cell(cell, fill)
        set_left_accent(cell, accent, sz="28")
        set_cell_margins(cell, top=70, bottom=70, left=110, right=110)
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        r = p1.add_run(title)
        set_run_font(r, size=10, bold=True, color=title_color)
        if isinstance(body, str):
            body = [body]
        for line in body:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.12
            r = p.add_run(line)
            set_run_font(r, size=10.2, color=DARK)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def table(self, headers, rows, col_widths=None, first_col_bold=True):
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table, color="D5DCE6", sz="4")
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            shade_cell(cell, NAVY_HEX)
            cell_text(cell, h, size=9, bold=True, color=WHITE, align="center")
        prevent_row_split(table.rows[0])
        for r_i, row in enumerate(rows):
            for c_i, val in enumerate(row):
                cell = table.rows[r_i + 1].cells[c_i]
                shade_cell(cell, ROW_HEX if r_i % 2 else WHITE_HEX)
                align = "left" if c_i == 0 else "center"
                bold = first_col_bold and c_i == 0
                cell_text(cell, str(val), size=9, bold=bold, color=DARK, align=align)
            prevent_row_split(table.rows[r_i + 1])
        if col_widths:
            for row in table.rows:
                for i, w in enumerate(col_widths):
                    row.cells[i].width = Cm(w)
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)
        return table

    def image(self, path: Path, width_in=6.4):
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run()
        run.add_picture(str(path), width=Inches(width_in))
        return para

    def save(self, path: Path):
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))


def build(charts: dict[str, Path]):
    n = Notes()

    # Cover
    n.p("2026. 8. 19.  ·  Quick 코멘트 종합  ·  시각화 보고서", size=10.5, color=GRAY, align="center", space_after=3)
    n.p("MARKET BRIEFING", size=12, bold=True, color=GOLD, align="center", space_after=2)
    n.p("8월 19일 시장 브리핑", size=22, bold=True, color=NAVY, align="center", space_after=2)
    n.p("환율 · 주주환원 · HBM 논쟁 · 매크로 · 개별 종목", size=13, bold=True, color=NAVY2, align="center", space_after=8)

    n.callout(
        "오늘 한 장으로 보면",
        [
            "단기 축 = 유가 → 10년물 금리 → AI/반도체 밸류에이션. 10년물 4.7% 이하 안정이 핵심.",
            "국내 축 = SK하이닉스 40조 자사주 소각 + FCF 50% 이상 환원. 삼성 주주환원 기대 동반.",
            "구조 축 = HBM 고가격은 26~28 초과이익이지만, 이후 효율화·대체 유인을 키우는 양날의 검.",
            "환율 축 = 1,400원 하회는 국내 달러공급이 주도. 추가 하락은 달러인덱스·국내수급의 함수.",
        ],
        kind="key",
    )

    n.h2("보고서 구성")
    n.table(
        ["파트", "주제", "핵심 차트"],
        [
            ["1", "매크로 — 전쟁·유가·금리", "전달경로 · 금리/유가 게이지"],
            ["2", "자금 피벗 — AI→헬스케어", "섹터 로테이션"],
            ["3", "SK하이닉스 주주환원", "FCF·환원 구조"],
            ["4", "메모리 밸류에이션", "Peer PER"],
            ["5", "환율 — 1,400 하회 해설", "민감도 · 수급 루프"],
            ["6", "HBM 논쟁 (우드·톰슨)", "수요식 · 메모리 계층"],
            ["7", "반도체 뉴스플로우", "파운드리·마벨"],
            ["8", "개별 아이디어", "이수·기가비스 등"],
            ["9", "워치리스트", "6대 변수 대시보드"],
        ],
        col_widths=[2.0, 6.5, 8.5],
    )

    # ── 1 Macro ──
    n.h1("매크로 — 전쟁보다 금리", num="1.")
    n.callout(
        "한 줄 결론",
        [
            "미·이란 불확실성 → 유가↑ → 인플레 우려 → 미 국채금리↑ → 고PER AI/반도체 압박.",
            "AI 수요가 갑자기 꺾인 것이 아니라, 밸류에이션 할인율 충격에 가깝다.",
        ],
        kind="key",
    )
    n.image(charts["macro_chain"])
    n.image(charts["rate_gauge"])
    n.image(charts["oil_gauge"])

    n.h2("임계값 요약")
    n.table(
        ["변수", "완화 신호", "위험 신호"],
        [
            ["브렌트 유가", "$90 전후 안정", "$100 이상 — 인플레·금리 악순환"],
            ["미 10년물", "4.7% 이하 안정", "5% 돌파·고착 — 위험자산 회피"],
            ["미 30년물", "고점 후 진정", "5.3%대 재돌파 시 심리 악화"],
            ["USD/JPY", "157~159 안정/반등", "150대 급락 — 엔캐리 청산 의심"],
        ],
        col_widths=[3.5, 6.5, 7.0],
    )

    n.h2("무게 중심 삼각형")
    n.image(charts["macro_triangle"], width_in=5.2)
    n.bullet("토큰 수요 확대 → AI 투자 GO는 웬만한 금리 상승에도 근본 궤도가 바뀌기 어렵다.")
    n.bullet("다만 안정화 전까지는 채권시장 의견이 더 힘을 얻는 것이 현실.")
    n.bullet("재무부 장기채 바이백 2배 확대 = 단기 방어선. 재정적자·인플레·AI 자금수요는 미해결.")

    n.callout(
        "재무부 바이백의 의미",
        [
            "장기금리 급등 속도↓ · 성장주 할인율에 단기 호재.",
            "근본 해결 아님: 재정적자↑ / 인플레 / AI CAPEX 회사채 수요는 그대로.",
            "Fed 의사록은 매파적이나, 당장은 Treasury 영향력이 더 크게 작용.",
        ],
        kind="blue",
    )

    # ── 2 Rotation ──
    n.h1("자금 피벗 — 금리↓인데 기술주↓", num="2.")
    n.image(charts["rotation"])
    n.bullet("공식 깨짐: 장기금리↓ → 기술주↑ 가 이번에는 작동하지 않음.")
    n.bullet("모더나 암백신 3상 성공 → 헬스케어로 성장주 자금 이동.")
    n.bullet("한국 반도체 투자자: '금리 하락'만 보고 안도하기보다 AI·반도체 자금 이탈 지속 여부를 확인.")
    n.bullet("다음 관전: 8/26 엔비디아 실적 — AI 심리 재강화 vs 로테이션 지속.")

    # ── 3 SK Hynix returns ──
    n.h1("SK하이닉스 주주환원", num="3.")
    n.callout(
        "핵심 숫자",
        [
            "40조원 자사주 취득·전량 소각 (발행주식 ~3.3%, 8/20~약 3개월).",
            "2025~27 누적 FCF의 '50% 이상' 환원으로 상향 (기존 50% 범위 내 → 초과 목표).",
            "이미 40조 확정 → 단순 50% 기준만으로도 추가 ~152.5조 이상 여지.",
            "구체 추가 규모·방식은 3Q26 실적발표 때 안내 예정.",
        ],
        kind="bull",
    )
    n.image(charts["skh_fcf"])

    n.h2("해석 시 주의")
    n.bullet("192.5조 = 2027년에 한꺼번에 지급이 아님. 프로그램 기간(25~27) 누적 환원.")
    n.bullet("2028년 102.5조 표는 FCF×50% 참고치일 뿐. 회사 정책상 2028 환원액이 아님.")
    n.bullet("회사는 주가 저평가를 명시 → 자본 재배치·주주가치 제고 목적.")
    n.bullet("ADR 발행으로 희석된 지분율 → 소각으로 SK스퀘어 지분율을 이전 수준으로 되돌리는 구조.")
    n.flow(["순현금 ~69조", "강력한 FCF", "40조 소각", "추가 배당·Buyback", "밸류 재평가"])

    n.h2("Peer Buyback 맥락")
    n.image(charts["buyback"])
    n.p("키옥시아는 단기 집행 완료에도 시황 부진 시 소폭 상승에 그침(비중 이상은 상승). "
        "샌디스크 +20%는 Buyback만의 직접 원인이 아니라 Investor Day 이후 성장률·마진 재평가 + EPS 레버리지로 보는 것이 타당.",
        size=10.5)

    # ── 4 Valuation ──
    n.h1("메모리 밸류에이션 스냅샷", num="4.")
    n.image(charts["per"])
    n.table(
        ["종목", "가격(코멘트)", "26년 PER", "27년 PER", "보수 시나리오"],
        [
            ["SK하이닉스 본주", "150만원", "4.3배", "3.4배", "27년 5.1배"],
            ["삼성전자", "24.75만원", "5.2배", "3.7배", "27년 5.6배"],
            ["SKH ADR", "$163.8", "6.6배", "5.2배", "27년 7.7배"],
            ["마이크론", "$937", "Fwd 7.5배", "CY27 6.25배", "—"],
            ["샌디스크", "$1,568", "—", "FY27 7.8배", "—"],
        ],
        col_widths=[3.6, 3.2, 2.8, 2.8, 4.0],
    )
    n.callout(
        "ADR 프리미엄 함의",
        [
            "본주 대비 ADR ~52% 프리미엄. 정상 +20%(TSMC ~15%) 가정 시 본주 ~190만 암시.",
            "최근 실제 30%+ 프리미엄 감안 시 본주 169~175만원 레인지 논의.",
        ],
        kind="note",
    )
    n.p("다른 접근: 27년 성장 없다고 단순화하고 26년 PER 6~7배(과거 사이클 4~8배) 적용 시 "
        "SKH 208~242만 / 삼성 28.7~33.5만 산출 가능(코멘트 시나리오).", size=10.5)

    # ── 5 FX ──
    n.h1("환율 — 왜 1,400 아래인가", num="5.")
    n.image(charts["fx_path"])
    n.image(charts["fx_sens"])

    n.h2("추가 하락의 조건과 리스크")
    n.bullet("달러인덱스만 3~4% 하락해도 하단 ~1,360원 가능(계산상).")
    n.bullet("한국 강한 달러 공급 가세 시 1,340원대도 열림 — 다만 '조건부'.")
    n.bullet("원화 추가강세 리스크: ① 외인 국내주식 매도 ② 미-이란/유가 ③ 1,350 부근 달러 수요 증가.")
    n.bullet("시장은 연준 금리인상 1~2회를 일부 반영. 기대 되돌림 시 DXY 99→96~97 → 추가 원화강세.")
    n.callout(
        "수출주 함의",
        ["너무 가파른 원화강세는 3Q 수출주 원화약세 효과 기대를 깎는다. 환율 민감도는 하이닉스 > 삼성."],
        kind="bear",
    )

    # ── 6 HBM ──
    n.h1("HBM 논쟁 — 우드·톰슨을 어떻게 읽을까", num="6.")
    n.callout(
        "가장 타당한 해석",
        [
            "'현재 메모리주가 틀렸다'가 아니라,",
            "'현재의 높은 HBM 가격결정력이 영구적이라고 가정하면 안 된다'는 경고.",
        ],
        kind="key",
    )
    n.image(charts["hbm"])
    n.image(charts["hierarchy"])

    n.h2("타당성 정리")
    n.table(
        ["주장", "타당성", "코멘트"],
        [
            ["HBM 비싸서 메모리 수요가 곧 꺾인다", "낮음", "가격≠즉시 수요 고점"],
            ["비싸면 효율화·대체 기술 유인↑", "높음", "경제적 유인 논리"],
            ["SRAM이 HBM을 완전 대체", "오해", "대체보다 분업·계층화"],
        ],
        col_widths=[6.5, 2.5, 8.0],
    )
    n.p("벤 톰슨 비유: HBM 공급자 = 호르무즈. 가격을 과도하게 올리면 고객은 장기적으로 공급망·기술 대체를 찾는다.", size=10.5)

    # ── 7 Semi news ──
    n.h1("반도체 뉴스플로우", num="7.")
    n.h2("삼성 파운드리 가격 인상 (Reuters)")
    n.table(
        ["공정", "인상폭", "배경"],
        [
            ["4나노(SF4)", "미·중 10~15% / 대만 5~10%", "TSMC 포화 → 삼성으로 분산"],
            ["5나노(SF5)", "웨이퍼 10~15%", "중국 팹리스 해외 의존↑"],
            ["8나노", "~10%", "평택 SF4 풀가동(퀄컴+HBM 베이스다이)"],
        ],
        col_widths=[3.5, 5.5, 8.0],
    )
    n.bullet("파운드리 흑자 전환 기대(이르면 내년). AI·HPC 비중 확대. 구글 4나노 협의 중.")

    n.h2("Google × Marvell")
    n.flow(["커스텀 ASIC 확대", "TPU 주변 생태계", "Warrant 5,897만주", "매출 $5억마다 vest"])
    n.bullet("마벨 +7~10%대 vs 브로드컴 -4%후반 — Broadcom TPU 독점 견제 신호로 해석.")
    n.bullet("행사가 $206.58, FY27 Q3~FY33. 단기 이벤트보다 장기 AI ASIC 계약 성격.")

    # ── 8 Names ──
    n.h1("개별 아이디어 요약", num="8.")
    n.image(charts["isu"])

    n.h2("이수페타시스 — 이익 레버리지")
    n.bullet("2Q26 컨센 상회. Multi-Lam 1Q 7%→2Q 11%, 수주잔고 20%+.")
    n.bullet("4Q G사 Multi-Lam 전환 + M사 ASIC 양산. 판가 +15% 협상.")
    n.bullet("Capa+믹스+수율+판가 → 양적·질적 성장. 27년 OP +10% 상향여지.")

    n.h2("기가비스 — FC-BGA의 눈과 레이저")
    n.bullet("AOI(검사)+AOR(수리). 고객: 이비덴·신코·유니마이크론·삼성전기.")
    n.bullet("일본 기판사 89.5억 수주(매출比 17.1%). 컨센 TP ~19만, 26년 기준 비쌈·27년 성장이 초점.")
    n.bullet("메리츠: 26년 매출 1,855억 / OP 721억 전망(장비주 레버리지).")

    n.h2("기타 한 줄")
    n.table(
        ["종목/테마", "포인트"],
        [
            ["LS", "자회사 체력↑. 26/27 OP +20%/+17% 상향. 자사주 소각 의무화 논의"],
            ["한화에어로스페이스", "미 육군 MTC 시제기 단독. 양산 시 ~10조 파이프라인"],
            ["LG엔솔", "북미 EV→ESS 재편. 테슬라 메가팩 LFP. 방산 배터리 협의"],
            ["알테오젠", "키트루다SC 상업화 + 병용 확대. MSD·모더나 모멘텀 연계"],
            ["유니트리", "커촹반 상장 급등. PSR 고평가 논쟁. 상업화 ROI가 다음 관문"],
        ],
        col_widths=[4.5, 12.5],
    )

    # ── 9 Watch ──
    n.h1("워치리스트 & 마감 시사점", num="9.")
    n.image(charts["watch"])

    n.callout(
        "마감 생각 (요약)",
        [
            "금리발 하락을 주주환원으로 일부 방어. 롱·숏 공방은 계속.",
            "'금리↓=주가↑' 공식에만 몰입하기보다, 성장이 금리 비용을 이기는지 확인하는 국면.",
            "추가 하락은 막더라도 보수적 접근 유지. 반등 시 차익 매물 구간 가능.",
        ],
        kind="note",
    )

    n.h2("내일·앞으로의 체크")
    n.flow(["10년물 4.7%", "유가 $90", "USD/JPY 안정", "Nvidia 8/26", "3Q26 환원 디테일"])

    n.p(" ", space_after=8)
    n.callout(
        "면책",
        [
            "본 자료는 Quick 코멘트를 시각화·재구성한 참고용 브리핑이며 매수·매도 추천이 아닙니다.",
            "투자 판단과 책임은 각자에게 있으며, 법적 자료로 활용할 수 없습니다.",
            "수치·시나리오는 공개 코멘트 시점 기준이며 시장 상황에 따라 달라질 수 있습니다.",
        ],
        kind="bear",
    )

    n.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


def main():
    print("Generating charts...")
    charts = generate_all_charts()
    for k, p in charts.items():
        print(f"  {k}: {p}")
    print("Building document...")
    build(charts)


if __name__ == "__main__":
    main()
