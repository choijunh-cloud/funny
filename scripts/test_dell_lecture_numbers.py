#!/usr/bin/env python3
"""Dell Q2 FY27 강의노트에 들어간 핵심 숫자가 보도자료·콜과 맞는지 검증."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

from generate_sep01_dell_lecture import OUT_PATH, build

# 공식 보도자료 (단위: 백만 달러, EPS는 달러)
PR = {
    "q2_rev": 46971,
    "q2_ai": 16401,
    "q2_trad": 10531,
    "q2_storage": 4850,
    "q2_isg": 31782,
    "q2_csg": 15034,
    "q2_commercial": 13192,
    "q2_consumer": 1842,
    "q2_isg_oi": 4781,
    "q2_csg_oi": 1142,
    "h1_rev": 90813,
    "h1_ai": 32533,
    "h1_trad": 19074,
    "h1_storage": 9184,
    "h1_isg": 60791,
    "h1_csg": 29643,
    "h1_isg_oi": 7836,
    "h1_csg_oi": 2312,
    "q2_gaap_eps": 6.34,
    "q2_nongaap_eps": 7.04,
    "h1_nongaap_eps": 11.90,
    "q2_nongaap_oi": 5929,
    "h1_nongaap_oi": 10164,
    "q2_cfo": 2225,
    "h1_cfo": 6306,
    "q2_adj_fcf": 8149,
    "fy_rev_old": 167000,
    "fy_rev_new": 192000,
    "fy_ai_old": 60000,
    "fy_ai_new": 74000,
    "fy_eps_old": 17.90,
    "fy_eps_new": 25.50,
    "q3_rev": 49000,
    "q3_eps": 6.50,
    "q3_ai": 19000,
    "q1_backlog": 51300,  # 전 분기 콜/보도
    "q2_orders": 60900,
    "q2_backlog": 95000,
}


def derived():
    q1_rev = PR["h1_rev"] - PR["q2_rev"]
    q1_ai = PR["h1_ai"] - PR["q2_ai"]
    q1_trad = PR["h1_trad"] - PR["q2_trad"]
    q1_storage = PR["h1_storage"] - PR["q2_storage"]
    q1_isg = PR["h1_isg"] - PR["q2_isg"]
    q1_isg_oi = PR["h1_isg_oi"] - PR["q2_isg_oi"]
    q4_rev = PR["fy_rev_new"] - PR["h1_rev"] - PR["q3_rev"]
    q4_ai = PR["fy_ai_new"] - PR["h1_ai"] - PR["q3_ai"]
    q4_eps = round(PR["fy_eps_new"] - PR["h1_nongaap_eps"] - PR["q3_eps"], 2)
    h2_rev = PR["fy_rev_new"] - PR["h1_rev"]
    backlog_check = PR["q1_backlog"] + PR["q2_orders"] - PR["q2_ai"]
    return {
        "q1_rev": q1_rev,  # 43842
        "q1_ai": q1_ai,  # 16132
        "q1_trad": q1_trad,  # 8543
        "q1_storage": q1_storage,  # 4334
        "q1_isg": q1_isg,  # 29009
        "q1_isg_oi_rate": q1_isg_oi / q1_isg,  # ~10.5%
        "q2_isg_oi_rate": PR["q2_isg_oi"] / PR["q2_isg"],  # 15.0%
        "q4_rev": q4_rev,  # 22187? wait 192000-90813-49000=52187
        "q4_ai": q4_ai,  # 22467
        "q4_eps": q4_eps,  # 7.10
        "h2_rev": h2_rev,  # 101187
        "non_ai_raise": (PR["fy_rev_new"] - PR["fy_rev_old"]) - (PR["fy_ai_new"] - PR["fy_ai_old"]),
        "backlog_check": backlog_check,  # 95799
        "q2_ai_qoq": PR["q2_ai"] / q1_ai - 1,
    }


def doc_text(path: Path) -> str:
    d = Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def must_contain(blob: str, needles: list[str]) -> list[str]:
    missing = []
    for n in needles:
        if n not in blob:
            missing.append(n)
    return missing


def main() -> int:
    build()
    assert OUT_PATH.exists(), f"missing {OUT_PATH}"
    blob = doc_text(OUT_PATH)
    d = derived()

    # 산식 가드 — 강의에 넣은 역산이 보도자료와 맞는지
    assert abs(d["q1_rev"] - 43842) < 2
    assert abs(d["q1_ai"] - 16132) < 2
    assert abs(d["q4_rev"] - 52187) < 2
    assert abs(d["q4_ai"] - 22467) < 2
    assert d["q4_eps"] == 7.10
    assert abs(d["non_ai_raise"] - 11000) < 1
    assert 95000 <= d["backlog_check"] <= 96000
    assert d["q2_ai_qoq"] < 0.03  # AI 매출 QoQ 거의 평탄
    assert d["q2_isg_oi_rate"] > 0.149
    assert 0.104 < d["q1_isg_oi_rate"] < 0.106

    missing = must_contain(
        blob,
        [
            "$470억",
            "$7.04",
            "$609억",
            "$164억",
            "$950억",
            "$105억",
            "+122%",
            "$49억",
            "+26%",
            "$1,920억",
            "$25.50",
            "$6.50",
            "DRAM",
            "NAND",
            "Vera Rubin",
            "14G",
            "18G",
            "Project Lightning",
            "120만",
            "6,500",
        ],
    )
    if missing:
        print("MISSING STRINGS:", missing)
        return 1

    # 흔한 단위 실수: $47억이라고 쓰면 $4.7B가 됨
    if re.search(r"\$47억\b", blob):
        print("BAD UNIT: $47억 (should be $470억 for $47B)")
        return 1
    if re.search(r"\$60\.9억", blob):
        print("BAD UNIT: $60.9억")
        return 1

    print("OK")
    print(f"docx={OUT_PATH} bytes={OUT_PATH.stat().st_size}")
    print(
        "derived",
        {
            "q1_rev_b": round(d["q1_rev"] / 1000, 1),
            "q1_ai_b": round(d["q1_ai"] / 1000, 1),
            "q4_rev_b": round(d["q4_rev"] / 1000, 1),
            "q4_ai_b": round(d["q4_ai"] / 1000, 1),
            "q4_eps": d["q4_eps"],
            "non_ai_raise_b": d["non_ai_raise"] / 1000,
            "backlog_check_b": round(d["backlog_check"] / 1000, 1),
            "q2_ai_qoq": round(d["q2_ai_qoq"] * 100, 1),
        },
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
