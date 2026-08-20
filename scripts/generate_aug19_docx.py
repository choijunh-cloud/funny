#!/usr/bin/env python3
"""8월 19일 시장상황 시각화 보고서(.docx) — 차트 삽입."""

from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Mm, Pt, RGBColor

import aug19_data as d
from charts_aug19 import render_all

OUT_PATH = Path("/workspace/lectures/8월 19일 시장상황 시각화 보고서.docx")
CHART_DIR = Path("/workspace/reports/charts")

KR_FONT = "맑은 고딕"
NAVY = RGBColor(0x0F, 0x20, 0x43)
NAVY2 = RGBColor(0x1E, 0x40, 0x7C)
GOLD = RGBColor(0xB8, 0x94, 0x3A)
GRAY = RGBColor(0x4B, 0x55, 0x63)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x16, 0x65, 0x34)
RED = RGBColor(0x99, 0x1B, 0x1B)
AMBER = RGBColor(0x7A, 0x5C, 0x12)

NAVY_HEX = "0F2043"
NAVY2_HEX = "1E407C"
GOLD_HEX = "B8943A"
LIGHT_HEX = "EEF2F8"
GREEN_HEX = "E8F5E9"
RED_HEX = "FDECEA"
AMBER_HEX = "FFF8E7"
BLUE_HEX = "E8F1FB"
ROW_HEX = "F7F9FC"
WHITE_HEX = "FFFFFF"


def set_run_font(run, size=11, bold=False, color=DARK, font=KR_FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), KR_FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def shade_cell(cell, fill):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill}"/>')
    )


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'<w:top w:w="{top}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/>'
            f'<w:bottom w:w="{bottom}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/>'
            f"</w:tcMar>"
        )
    )


def set_table_borders(table, color="D0D7E2", sz="4"):
    tbl_pr = table._tbl.tblPr if table._tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tbl_pr.append(
        parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f"</w:tblBorders>"
        )
    )


def set_left_accent(cell, color=NAVY_HEX, sz="28"):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="nil"/><w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="nil"/><w:right w:val="nil"/>'
            f"</w:tcBorders>"
        )
    )


def prevent_row_split(row):
    row._tr.get_or_add_trPr().append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))


def cell_text(cell, text, size=10, bold=False, color=DARK, align="left"):
    cell.text = ""
    align_enum = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    for i, line in enumerate(str(text).split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = align_enum
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1 if i < len(str(text).split("\n")) - 1 else 0)
        run = p.add_run(line)
        set_run_font(run, size=size, bold=bold, color=color)
    set_cell_margins(cell)


class Notes:
    def __init__(self, charts: dict[str, Path]):
        self.doc = Document()
        self.charts = charts
        self._setup()

    def _setup(self):
        sec = self.doc.sections[0]
        sec.page_width = Mm(210)
        sec.page_height = Mm(297)
        sec.left_margin = Mm(15)
        sec.right_margin = Mm(15)
        sec.top_margin = Mm(15)
        sec.bottom_margin = Mm(15)

        normal = self.doc.styles["Normal"]
        normal.font.name = KR_FONT
        normal.font.size = Pt(11)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), KR_FONT)

        hp = sec.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = hp.add_run("8/19 시장상황 시각화 보고서  ·  매크로 · 환율 · 환원 · HBM · 마벨")
        set_run_font(r, size=8.5, color=GRAY)

        fp = sec.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run("공개 퀵코멘트 재구성  ·  매수·매도 추천 아님  ·  ")
        set_run_font(r, size=8, color=GRAY)
        fp._p.append(
            parse_xml(
                f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE ">'
                f'<w:r><w:rPr><w:sz w:val="16"/><w:color w:val="4B5563"/>'
                f'<w:rFonts w:ascii="{KR_FONT}" w:hAnsi="{KR_FONT}" w:eastAsia="{KR_FONT}"/></w:rPr>'
                f"<w:t></w:t></w:r></w:fldSimple>"
            )
        )
        core = self.doc.core_properties
        core.title = "8월 19일 시장상황 시각화 보고서"
        core.author = "준혁"
        core.subject = "매크로, 환율, SK하이닉스 환원, HBM, 마벨, 파운드리"

    def p(self, text, size=11, bold=False, color=DARK, space_after=6, align="left"):
        para = self.doc.add_paragraph()
        para.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }[align]
        para.paragraph_format.space_after = Pt(space_after)
        run = para.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
        return para

    def h1(self, text, num=None):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
        if num:
            run = para.add_run(f"{num}  ")
            set_run_font(run, size=15, bold=True, color=GOLD)
        run = para.add_run(text)
        set_run_font(run, size=15, bold=True, color=NAVY)
        para._p.get_or_add_pPr().append(
            parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="12" w:space="3" w:color="{NAVY_HEX}"/>'
                f"</w:pBdr>"
            )
        )

    def h2(self, text):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(3)
        run = para.add_run(text)
        set_run_font(run, size=12.5, bold=True, color=NAVY2)

    def bullet(self, text, size=10.5):
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.5)
        para.paragraph_format.first_line_indent = Cm(-0.32)
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run("• ")
        set_run_font(run, size=size, color=NAVY2)
        run = para.add_run(text)
        set_run_font(run, size=size, color=DARK)

    def flow(self, items):
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(8)
        for i, item in enumerate(items):
            if i:
                run = para.add_run("  →  ")
                set_run_font(run, size=10, bold=True, color=GOLD)
            run = para.add_run(item)
            set_run_font(run, size=10, bold=True, color=NAVY)

    def callout(self, title, body, kind="key"):
        palette = {
            "key": (NAVY_HEX, LIGHT_HEX, NAVY),
            "bull": ("166534", GREEN_HEX, GREEN),
            "bear": ("991B1B", RED_HEX, RED),
            "note": (GOLD_HEX, AMBER_HEX, AMBER),
            "blue": (NAVY2_HEX, BLUE_HEX, NAVY2),
        }
        accent, fill, title_color = palette[kind]
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        shade_cell(cell, fill)
        set_left_accent(cell, accent)
        set_cell_margins(cell, top=70, bottom=70, left=110, right=110)
        cell.text = ""
        p1 = cell.paragraphs[0]
        r = p1.add_run(title)
        set_run_font(r, size=10, bold=True, color=title_color)
        if isinstance(body, str):
            body = [body]
        for line in body:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(line)
            set_run_font(r, size=10.5, color=DARK)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def table(self, headers, rows, col_widths=None):
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table)
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            shade_cell(cell, NAVY_HEX)
            cell_text(cell, h, size=9, bold=True, color=WHITE, align="center")
        prevent_row_split(table.rows[0])
        for r_i, row in enumerate(rows):
            for c_i, val in enumerate(row):
                cell = table.rows[r_i + 1].cells[c_i]
                shade_cell(cell, ROW_HEX if r_i % 2 else WHITE_HEX)
                cell_text(cell, str(val), size=9, bold=c_i == 0, align="left" if c_i == 0 else "center")
            prevent_row_split(table.rows[r_i + 1])
        if col_widths:
            for row in table.rows:
                for i, w in enumerate(col_widths):
                    row.cells[i].width = Cm(w)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def image(self, key, width=16.4):
        path = self.charts[key]
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        p.add_run().add_picture(str(path), width=Cm(width))

    def save(self, path: Path):
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))


def build(chart_dir: Path = CHART_DIR) -> Path:
    charts = render_all(chart_dir)
    n = Notes(charts)

    n.p("2026. 8. 19.  ·  당일 퀵코멘트 재구성  ·  시각화 보고서", size=10.5, color=GRAY, align="center", space_after=4)
    n.p("매크로의 역습  ·  원화 강세  ·  메모리의 양날", size=20, bold=True, color=NAVY, align="center", space_after=4)
    n.p("금리·환율이 할인해 놓은 자리 위에 하이닉스 환원, 삼성 파운드리, 마벨–구글, HBM 효율화 논쟁이 동시에 올라왔습니다.", size=11, color=GRAY, align="center", space_after=10)

    n.callout(
        "오늘 한 장",
        [
            f"하락의 직접 원인은 AI 수요 붕괴가 아니라 할인율 + 차익실현입니다. 저녁 핵심은 달러-원이 국내 수급으로 1,400 아래를 먼저 깼다는 점.",
            f"1,520→1,420 가정 시 SK하이닉스 EPS 약 −{d.SKH_EPS_FX_HIT:.1f}%, 27년 순익 300~400조이면 약 {d.SKH_NI_ADJ_LOW:.0f}~{d.SKH_NI_ADJ_HIGH:.0f}조 조정. 2H26 환율 조정 {d.SKH_2H26_FX_ADJ}조 언급.",
            f"40조 소각은 3.3%(EPS +{d.SKH_EPS_UPLIFT:.1f}%)이지 환원의 끝이 아닙니다. 25~27 FCF {d.SKH_FCF_2527}조 × 50% = {d.SKH_RETURN_MIN:g}조, 추가 {d.SKH_RETURN_ADD:g}조. 2028년 102.5조는 정책이 아닙니다.",
        ],
        kind="key",
    )

    n.h2("보고서 구성")
    n.table(
        ["파트", "질문", "가져갈 숫자"],
        [
            ["1 매크로", "전쟁인가, 금리인가?", "10년 4.7% vs 5.0%  ·  바이백 $20→40억"],
            ["2 환율", "왜 1,400 아래인가?", f"β 삼성 0.4 / SKH 0.9  ·  하단 {d.KRW_FLOOR_DXY:,}~{d.KRW_FLOOR_SUPPLY:,}"],
            ["3 환원", "40조면 끝인가?", f"{d.SKH_RETURN_MIN:g}조 하한  ·  3Q26 추가 안내"],
            ["4 밸류", "본주와 ADR 중 무엇을 믿나?", f"본주 {d.SKH_PER_26}/{d.SKH_PER_27}배  ·  ADR 프리미엄 52%"],
            ["5 HBM", "가격이 수요를 죽이나?", "연산증가 − 효율개선  ·  26~28 vs 28+"],
            ["6 마벨", "브로드컴 독점의 균열?", f"워런트 {d.MRVL_WARRANT_M}만주  ·  ${d.MRVL_STRIKE}"],
            ["7 파운드리·NVDA", "비메모리 가격결정력?", "파운드리 +10~15%  ·  Rubin 3Q26"],
            ["8 소부장", "익스포저를 어디에?", "이수 Multi-Lam  ·  기가비스 수주"],
        ],
        col_widths=[3.4, 5.4, 8.8],
    )

    # 1
    n.h1("매크로 — 전쟁 자체보다 금리", num="1.")
    n.flow(["이란·호르무즈", "유가", "인플레", "미 국채금리", "고PER 할인", "AI·반도체"])
    n.image("macro_levels")
    n.table(
        ["변수", "안정", "위험", "당일 인쇄"],
        [
            ["미 10년물", "4.7% 이하", "5.0% 돌파·고착", "4.75 → 4.708 → 바이백 후 4.64"],
            ["미 30년물", "고점 이탈", "5.3%대 고착", "5.34 → 5.285 → 5.19"],
            ["브렌트", "$90 전후", "$100+ 악순환", "WTI $84대 (장초 코멘트)"],
            ["USD/JPY", "157~159 유지", "159→155→150 급락", "반등 = 2024.8형 아님"],
        ],
        col_widths=[3.2, 3.8, 4.4, 6.2],
    )
    n.image("carry_compare")
    n.callout(
        "재무부 바이백 = 방어선, 해결책 아님",
        [
            "장기채 매입 $20억→$40억. 당장은 Treasury가 Fed 매파(7월 의사록 3명 인상)보다 금리를 눌렀습니다.",
            "재정적자·인플레·AI 회사채 수요는 그대로입니다. 금리 하락만 보고 기술주를 사는 공식은 이번 세션에서 깨졌습니다(헬스케어 로테이션).",
        ],
        kind="note",
    )
    n.bullet("2024.8/5: Nikkei −19.5%, 코스피 −11.9%, 엔화 6%대 급등 후 8/6 즉시 반등 = 레버리지 청산.")
    n.bullet("현재 엔화가 반등 중이면 1차 충격은 엔캐리보다 글로벌 금리/밸류 압박.")
    n.bullet("FT 이란 기사: 유럽 공격 결정이 아니라, 확전 시 남동부 유럽 미군 시설·해저 인프라 옵션 검토.")

    # 2
    n.h1("환율 — 국내 수급이 먼저", num="2.")
    n.flow(["법인세·설비 원화수요", "수출 달러매도", "헤지 비중 ↑", "달러-원 하락", "고환율 잔여 매도", "하락 가속"])
    n.image("fx_ladder")
    n.image("fx_sensitivity")
    n.image("fx_ni_adj")
    n.callout(
        "민감도 산식 (원문 교정 포함)",
        [
            f"원/달러 {d.USD_KRW_FROM:,}→{d.USD_KRW_TO:,} = {d.USD_KRW_MOVE_PCT:.2f}%.",
            f"삼성전자 β {d.SEC_FX_BETA} → EPS 약 {abs(d.USD_KRW_MOVE_PCT)*d.SEC_FX_BETA:.1f}%.  SK하이닉스 β {d.SKH_FX_BETA} → EPS 약 −{d.SKH_EPS_FX_HIT:.1f}%.",
            f"27년 순익 300~400조 × {d.SKH_EPS_FX_HIT:.1f}% ≈ {d.SKH_NI_ADJ_LOW:.1f}~{d.SKH_NI_ADJ_HIGH:.1f}조. 코멘트 원문 ‘18~24조’와 일치.",
            "원문 ‘삼성전자 민감도 → SK하이닉스 EPS +0.4%’는 표기 오타로 보고 삼성 0.4 / 하이닉스 0.9로 그렸습니다.",
        ],
        kind="blue",
    )
    n.bullet(f"1,300원대 중반은 조건부: DXY만 3~4% 하락 시 약 {d.KRW_FLOOR_DXY:,}원, 한국 달러 공급 가세 시 {d.KRW_FLOOR_SUPPLY:,}원대.")
    n.bullet(f"추가 하락의 핵심 변수는 달러. 인상 기대 되돌림 → DXY {d.DXY_NOW} → {d.DXY_DOWNSIDE[0]}~{d.DXY_DOWNSIDE[1]}.")
    n.bullet("리스크: 외국인 국내주식 매도, 미국-이란/유가, 1,350 부근 달러 수요 증가.")
    n.bullet("수출주 원화약세 효과는 3Q에 가팔라진 강세의 역풍으로 전환될 수 있음.")

    # 3
    n.h1("SK하이닉스 주주환원", num="3.")
    n.image("skh_return")
    n.table(
        ["항목", "숫자"],
        [
            ["취득·소각", f"{d.SKH_BUYBACK_KRW_T}조원  ·  전일 종가 {d.SKH_BUYBACK_PX:,}원 기준 약 {d.SKH_BUYBACK_SHARES_M:.2f}백만주"],
            ["발행주식 대비", f"{d.SKH_BUYBACK_PCT:.2f}%  ·  EPS +{d.SKH_EPS_UPLIFT:.2f}% (주식수 감소)"],
            ["기간", f"8/20~11/19  ·  {d.SKH_DAYS}영업일  ·  약 {d.SKH_DAILY_KRW_100M:,.0f}억원/일"],
            ["정책", "25~27 누적 FCF ‘50% 이상’ + 자사주·배당 병행. 특별배당 검토. 3Q26 구체화"],
            ["현금", f"2Q 순현금 약 {d.SKH_NET_CASH_2Q}조"],
            ["구조", "ADR 희석 → 소각 → SK스퀘어 지분율 복원"],
        ],
        col_widths=[3.6, 14.0],
    )
    n.image("skh_fcf_ladder")
    n.callout(
        "오해 금지",
        [
            f"{d.SKH_RETURN_MIN:g}조는 2027년 일시 지급이 아니라 3년 프로그램 누적 하한입니다.",
            "2028년 102.5조는 회사 정책이 아니라 내부 FCF 모델에 50%를 적용한 참고치입니다.",
            f"내부 래더 179/242/237 → 보수 {d.SKH_FCF_CONSERVATIVE[0]}/{d.SKH_FCF_CONSERVATIVE[1]}/{d.SKH_FCF_CONSERVATIVE[2]} (합 {d.SKH_FCF_CONSERVATIVE_SUM}조)는 25~27 FCF {d.SKH_FCF_2527}조와 다른 프레임입니다. 섞어 쓰지 마세요.",
        ],
        kind="bear",
    )
    n.bullet("키옥시아 8,000억엔 매입: 46,500→49,950 누적 +7.4%. 시황이 나쁘면 소폭이지만 매입 비중 이상은 상승.")
    n.bullet("샌디스크 $140억 한도 확대는 8/5~8/7 −15.1%와 겹침. +20%의 직접 원인이 아님. 8/13 Investor Day가 본반등.")

    # 4
    n.h1("밸류에이션 맵", num="4.")
    n.image("per_map")
    n.image("adr_premium")
    n.image("samsung_band")
    n.table(
        ["", "26Y OP / EPS", "27Y OP / EPS", "PER"],
        [
            ["SKH 컨센", f"{d.SKH_OP_26}조 / {d.SKH_EPS_26:,}", f"{d.SKH_OP_27}조 / {d.SKH_EPS_27:,}", f"{d.SKH_PER_26} / {d.SKH_PER_27}"],
            ["SKH 보수", f"{d.SKH_OP_26_BEAR[0]}~{d.SKH_OP_26_BEAR[1]}조", f"EPS {d.SKH_EPS_26_BEAR[0]//1000}~{d.SKH_EPS_26_BEAR[1]//1000}K", f"27Y {d.SKH_PER_27_BEAR}"],
            ["삼성 컨센", f"{d.SEC_OP_26}조 / {d.SEC_EPS_26:,}", f"{d.SEC_OP_27}조 / {d.SEC_EPS_27:,}", f"{d.SEC_PER_26} / {d.SEC_PER_27}"],
            ["삼성 보수", f"{d.SEC_OP_26_BEAR[0]}~{d.SEC_OP_26_BEAR[1]}조", f"EPS {d.SEC_EPS_26_BEAR[0]//1000}~{d.SEC_EPS_26_BEAR[1]//1000}K", f"27Y {d.SEC_PER_27_BEAR}"],
            ["MU", f"주가 ${d.MU_PX:.2f}", f"CY27 EPS ${d.MU_CY27_EPS}", f"F12M {d.MU_F12_PER} / CY27 {d.MU_CY27_PER}"],
            ["SNDK", f"주가 ${d.SNDK_PX:.2f}", f"FY27 EPS ${d.SNDK_FY27_EPS}", f"{d.SNDK_FY27_PER}배"],
        ],
        col_widths=[3.2, 5.2, 5.4, 4.0],
    )
    n.callout(
        "ADR 52%는 과함",
        [
            f"고가 ${d.SKH_ADR_HIGH} × {d.SKH_ADR_FX_REF}원 ≈ 228만 환산. 종가 ${d.SKH_ADR_CLOSE}(+0.3%).",
            f"정상 프리미엄 +20%면 본주 약 {d.SKH_LOCAL_IF_PREM20/10000:.0f}만. 최근 30~35%면 {d.SKH_LOCAL_IF_PREM30/10000:.0f}~{d.SKH_LOCAL_IF_PREM35/10000:.0f}만.",
            f"26년 성장 0 + PER 6~7배: 하이닉스 {d.SKH_PER6/10000:.0f}~{d.SKH_PER7/10000:.0f}만, 삼성 {d.SEC_PER6/10000:.1f}~{d.SEC_PER7/10000:.1f}만.",
        ],
        kind="blue",
    )

    # 5
    n.h1("HBM 논쟁 — 가격결정력은 영구가 아니다", num="5.")
    n.image("hbm_net")
    n.callout(
        "두 문장의 타당성",
        [
            "낮음: HBM이 비싸졌으니 메모리 수요가 곧 꺾인다.",
            "높음: 가격이 높아질수록 AI 업계가 사용량을 줄이는 기술을 개발할 경제적 유인이 커진다.",
            "핵심 변수 = AI 연산 증가율 − 메모리 효율 개선률.",
        ],
        kind="key",
    )
    n.table(
        ["현상", "가시성", "의미"],
        [
            ["Cerebras 온칩 SRAM", "높음", "HBM 의존 축소. 대체라기보다 특정 workload"],
            ["Groq SRAM inference", "높음", "추론 특화"],
            ["NVIDIA + Groq", "높음", "GPU 진영도 추론용 SRAM 채택"],
            ["KV Cache 압축", "높음", "필요 HBM 용량 감소"],
            ["HBF / SSD 계층", "진행 중", "분업. SRAM+HBM+DRAM+SSD"],
        ],
        col_widths=[4.6, 2.8, 10.2],
    )
    n.bullet("벤 톰슨 호르무즈 비유: 가격을 과도하게 올리면 고객은 장기적으로 공급망·기술을 바꾼다.")
    n.bullet("26~28년은 초과이익, 28년 이후는 효율화·대체를 촉진하는 양날. ‘지금 틀렸다’가 아니라 ‘영구 가정 금지’.")
    n.bullet("이 시각이 헤게모니(메모리 vs 비메모리)로 퍼지면 수급이 버거워집니다.")

    # 6
    n.h1("Google × Marvell", num="6.")
    n.image("mrvl_avgo")
    n.table(
        ["항목", "내용"],
        [
            ["협력", "7/29 커스텀 반도체 확대. AI 추론 + Storage + NIC + Memory Interface + Near-memory"],
            ["워런트", f"최대 {d.MRVL_WARRANT_M}만주  ·  행사가 ${d.MRVL_STRIKE}"],
            ["Vesting", f"Google Custom Products 매출 ${d.MRVL_TRANCHE_USD_M}M마다 1 tranche"],
            ["기간", d.MRVL_WINDOW],
            ["해석", "장기 ASIC 계약. 브로드컴 TPU 독점 견제. 마벨 +7~10% / 브로드컴 −4%대"],
        ],
        col_widths=[3.2, 14.4],
    )

    # 7
    n.h1("삼성 파운드리 · NVIDIA / OpenAI", num="7.")
    n.table(
        ["공정", "인상", "배경"],
        [
            ["4nm SF4 중·미", "10~15%", "TSMC 첨단 포화, 중국 팹리스 해외 의존"],
            ["4nm SF4 대만", "5~10%", "고객 지역별 차등"],
            ["5nm SF5", "10~15%", "웨이퍼 기준"],
            ["8nm", "약 10%", "레거시"],
        ],
        col_widths=[4.0, 3.2, 10.4],
    )
    n.bullet("평택 SF4: 퀄컴 + HBM 베이스다이 풀가동. 내년 흑자 전환 기대. 첨단 50%+, AI/HPC 30%+.")
    n.bullet("고객: 테슬라·애플·브로드컴, 엔비디아 추론칩, 구글 4nm 협의.")
    n.image("nvda_openai")
    n.flow(["NVIDIA 금융/지분", "AI 기업·DC", "GPU 구매", "엔비디아 매출"])
    n.bullet("Q3 YoY +90% → $108B, Q4 +77%여도 $120B. 절대액은 계속 증가.")
    n.bullet("Rubin 3Q26, 추론 35배, AI Factory 10배, 랙 $7~8.5M. Top5 2027 CAPEX ≥ $1T(+33%).")
    n.bullet("OpenAI Q2 $6.7B(+18%), 손실 $9.3B→$12.3B. 순환금융 논쟁 → SOX 단기 우려.")

    # 8
    n.h1("소부장 — 이수페타시스 · 기가비스", num="8.")
    n.image("isu_mix")
    n.image("giga_leverage")
    n.table(
        ["종목", "핵심", "숫자"],
        [
            ["이수페타시스", "Capa가 아니라 Multi-Lam 이익 레버리지", f"2Q 매출 {d.ISU_REV:,}억(+{d.ISU_REV_BEAT}% beat) · OP {d.ISU_OP}억 · ML {d.ISU_ML['1Q']}→{d.ISU_ML['2Q']}% · 판가 +{d.ISU_ASP}%"],
            ["기가비스", "FC-BGA의 눈(AOI)+수리공(AOR)", f"일본 수주 {d.GIGA_CONTRACT}억({d.GIGA_SALES_PCT}%) · 25 {d.GIGA_25['rev']}/{d.GIGA_25['op']} → 26E {d.GIGA_26E['rev']}/{d.GIGA_26E['op']}"],
        ],
        col_widths=[3.4, 5.8, 8.4],
    )
    n.bullet("이수 Capa 월 1,200→1,500(27.2Q)→1,800억(28H2). 27년 OP +10% 여지. 멀티플 30.9→26.0배.")
    n.bullet(f"기가비스 컨센 TP {d.GIGA_TP//10000}만. 26년 실적 기준 비싸, 27년이 초점. 분할 접근 후보.")

    # 9
    n.h1("로테이션 · 기타", num="9.")
    n.bullet("장기금리 하락 + 기술주 약세 = 헬스케어 피벗. 모더나 암백신 3상, NBI 신고가. 한국 바이오로 바로 이식은 조심.")
    n.bullet("다음 확인: 8/26 엔비디아 실적 — AI 심리 재점화 vs 로테이션 연장.")
    n.bullet("유니트리 PSR 155배(60배 프레임의 2.6배). 상업화 초기, 1Q 순익 −47.7%.")
    n.bullet("앤트로픽 매출 최대 $1,200억 / EV $2조 관측. LG엔솔 ESS, 한화 MTC, LS 사상최대.")
    n.bullet("국내기관 고민은 AI 이탈이 아니라 대형주 익스포저 vs 변압기/소부장 믹스.")

    n.h1("클로징", num="10.")
    n.callout(
        "가져갈 세 문장",
        [
            "1) 오늘은 AI 수요 파괴가 아니라 할인율 + 차익실현 + 원화 강세가 겹친 날입니다.",
            f"2) 하이닉스 40조는 시작이고, 환율은 27년 순익에서 약 {d.SKH_NI_ADJ_LOW:.0f}~{d.SKH_NI_ADJ_HIGH:.0f}조를 움직일 수 있습니다.",
            "3) HBM 가격결정력을 영구로 두지 말 것. SRAM이 HBM을 지운다고 단정하지 말 것. 연산−효율과 2028년 이후를 보면 됩니다.",
        ],
        kind="key",
    )
    n.p("위 자료는 공개 퀵코멘트를 도표로 재구성한 참고용입니다. 매수·매도 추천이 아니며 투자 판단은 각 독자의 몫입니다.", size=9.5, color=GRAY)

    n.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    p = build()
    print(f"Wrote {p} ({p.stat().st_size} bytes)")
