#!/usr/bin/env python3
"""8월 21–22일 패널 5쟁점 p<0.05 검증 강의노트(.docx) 생성."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Mm, Pt, RGBColor

OUT_PATH = Path("/workspace/lectures/8월 21-22일 패널 5쟁점 p값 검증.docx")

KR_FONT = "맑은 고딕"
EN_FONT = "Calibri"

NAVY = RGBColor(0x0F, 0x20, 0x43)
NAVY2 = RGBColor(0x1E, 0x40, 0x7C)
GOLD = RGBColor(0xB8, 0x94, 0x3A)
GRAY = RGBColor(0x4B, 0x55, 0x63)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x16, 0x65, 0x34)
RED = RGBColor(0x99, 0x1B, 0x1B)
AMBER = RGBColor(0x7A, 0x5C, 0x12)

NAVY_HEX = "0F2043"
NAVY2_HEX = "1E407C"
GOLD_HEX = "B8943A"
LIGHT_HEX = "EEF2F8"
GREEN_HEX = "E8F5E9"
RED_HEX = "FDECEA"
AMBER_HEX = "FFF8E7"
BLUE_HEX = "E8F1FB"
ROW_HEX = "F7F9FC"
WHITE_HEX = "FFFFFF"


def set_run_font(run, size=11, bold=False, color=DARK, italic=False, font=KR_FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), KR_FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(
        parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill}"/>')
    )


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(
        parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'<w:top w:w="{top}" w:type="dxa"/>'
            f'<w:left w:w="{left}" w:type="dxa"/>'
            f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
            f'<w:right w:w="{right}" w:type="dxa"/>'
            f"</w:tcMar>"
        )
    )


def set_table_borders(table, color="D0D7E2", sz="4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f"</w:tblBorders>"
    )
    tbl_pr.append(borders)


def set_left_accent(cell, color=NAVY_HEX, sz="24"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(
        parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="nil"/>'
            f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="nil"/>'
            f'<w:right w:val="nil"/>'
            f"</w:tcBorders>"
        )
    )


def prevent_row_split(row):
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()
    tr_pr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))


def cell_text(cell, text, size=10, bold=False, color=DARK, align="left", font=KR_FONT):
    cell.text = ""
    align_enum = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    lines = str(text).split("\n")
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = align_enum
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1 if i < len(lines) - 1 else 0)
        p.paragraph_format.line_spacing = 1.12
        run = p.add_run(line)
        set_run_font(run, size=size, bold=bold, color=color, font=font)
    set_cell_margins(cell)


def add_runs(paragraph, parts, size=11, color=DARK):
    for part in parts:
        if isinstance(part, str):
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)
        else:
            text, bold, *rest = part
            c = rest[0] if rest else color
            run = paragraph.add_run(text)
            set_run_font(run, size=size, bold=bold, color=c)


class Notes:
    def __init__(self):
        self.doc = Document()
        self._setup()

    def _setup(self):
        sec = self.doc.sections[0]
        sec.page_width = Mm(210)
        sec.page_height = Mm(297)
        sec.left_margin = Mm(16)
        sec.right_margin = Mm(16)
        sec.top_margin = Mm(16)
        sec.bottom_margin = Mm(16)
        sec.header_distance = Mm(8)
        sec.footer_distance = Mm(8)

        normal = self.doc.styles["Normal"]
        normal.font.name = KR_FONT
        normal.font.size = Pt(11)
        normal.font.color.rgb = DARK
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), KR_FONT)
        pf = normal.paragraph_format
        pf.space_after = Pt(6)
        pf.space_before = Pt(0)
        pf.line_spacing = 1.18

        header = sec.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = hp.add_run("8/21–22 패널 5쟁점  ·  p<0.05 검증  ·  강의노트")
        set_run_font(r, size=8.5, color=GRAY)

        footer = sec.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run("대화록 재구성  ·  통계는 Yahoo/FRED 재현  ·  통과 항목 위주  ·  ")
        set_run_font(r, size=8, color=GRAY)
        fld = parse_xml(
            f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE ">'
            f'<w:r><w:rPr><w:sz w:val="16"/><w:color w:val="4B5563"/>'
            f'<w:rFonts w:ascii="{KR_FONT}" w:hAnsi="{KR_FONT}" w:eastAsia="{KR_FONT}"/></w:rPr>'
            f"<w:t></w:t></w:r></w:fldSimple>"
        )
        fp._p.append(fld)

        core = self.doc.core_properties
        core.title = "8월 21-22일 패널 5쟁점 p값 검증"
        core.author = "준혁"
        core.subject = "삼성·하이닉스 환원, 카카오 분할, 매크로, AI 토큰, 홀수해·할로윈"

    def p(self, text, size=11, bold=False, color=DARK, space_after=6, space_before=0, align="left"):
        para = self.doc.add_paragraph()
        para.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }[align]
        para.paragraph_format.space_after = Pt(space_after)
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.line_spacing = 1.18
        run = para.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
        return para

    def rich(self, parts, size=11, space_after=6, space_before=0):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_after = Pt(space_after)
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.line_spacing = 1.18
        add_runs(para, parts, size=size)
        return para

    def h1(self, text, num=None):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing = 1.1
        if num:
            run = para.add_run(f"{num}  ")
            set_run_font(run, size=16, bold=True, color=GOLD)
        run = para.add_run(text)
        set_run_font(run, size=16, bold=True, color=NAVY)
        pPr = para._p.get_or_add_pPr()
        pPr.append(
            parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="12" w:space="4" w:color="{NAVY_HEX}"/>'
                f"</w:pBdr>"
            )
        )
        return para

    def h2(self, text):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(text)
        set_run_font(run, size=13, bold=True, color=NAVY2)
        return para

    def h3(self, text):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(3)
        run = para.add_run(text)
        set_run_font(run, size=11.5, bold=True, color=NAVY)
        return para

    def bullet(self, text, level=0, bold_lead=None, size=11):
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.55 + level * 0.45)
        para.paragraph_format.first_line_indent = Cm(-0.35)
        para.paragraph_format.space_after = Pt(2.5)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.line_spacing = 1.15
        mark = "• " if level == 0 else "– "
        run = para.add_run(mark)
        set_run_font(run, size=size, color=NAVY2 if level == 0 else GRAY)
        if bold_lead:
            run = para.add_run(bold_lead)
            set_run_font(run, size=size, bold=True, color=DARK)
            run = para.add_run(text)
            set_run_font(run, size=size, color=DARK)
        else:
            run = para.add_run(text)
            set_run_font(run, size=size, color=DARK)
        return para

    def flow(self, items, size=11):
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing = 1.2
        for i, item in enumerate(items):
            if i:
                run = para.add_run("   →   ")
                set_run_font(run, size=size, bold=True, color=GOLD)
            run = para.add_run(item)
            set_run_font(run, size=size, bold=True, color=NAVY)
        return para

    def callout(self, title, body, kind="key"):
        palette = {
            "key": (NAVY_HEX, LIGHT_HEX, NAVY),
            "bull": ("166534", GREEN_HEX, GREEN),
            "bear": ("991B1B", RED_HEX, RED),
            "note": (GOLD_HEX, AMBER_HEX, AMBER),
            "blue": (NAVY2_HEX, BLUE_HEX, NAVY2),
        }
        accent, fill, title_color = palette[kind]
        table = self.doc.add_table(rows=1, cols=1)
        table.autofit = True
        table.allow_autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        shade_cell(cell, fill)
        set_left_accent(cell, accent, sz="28")
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        p1.paragraph_format.space_before = Pt(0)
        r = p1.add_run(title)
        set_run_font(r, size=10, bold=True, color=title_color)
        if isinstance(body, str):
            body = [body]
        for line in body:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(line)
            set_run_font(r, size=10.5, color=DARK)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def table(self, headers, rows, col_widths=None, first_col_bold=True):
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        set_table_borders(table, color="D5DCE6", sz="4")
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            shade_cell(cell, NAVY_HEX)
            cell_text(cell, h, size=9.5, bold=True, color=WHITE, align="center")
        prevent_row_split(table.rows[0])
        for r_i, row in enumerate(rows):
            for c_i, val in enumerate(row):
                cell = table.rows[r_i + 1].cells[c_i]
                if r_i % 2 == 1:
                    shade_cell(cell, ROW_HEX)
                else:
                    shade_cell(cell, WHITE_HEX)
                align = "left" if c_i == 0 else "center"
                bold = first_col_bold and c_i == 0
                cell_text(cell, str(val), size=9.5, bold=bold, color=DARK, align=align)
            prevent_row_split(table.rows[r_i + 1])
        if col_widths:
            for row in table.rows:
                for i, w in enumerate(col_widths):
                    row.cells[i].width = Cm(w)
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(8)
        spacer.paragraph_format.space_before = Pt(2)
        return table

    def spacer(self, pt=4):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(pt)
        p.paragraph_format.space_before = Pt(0)

    def save(self, path: Path):
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))


def build():
    n = Notes()

    n.p("2026. 8. 22. 검증노트  ·  대화록 5분야 × p<0.05 재현", size=10.5, color=GRAY, align="center", space_after=4)
    n.p("PASSED FIRST", size=13, bold=True, color=GOLD, align="center", space_after=2)
    n.p("패널 5쟁점 통계 검증", size=22, bold=True, color=NAVY, align="center", space_after=2)
    n.p("삼성·하이닉스 환원  ·  카카오 분할  ·  매크로  ·  토큰  ·  홀수해/할로윈", size=13, bold=True, color=NAVY2, align="center", space_after=8)
    n.p("박세익 · 조 부장 · 김장열 논리를 공시·시세·시계열로 다시 검산  |  기준일 2026-08-22", size=11, color=GRAY, align="center", space_after=10)

    n.callout(
        "한 장으로 보면 — 강의에서 숫자로 말해도 되는 것",
        [
            "p<0.05로 남는 축은 세 개입니다. ① 한국 증시 홀수해>짝수해  ② 한국 증시 겨울(11–4월)>여름(5–10월)  ③ 주가 YoY가 수출 YoY를 3~4개월 선행.",
            "삼전/하닉/카카오는 가설검정의 대상이 아닙니다. 금산법·소각 산식·분할비율은 구조 팩트입니다. 위스퍼 넘버와 일평균 1조 원 매수는 기각했습니다.",
            "조건부 할로윈 ‘승률 100%’는 표본 5개라 양측 이항은 탈락(p=0.0625)합니다. 평균 수익>0 검정만 통과합니다. 2026년 공식은 아직 8월 종가가 없습니다.",
        ],
        kind="key",
    )

    n.h2("검증 설계")
    n.table(
        ["항목", "내용"],
        [
            ["데이터", "Yahoo Finance 종가(코스피·코스닥·S&P·나스닥·삼전·하닉·카카오)\nFRED VALEXPKRM052N 한국 상품수출(월, 반도체 단독 시계열 아님)"],
            ["수익률", "연간: 직전 12월 종가 → 당해 12월 종가\n할로윈: 10월 월말 → 익년 4월 월말 / 4월 월말 → 10월 월말"],
            ["검정", "Welch t, Mann–Whitney U, 정확 이항, Pearson 시차상관\n기본 양측 p<0.05 · 방향 가설은 단측도 병기 · BH-FDR 18개 가족"],
            ["코스닥 표본", "Yahoo ^KQ11이 2000-10부터라 1998–2000 연간수익률은 제외\n짝수 12년 · 홀수 13년 (패널의 ‘14번’보다 짧음)"],
            ["선행성 한계", "관세청 10일 반도체 수출 원시계열은 미확보\n전체 수출 YoY로 대체. 방향은 통과, 품목 특정은 보류"],
        ],
        col_widths=[3.4, 14.2],
    )

    n.h2("강의 구성")
    n.table(
        ["파트", "성격", "오늘 가져갈 것"],
        [
            ["0", "검정표", "통과 12 · 기각/표본부족 · 구조 팩트"],
            ["1", "구조", "삼전 배당 vs 하닉 소각, 금산법 10%"],
            ["2", "구조", "카카오AI 0.36 / X 0.64, 장중 −13.2%"],
            ["3", "혼합", "금리 팩트 + 병목 서술. 세컨티어 P/E는 미검증"],
            ["4", "팩트", "OpenAI 분기 정정, Anthropic 116억 달러"],
            ["5", "통계 핵심", "홀수해·겨울장·3~4개월 선행 — 통과분만"],
        ],
        col_widths=[2.0, 3.2, 12.4],
    )

    n.h2("오프닝 멘트 (녹화용)")
    n.p(
        "오늘은 대화록을 다시 읽지 않습니다. 숫자만 다시 셉니다. "
        "패널이 말한 다섯 덩어리 가운데 p값으로 남는 것과, 공시로만 남는 것과, 버려야 하는 것을 나눕니다. "
        "강의에서 자신 있게 쓸 문장은 통과한 것만입니다."
    )

    # ── 0. Scoreboard ──────────────────────────────────────
    n.h1("통과 점수판", num="0.")
    n.callout(
        "판정 규칙",
        [
            "PASS = 양측(또는 사전 방향 단측) p<0.05. BH = 18개 핵심 검정에 벤야미니–호흐베르크.",
            "구조 산식(소각 EPS, 시총 대비 환원)은 p값이 아니라 항등식입니다. 따로 ‘산식 확인’으로 표시합니다.",
        ],
        kind="blue",
    )

    n.h3("통계 PASS — 강의에서 숫자로 써도 되는 행")
    n.table(
        ["가설", "결과", "p", "BH"],
        [
            ["코스닥 홀수해 평균 > 짝수해 (Welch)", "+24.2% vs −12.3%\n차이 +36.5%p, t=3.62", "0.0014", "PASS"],
            ["코스피 홀수해 평균 > 짝수해 (Welch)", "+28.9% vs −2.0%\n차이 +31.0%p, t=2.95", "0.0067", "PASS"],
            ["코스피 홀수해 평균 > 0 (단측 t)", "n=14, 13승 1패(2011)", "0.0011", "PASS"],
            ["코스닥 홀수해 상승확률 ≠ 50%", "11/13 = 84.6%", "0.0225", "PASS"],
            ["코스닥 짝수해 상승확률 ≠ 50%", "2/12 = 16.7%\n상승: 2014, 2020", "0.0386", "BH 0.053\n원 p만"],
            ["코스피 겨울(11–4) > 여름(5–10)", "+11.6% vs −1.2%\n1998–2024, n=27", "0.0233", "PASS"],
            ["코스닥 겨울 > 여름", "+9.6% vs −6.2%", "0.0039", "PASS"],
            ["조건부 할로윈 S&P 평균>0\n(8월>10월일 때만, 2015+)", "5/5, 평균 +15.3%", "0.0081\n단측 t", "PASS"],
            ["조건부 할로윈 나스닥 평균>0", "5/5, 평균 +17.7%", "0.0028\n단측 t", "PASS"],
            ["미 중간선거 집권당 하원 의석 손실", "1946–2022 18/20=90%\n평균 −25.6석", "0.00040\n/ 2.0e-5", "PASS"],
            ["하닉 주가 YoY가 한국 수출 YoY 선행", "최강 시차 +3개월\nr=0.631, n=231", "4.7e-27", "PASS"],
            ["코스피 YoY가 한국 수출 YoY 선행", "최강 시차 +4개월\nr=0.791, n=230", "1.2e-50", "PASS"],
        ],
        col_widths=[5.4, 5.6, 3.2, 3.4],
    )

    n.h3("통계 FAIL 또는 표본 부족 — 강의에서 확정처럼 말하지 말 것")
    n.table(
        ["패널 문장", "재현", "판정"],
        [
            ["코스닥 짝수해 평균 −17.8%, 평균<0", "평균 −12.3%, 단측 t p=0.059", "기각 (경계)"],
            ["코스닥 홀수해 평균 +39.7%, 14번 중 12번", "평균 +24.2%, 13번 중 11번", "방향만 통과\n점추정 불일치"],
            ["조건부 할로윈 S&P/나스닥 승률 100%", "5/5는 맞음. 양측 이항 p=0.0625", "승률 검정 기각\n평균 검정만 통과"],
            ["조건부 할로윈 코스피 88%(8중 7),\n실패=2020.3 코로나", "6/7=85.7%, 실패연도는 2021(−9.3%)\n2020 동 규칙은 +38.8%", "승률 검정 기각\n실패연도 오인"],
            ["S&P 겨울 vs 여름 (비조건부)", "1998–2024 차이 p=0.15", "기각"],
            ["약세장을 구조적(−50%) vs 이벤트(−35~40%)로 분류", "2000 −55.7%, 2008 −54.5%\n2020 −35.7%, 2026.7 −38.6%", "사례 기술만\nn=2+2라 유형 검정 불가"],
        ],
        col_widths=[6.2, 6.6, 4.8],
    )

    n.callout(
        "다중검정 메모",
        "핵심 가족 18개 중 원 p 통과 13, BH 통과 12. 코스닥 짝수해 승률(p=0.039)은 FDR에서 탈락합니다. 강의에서는 ‘원 p 통과, 다중검정에는 약하다’고 한 줄 붙이면 됩니다.",
        kind="note",
    )

    # ── 1. Samsung vs Hynix ────────────────────────────────
    n.h1("삼성전자 vs SK하이닉스 환원", num="1.")
    n.callout(
        "패널 논리",
        [
            "같은 ‘역대급 환원’이라도 집행 장치가 다릅니다. 삼전은 현금배당 집중, 하닉은 장내 매수 후 전량 소각.",
            "시장이 삼전을 판 이유는 규모 부족이 아니라 (i) 위스퍼와의 괴리 (ii) 소각 불가에 가까운 금산법 (iii) 발표 전 오버슛입니다.",
        ],
        kind="key",
    )

    n.h2("1) 공시 비교 — 산식 확인")
    n.table(
        ["항목", "삼성전자 (8/21 이사회)", "SK하이닉스 (8/19 이사회)"],
        [
            ["환원 규모", "2026년 재원 90~110조\n(확정 상한 아님, 실적·FCF 연동)", "1차 자사주 40조 434억\n추가 환원은 3Q 실적 때 안내"],
            ["FCF 규칙", "2024–26 누적 FCF 50%\n상방 개런티 없음 — 패널 맞음", "누적 FCF 50% 이내 → 50% 이상\n개런티로 상향 — 패널 맞음"],
            ["즉시 집행", "3Q 정규배당 포함 현금 ~30조\n10월 말 이사회에서 확정", "8/20–11/19 장내 매수 2,407만 주\n발행주식 7.3049억 주의 3.3%"],
            ["소각", "잔여 60~80조는 내년 1월에\n배당+매입·소각 조합으로 결정", "취득 종료 후 전량 소각\n‘매입 기간 3개월+소각’"],
            ["별도", "임직원 보상용 자사주 15조\n환원 재원과 분리, 소각 아님", "고정·특별배당 확대 검토"],
        ],
        col_widths=[3.2, 7.2, 7.2],
    )

    n.h3("주당 효과 — 패널 숫자 재계산")
    n.bullet("하닉 소각 3.3%의 영구 EPS 상향은 1/(1−0.033)−1 = +3.41%입니다. 패널의 ‘약 3%’는 산식 확인.", bold_lead="EPS: ")
    n.bullet("3Q 30조를 Yahoo 보통주 57.64억 주로 나누면 주당 약 5,205원. 우선주 포함 67.93억 주면 4,417원. 패널 DPS 5,570원은 회사 가이드가 아니라, 30조에 잔여 정규배당을 더한 추정치에 가깝습니다.", bold_lead="DPS: ")
    n.bullet("우선주 포함 시총 약 1,851조(8/21). 160조 환원을 시총으로 나누면 +8.6%. 패널의 ‘160조여도 +9%’ 항등식은 확인.", bold_lead="밸류: ")
    n.bullet("8/19 종가 247,500 → 8/21 종가 281,500 = +13.7%. ‘발표 직전 이틀 +13%’는 8/19 급락 종가 기준이면 맞습니다. +8.6% 여력 대비 오버슛 약 +5%p (패널 +4%p).", bold_lead="오버슛: ")

    n.h2("2) 금산법 제24조 — 삼전이 소각을 크게 못 여는 이유")
    n.p("금융산업의 구조개선에 관한 법률상 금융 계열사는 비금융 계열사 지분을 합산 10% 초과 보유할 수 없습니다. 6월 말 삼성생명 8.51% + 삼성화재 1.49% = 10.00%로 한도가 이미 찼습니다.")
    n.flow(["자사주 전량 소각", "발행주식 분모 축소", "금융 계열 지분율 자동 상승", "한도 초과분 시장 매도/블록딜"])
    n.bullet("2026년 3월 소각 때 이미 약 1.5조 원 블록딜로 지분율을 맞춘 전례가 있습니다. 60~100조 전량 소각은 오버행을 키우는 역설입니다.")
    n.bullet("하닉은 반대로 SK스퀘어가 공정거래법상 지분 하한을 신경 쓰므로 소각(분자 유지, 분모 축소)이 지배구조와 같은 방향입니다. 패널의 ‘왜 배당 vs 소각인가’ 구조 설명은 맞습니다.")

    n.h2("3) 수급 집행 — 여기를 고쳐 말해야 합니다")
    n.callout(
        "기각: 일평균 1조 원 · 65만 주",
        [
            "40조 ÷ 약 60거래일 ≈ 0.67조 원, 2,407만 주 ÷ 60 ≈ 40.1만 주. 65거래일이면 0.62조 · 37.0만 주.",
            "공시상 하루 주문 한도는 240.7만 주입니다. 한도를 매일 소진한다는 가정이 패널 숫자의 출처로 보이며, 금액·주수 모두 과대입니다.",
            "그래도 메커니즘은 유효합니다. 하닉은 3개월간 장내 매수가 매일 호가창에 들어가고, 삼전 30조는 배당락까지 시간차가 있습니다.",
        ],
        kind="bear",
    )

    n.h3("8/18–8/21 시세 (정규장)")
    n.table(
        ["날짜", "삼성전자", "SK하이닉스", "코스피"],
        [
            ["8/18", "268,500", "1,662,000", "6,870"],
            ["8/19 (하닉 −9.8%)", "247,500", "1,500,000", "6,471"],
            ["8/20", "271,000", "1,691,000", "6,853"],
            ["8/21 (삼전 공시일)", "281,500 (+3.87%)", "1,730,000", "6,913"],
        ],
        col_widths=[4.4, 4.4, 4.4, 4.4],
    )
    n.p("삼전은 정규장 +3.87% 뒤 넥스트레이드에서 한때 266,000(종가 대비 −5.5%)까지 밀렸습니다. 하닉 시간외 174~176만 원대는 거래소 정규 시계열로 재현하지 못했습니다. 방향(삼전 sell-on, 하닉 지지)만 패널과 일치합니다.")

    n.h2("4) 위스퍼 · 과세 · 김장열 오버슛")
    n.bullet("시장 일부는 FCF 263조 × 50% ≈ 130조, 언론 단독은 100조+, 위스퍼는 150~200조까지 올라갔습니다. 90~110조는 ‘사상 최대’이되 기대의 하단입니다.", bold_lead="괴리: ")
    n.bullet("고액 자산가 금융소득종합과세(최고 45%+)는 배당 확대 시 배당락 전 매도를 자극합니다. 우선주가 당일 +8.26%로 괴리율 30%→23%를 좁힌 것은 이 채널과 맞습니다.", bold_lead="과세: ")
    n.bullet("시총 1,851조 기준 160조 환원의 기계적 여력 +8.6% vs 8/19→21 +13.7%. 차익 실현이 나온 자리는 통계가 아니라 산술입니다.", bold_lead="오버슛: ")

    # ── 2. Kakao ───────────────────────────────────────────
    n.h1("카카오 인적분할", num="2.")
    n.callout(
        "패널 논리",
        "인적분할 공시 자체를 ‘또 한 번의 쪼개기 상장’으로 읽었습니다. 사측은 지주사 전환을 부인하지만, 시장은 톡/페이 락인 해체와 AI 멀티플을 먼저 깎았습니다.",
        kind="key",
    )

    n.h2("1) 분할 구조 — 공시 확인")
    n.table(
        ["법인", "비율", "사업", "환원 약속"],
        [
            ["카카오AI (신설)", "순자산 0.36", "톡 · AI · AI커머스 · 광고 · 맵", "별도 조정 FCF 20~35%"],
            ["카카오X (존속)", "순자산 0.64", "뱅크·페이·증권 · 모빌리티 · 콘텐츠 · 벤처스", "자회사 배당 30% + 투자차익 30%"],
        ],
        col_widths=[4.0, 3.2, 6.0, 4.4],
    )
    n.bullet("주총 2026-12-17, 분할기일 2027-01-01, AI 재상장·X 변경상장 1/27. 기존 주주는 양사 주식을 비율대로 받습니다(물적 분할이 아님).")
    n.bullet("두나무 매각 차익으로 3,000억 자사주 매입·소각을 같이 내놨습니다. 급락을 막지 못했습니다.")

    n.h2("2) 8/21 가격 — 두 자릿수는 장중")
    n.p("전일 38,700 → 종가 35,800(−7.5%) · 장중 저가 33,600(−13.2%). 패널의 ‘두 자릿수 급락’은 종가가 아니라 장중 저가 기준입니다. 3월 4일 이란전 장중 −16.3% 이후 첫 두 자릿수 장중 하락입니다.")

    n.h2("3) 시장이 판 세 가지 이유 — 논리 정리, 통계 대상 아님")
    n.bullet("페이·뱅크·게임즈 연속 상장의 학습. 인적분할 공시 = 지주사 디스카운트 전조. 사측 ‘X는 지주사 전환 안 함’과 시장의 NAV 60~70% 할인 우려가 충돌.", bold_lead="1) 쪼개기 불신: ")
    n.bullet("5,000만 톡 이용자의 메신저 안 결제·금융이 락인이었습니다. 플랫폼과 테크핀을 법인으로 자르면 내부 거래·데이터·브랜드 로열티를 계약으로 다시 붙여야 합니다.", bold_lead="2) 톡–페이 분리: ")
    n.bullet("선행 P/E ~20배 vs 네이버 15~16배. 하이퍼클로바X·자체 DC가 있는 네이버 대비 API 연동 중심이라는 평가. 멀티플 정당화 실패는 의견이지 검정 결과가 아닙니다.", bold_lead="3) AI 멀티플: ")

    # ── 3. Macro / supply ──────────────────────────────────
    n.h1("매크로와 반도체 밸류체인", num="3.")
    n.callout(
        "통과에 가까운 팩트 vs 서술",
        [
            "금리 레벨은 확인됩니다. 8/22 미 10년 4.74%, 30년 5.28%, 2026 고점 30년 5.31%. 패널의 4.7~5.0 / 5.2 근방은 맞습니다.",
            "빅테크 장기채가 국채 수요를 밀어낸다는 채널, GPU→CPU 비율, 님비 70%는 시계열 검정이 아닙니다. 병목 이동은 실적으로 우회 확인합니다.",
        ],
        kind="blue",
    )

    n.h2("1) 금리와 빅테크 채권 — 패널 메커니즘")
    n.p("하이퍼스케일러가 AI CapEx를 10·30년 회사채로 조달하면, 민간은 ‘국채+스프레드’를 국채 대신 삽니다. 장기 금리가 튀고 성장주 할인율이 올라갑니다. 명목 GDP가 장기 금리보다 높으면 즉시 시스템 위기는 아니라는 안도 논리 — 성장률 공식 검정은 이 노트에서 하지 않았습니다.")

    n.h2("2) 병목 이동 — 실적으로 확인된 부분")
    n.bullet("이비덴 FY27 1Q 매출 1,232억 엔(+26.4%), OP 269억 엔(+52.4%, 컨센서스 +28.9%). 연간 OP 가이던스 900억→1,270억 엔(+41%). 삼성전기가 이 서프라이즈에 4거래일 +54% 반등한 것은 가격 반응입니다.", bold_lead="FC-BGA: ")
    n.bullet("삼성전기는 8/1부터 MLCC 전 제품 +30%. 무라타·다이요유덴도 인상 대열. AI 랙당 고사양 MLCC 수요가 소비재를 구축합니다.", bold_lead="MLCC: ")
    n.bullet("구글 TPU + 브로드컴 메인이라는 구도에 마벨이 연결 칩으로 들어간다는 서술은 산업 코멘트입니다. 비중 데이터로 검정하지 않았습니다.", bold_lead="ASIC: ")
    n.bullet("에이전트 AI → 제어·오케스트레이션 CPU 비중 확대, GPU:CPU 8:1→1:1. 공개 시계열로 p값을 내지 못했습니다. 강의에서는 ‘방향 가설’로만 쓰십시오.", bold_lead="GPU→CPU: ")

    n.callout(
        "세컨티어 P/E 상대평가 — 미검증, 논리만 유지",
        "사이클 후반에 2·3차 벤더를 ‘대장 대비 싸다’고 사는 전략은 상위 멀티플 수축 + 하위 실적 레버리지가 겹치면 낙폭이 더 큽니다. 이 노트는 그 가설을 종목 유니버스로 검정하지 않았습니다. 경고로만 남깁니다.",
        kind="note",
    )

    # ── 4. AI tokens ───────────────────────────────────────
    n.h1("AI 모델과 토큰 이코노믹스", num="4.")
    n.callout(
        "정정하고 쓸 숫자",
        [
            "패널 ‘2분기 매출 57억 달러’는 1분기 숫자입니다. WSJ: OpenAI 1Q 57억 → 2Q 67억(+18%), 영업손실 93억 → 123억. 67+123=190억으로 ‘분기 비용 ~180억’ 스케일은 맞습니다.",
            "Anthropic 2Q 매출 116억 달러는 확인. 조정 영업이익 +5.59억 달러(SBC 제외). 패널의 ‘BEP 근접’보다 한 걸음 더 가 있습니다. 하반기 IPO는 계획 서술입니다.",
        ],
        kind="bear",
    )

    n.h2("패널이 그린 수익성 격차")
    n.table(
        ["회사", "2Q26 매출", "손익", "강의 포인트"],
        [
            ["OpenAI", "67억 달러 (1Q 57억)", "영업손실 123억", "캐시번 + 인력 이탈 프레임. 손실이 매출보다 빠름"],
            ["Anthropic", "116억 달러 (1Q 47억)", "조정 OP +5.59억", "B2B·코딩(Claude)으로 분기 매출 역전"],
        ],
        col_widths=[3.4, 5.0, 4.4, 4.8],
    )
    n.bullet("DeepSeek·Kimi 등 오픈웨이트가 범용 토큰 가격을 깎고, 구형 A100 추론 수요가 감가 부실을 완화한다는 논리 — 가격 지수로 미검증.")
    n.bullet("서부 공업용수·냉각탑 소음·주거 전기요금 → 반대 70%+, 중간선거 인허가. 님비 비율은 출처 미재현. 정치 캘린더는 5장의 중간선거 통계와 연결해 쓰십시오.")

    # ── 5. Quant ───────────────────────────────────────────
    n.h1("수급 왜곡과 퀀트 — 통과분만", num="5.")
    n.callout(
        "박세익·조 부장이 맞았던 뼈대",
        [
            "코스피 홀수해 평균 +28.9%는 1998–2025 재현에서 소수점까지 같습니다. 이 한 줄이 대화록에서 가장 깨끗한 통계입니다.",
            "코스닥 짝수해에 돈 벌기 어렵다는 방향도 통과합니다. 다만 평균 −17.8%·14번 중 2번은 표본·정의를 고쳐 말해야 합니다.",
        ],
        kind="bull",
    )

    n.h2("1) 홀수해 vs 짝수해")
    n.p("정의: 직전 12월 종가 대비 당해 12월 종가. 코스피 1998–2025(홀수 14, 짝수 14). 코스닥은 Yahoo가 2000-10부터라 2001–2025(홀수 13, 짝수 12).")

    n.table(
        ["지수", "홀수해", "짝수해", "차이 검정"],
        [
            ["코스피", "+28.9%  /  13승 1패\n유일한 음수: 2011 −11.0%", "−2.0%  /  7승 7패", "Welch p=0.0067\nMW p=0.018"],
            ["코스닥", "+24.2%  /  11승 2패\n음수: 2011, 2019", "−12.3%  /  2승 10패\n상승: 2014, 2020만", "Welch p=0.0014\nMW p=0.0018"],
            ["S&P 500 (참고)", "+14.8%  /  11/14", "+2.7%  /  9/14", "한국 효과만큼\n뚜렷하지 않음"],
        ],
        col_widths=[3.4, 5.0, 5.0, 4.2],
    )

    n.h3("패널 문장을 이렇게 고치십시오")
    n.bullet("유지: “홀수해 코스피 평균 +28.9%”. 재현 일치.", bold_lead="그대로: ")
    n.bullet("수정: “코스닥 홀수 +39.7%, 14중 12” → “2001–2025 홀수 +24.2%, 13중 11, p=0.022”.", bold_lead="고침: ")
    n.bullet("수정: “짝수 −17.8%, 14중 2(2014·2020)” → “평균 −12.3%(p=0.059로 0 하회는 실패), 승률 2/12는 p=0.039로 통과. 상승 연도 2014·2020은 동일”.", bold_lead="고침: ")

    n.h3("왜 짝수해가 약한가 — 중간선거 통계는 통과, 한국 연결은 절반")
    n.p("1946–2022 중간선거 20번 중 집권당 하원 의석 감소 18번(90%, 예외 1998·2002). 이항 p=0.00040, 평균 −25.6석 단측 p=2.0×10⁻⁵. ‘패배 확률 90%’는 의석 손실로 읽으면 정확합니다. 다수당 상실과는 다른 말입니다.")
    n.p("그 해가 한국 짝수해 약세의 원인이냐는 별 검정입니다. 코스닥 홀수 vs 중간선거 해 p=0.0018(통과), 홀수 vs 대선 해 p=0.055(경계). 코스피 홀수 vs 대선 해 p=0.022, 홀수 vs 중간선거 해 p=0.054. ‘짝수해=중간선거’로 단정하지 말고, 미국 정치 사이클이 한국 위험자산과 같이 움직인 흔적으로만 쓰십시오.")

    n.h2("2) 할로윈 — 조건부와 비조건부를 분리")
    n.p("패널 공식: 8월 종가 > 10월 종가인 해에만 11월→익년 4월 보유. 2015년 이후.")

    n.table(
        ["시장", "신호 연도", "승 / 평균", "무엇을 통과했나"],
        [
            ["S&P 500", "2016·18·20·22·23", "5/5  /  +15.3%", "평균>0 p=0.008. 승률 양측 p=0.0625 탈락"],
            ["나스닥", "동일 5개년", "5/5  /  +17.7%", "평균>0 p=0.0028. 승률 동일 탈락"],
            ["코스피", "7개년", "6/7  /  +10.7%", "승률 p=0.125, 평균 단측 p=0.054 탈락\n실패=2021 −9.3% (2020은 +38.8%)"],
        ],
        col_widths=[3.2, 4.6, 4.0, 5.8],
    )

    n.callout(
        "2020년을 실패 사례로 쓰면 안 됩니다",
        "패널은 ‘실패한 1번=2020년 3월 코로나’라고 했습니다. 그 문장은 5–10월에 파는 고전 할로윈의 2020년 오판입니다. 조건부(8→10월 하락 후 겨울 보유) 규칙에서 2020 코스피는 +38.8%로 가장 좋은 해입니다. 실제 실패는 2021년 겨울입니다.",
        kind="bear",
    )

    n.p("표본을 늘리면 다른 것이 통과합니다. 비조건부 겨울 vs 여름(1998–2024): 코스피 차이 +12.8%p p=0.023, 코스닥 +15.8%p p=0.0039. S&P는 같은 기간 p=0.15로 기각. ‘할로윈은 한국에서 더 진하다’가 데이터와 맞습니다.")
    n.p("2026년 적용: 8월이 아직 끝나지 않았습니다. 10월 종가가 8월 종가보다 낮을 때만 공식이 켜집니다. 지금 자리에서 미리 100%를 말하면 검정 밖입니다.")

    n.h2("3) 주가 3~6개월 선행 — 전체 수출로 재현, 통과")
    n.p("관세청 10일 반도체 수출 원자료는 없습니다. 대체 시계열은 FRED 한국 상품수출(월). YoY 상관:")
    n.table(
        ["선행 시차 (주가가 수출을 이끄는 개월)", "하닉 vs 수출 r (p)", "코스피 vs 수출 r (p)"],
        [
            ["0 (동행)", "0.522  (9.9e-18)", "0.626  (8.1e-27)"],
            ["+2", "0.621  (3.6e-26)", "0.770  (9.4e-47)"],
            ["+3", "0.631  (4.7e-27)  ★하닉 최강", "0.789  (2.9e-50)"],
            ["+4", "0.616  (2.1e-25)", "0.791  (1.2e-50)  ★코스피 최강"],
            ["+6", "0.515  (7.6e-17)", "0.693  (5.6e-34)"],
        ],
        col_widths=[6.2, 5.7, 5.7],
    )
    n.bullet("3~6개월 구간의 상관이 동행보다 높습니다. 패널 구간과 일치합니다.")
    n.bullet("한계: 반도체만이 아니라 한국 전체 수출입니다. 2026년 반도체 비중이 수출의 40%대라 방향은 쓸 수 있고, 관세청 10일 지표 그대로의 재현은 아닙니다.")
    n.bullet("8월 1–20일 반도체 수출 260억 달러(+198.8%, 사상 최대)는 아직 피크아웃이 아닙니다. ‘피크 전 조정을 분할 매수’는 이 선행 구조 위에서만 논리입니다.")

    n.h2("4) 약세장 유형 — 기술, 유형 검정은 불가")
    n.table(
        ["에피소드", "고점 → 저점", "낙폭", "패널 유형"],
        [
            ["2000 닷컴", "2000-01-04 1,059 → 2001-09-17 469", "−55.7%", "구조적 (≥50%) 부합"],
            ["2008 GFC", "2007-10-31 2,065 → 2008-10-24 939", "−54.5%", "구조적 부합"],
            ["2020 코로나", "2020-01-22 2,267 → 2020-03-19 1,458", "−35.7%", "이벤트 (−35~40%) 부합"],
            ["2026 7월", "2026-06-22 9,115 → 2026-07-30 5,594", "−38.6%", "이벤트 부합"],
        ],
        col_widths=[3.2, 7.4, 2.8, 4.2],
    )
    n.p("숫자는 맞습니다. ‘유형이 반복된다’는 n=2+2로는 p값을 줄 수 없습니다. 1987은 코스피 일별이 Yahoo에 없습니다.")

    n.h2("5) 레버리지 ETF · 호가창 — 팩트 확인, p값 없음")
    n.bullet("신한투자 집계: 삼전·하닉 롱 레버리지 14종 AUM 6/25 17.38조 → 8/10 5.44조(−69%). 패널의 ‘18조→5~6조, 3~4조가 정상’ 중 출발·현재는 확인, 3~4조는 의견.", bold_lead="AUM: ")
    n.bullet("VKOSPI 6/29 장중 97.99 → 8/13 종가 55.28. ‘50대, 정상 30대’는 레벨 서술로 타당합니다.", bold_lead="VKOSPI: ")
    n.bullet("코스피 7,500~8,500 개인 대기 매물 55조(상승 30·하락 27)는 증권사 집계 원자료를 못 봤습니다. 미검증.", bold_lead="매물대: ")

    # ── Close ──────────────────────────────────────────────
    n.h1("클로징 — 강의에서 쓸 다섯 문장", num="6.")
    n.callout(
        "통과한 것만 문장으로",
        [
            "1) 삼전은 금산법 때문에 소각 대신 30조 배당을 열었고, 8/19→21 +13.7%가 +8.6% 산술 여력을 넘긴 자리에서 시간외가 되돌렸습니다.",
            "2) 하닉 40조·3.3% 소각의 EPS 효과는 +3.41%입니다. 일평균 1조·65만 주는 틀렸고, 0.6조대·40만 주가 맞습니다.",
            "3) 카카오 분할비율 0.36/0.64는 공시 그대로입니다. 두 자릿수 하락은 장중 −13.2%입니다.",
            "4) 코스피 홀수해 +28.9%(p=0.001), 홀수>짝수(p=0.007). 코스닥 홀수>짝수(p=0.001). 짝수해 평균이 0보다 작다고는 말 못 합니다(p=0.059).",
            "5) 주가 YoY는 수출 YoY를 3~4개월 선행합니다(r=0.63~0.79). 8–10월 하락이 확인되면 겨울 보유는 한국에서 더 강합니다. 2015년 이후 100%는 n=5입니다.",
        ],
        kind="key",
    )

    n.h2("클로징 멘트 (녹화용)")
    n.p(
        "대화록은 다섯 개 분야로 넓습니다. 통계가 받아 주는 부분은 좁습니다. "
        "환원과 분할은 구조로 설명하고, 홀수해·겨울장·수출 선행만 숫자로 밀면 됩니다. "
        "100%와 −17.8%와 일평균 1조는 버려도 강의가 약해지지 않습니다."
    )

    n.spacer(8)
    n.p(
        "— 8월 21–22일 패널 검증노트. 대화록을 5분야로 재구성한 뒤 Yahoo/FRED로 p<0.05 재현. 스크립트: scripts/verify_panel_claims.py",
        size=9.5,
        color=GRAY,
        align="right",
    )

    n.save(OUT_PATH)
    print(f"Wrote {OUT_PATH} ({OUT_PATH.stat().st_size} bytes)")


if __name__ == "__main__":
    build()
