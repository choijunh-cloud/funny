#!/usr/bin/env python3
"""8/18–20 통합 시각화 보고서 워드."""

from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Mm, Pt, RGBColor

import aug19_data as a
import integrated_data as d
from charts_integrated import render_all

OUT_PATH = Path("/workspace/lectures/8월 18-20일 통합 시장 시각화 보고서.docx")
CHART_DIR = Path("/workspace/reports/charts_integrated")

KR = "맑은 고딕"
NAVY = RGBColor(0x0F, 0x20, 0x43)
NAVY2 = RGBColor(0x1E, 0x40, 0x7C)
GOLD = RGBColor(0xB8, 0x94, 0x3A)
GRAY = RGBColor(0x4B, 0x55, 0x63)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x16, 0x65, 0x34)
RED = RGBColor(0x99, 0x1B, 0x1B)
AMBER = RGBColor(0x7A, 0x5C, 0x12)
NAVY_H, LIGHT, GREEN_H, RED_H, AMBER_H, BLUE_H, ROW, WH = (
    "0F2043", "EEF2F8", "E8F5E9", "FDECEA", "FFF8E7", "E8F1FB", "F7F9FC", "FFFFFF",
)


def font(run, size=11, bold=False, color=DARK):
    run.font.name = KR
    run._element.rPr.rFonts.set(qn("w:eastAsia"), KR)
    run._element.rPr.rFonts.set(qn("w:ascii"), KR)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), KR)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def shade(cell, fill):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill}"/>')
    )


def margins(cell, t=60, b=60, l=80, r=80):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'<w:top w:w="{t}" w:type="dxa"/><w:left w:w="{l}" w:type="dxa"/>'
            f'<w:bottom w:w="{b}" w:type="dxa"/><w:right w:w="{r}" w:type="dxa"/></w:tcMar>'
        )
    )


def borders(table):
    table._tbl.tblPr.append(
        parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="D0D7E2"/>'
            f'<w:left w:val="single" w:sz="4" w:space="0" w:color="D0D7E2"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="D0D7E2"/>'
            f'<w:right w:val="single" w:sz="4" w:space="0" w:color="D0D7E2"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="D0D7E2"/>'
            f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="D0D7E2"/></w:tblBorders>'
        )
    )


def accent(cell, color=NAVY_H):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="nil"/><w:left w:val="single" w:sz="28" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="nil"/><w:right w:val="nil"/></w:tcBorders>'
        )
    )


def cell_text(cell, text, size=9, bold=False, color=DARK, align="left"):
    cell.text = ""
    al = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER}[align]
    for i, line in enumerate(str(text).split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = al
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        font(r, size=size, bold=bold, color=color)
    margins(cell)


class R:
    def __init__(self, charts):
        self.doc = Document()
        self.charts = charts
        sec = self.doc.sections[0]
        sec.page_width, sec.page_height = Mm(210), Mm(297)
        sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Mm(14)
        hp = sec.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = hp.add_run("8/18–20 통합 시각화 보고서  ·  11개 워드 재구성")
        font(r, 8.5, color=GRAY)
        fp = sec.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run("공개 자료 기준  ·  매수·매도 추천 아님  ·  ")
        font(r, 8, color=GRAY)
        fp._p.append(
            parse_xml(
                f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE ">'
                f'<w:r><w:rPr><w:sz w:val="16"/><w:color w:val="4B5563"/>'
                f'<w:rFonts w:ascii="{KR}" w:eastAsia="{KR}"/></w:rPr>'
                f"<w:t></w:t></w:r></w:fldSimple>"
            )
        )
        self.doc.core_properties.title = "8월 18-20일 통합 시장 시각화 보고서"
        self.doc.core_properties.author = "준혁"

    def p(self, t, size=11, bold=False, color=DARK, after=6, align="left"):
        para = self.doc.add_paragraph()
        para.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER}[align]
        para.paragraph_format.space_after = Pt(after)
        r = para.add_run(t)
        font(r, size, bold, color)

    def h1(self, t, num=None):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(11)
        para.paragraph_format.space_after = Pt(5)
        if num:
            r = para.add_run(f"{num}  ")
            font(r, 14.5, True, GOLD)
        r = para.add_run(t)
        font(r, 14.5, True, NAVY)
        para._p.get_or_add_pPr().append(
            parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="12" w:space="3" w:color="{NAVY_H}"/></w:pBdr>'
            )
        )

    def h2(self, t):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(7)
        para.paragraph_format.space_after = Pt(3)
        r = para.add_run(t)
        font(r, 12, True, NAVY2)

    def bullet(self, t):
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.48)
        para.paragraph_format.first_line_indent = Cm(-0.3)
        para.paragraph_format.space_after = Pt(2)
        r = para.add_run("• ")
        font(r, 10.5, color=NAVY2)
        r = para.add_run(t)
        font(r, 10.5)

    def flow(self, items):
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(7)
        for i, item in enumerate(items):
            if i:
                r = para.add_run("  →  ")
                font(r, 10, True, GOLD)
            r = para.add_run(item)
            font(r, 10, True, NAVY)

    def call(self, title, body, kind="key"):
        pal = {
            "key": (NAVY_H, LIGHT, NAVY),
            "bull": ("166534", GREEN_H, GREEN),
            "bear": ("991B1B", RED_H, RED),
            "note": (GOLD, AMBER_H, AMBER) if False else ("B8943A", AMBER_H, AMBER),
            "blue": ("1E407C", BLUE_H, NAVY2),
        }
        ac, fill, tc = pal[kind]
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        shade(cell, fill)
        accent(cell, ac)
        margins(cell, 70, 70, 110, 110)
        cell.text = ""
        r = cell.paragraphs[0].add_run(title)
        font(r, 10, True, tc)
        if isinstance(body, str):
            body = [body]
        for line in body:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(line)
            font(r, 10.5)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(5)

    def table(self, headers, rows, widths=None):
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        borders(table)
        for i, h in enumerate(headers):
            c = table.rows[0].cells[i]
            shade(c, NAVY_H)
            cell_text(c, h, 9, True, WHITE, "center")
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                c = table.rows[ri + 1].cells[ci]
                shade(c, ROW if ri % 2 else WH)
                cell_text(c, str(val), 9, ci == 0, DARK, "left" if ci == 0 else "center")
        if widths:
            for row in table.rows:
                for i, w in enumerate(widths):
                    row.cells[i].width = Cm(w)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(5)

    def img(self, key, w=16.4):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(7)
        p.add_run().add_picture(str(self.charts[key]), width=Cm(w))

    def save(self, path: Path):
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))


def build(chart_dir: Path = CHART_DIR) -> Path:
    charts = render_all(chart_dir)
    n = R(charts)
    n.p("2026. 8. 18–20  ·  업로드 워드 11개 통합", 10.5, color=GRAY, after=3, align="center")
    n.p("금리와 환율이 할인한 3일, 환원·파운드리·클라우드가 버틴 3일", 18, True, NAVY, 6, "center")
    n.p("이란 장기전 · 엔캐리 · 유가발 금리 · 원화 강세 · 하이닉스 환원 · HBM 논쟁 · 마벨–구글 · 삼성 파운드리 · NVIDIA · OpenAI · 알리바바 · 울프스피드 · 유니트리 · 실리콘투", 10, color=GRAY, after=8, align="center")

    n.call(
        "관통 문장",
        [
            "사흘의 공통분모는 AI 수요 파괴가 아닙니다. 할인율(금리·환율)이 가격을 깎고, 순환금융·효율화가 내러티브를 흔들고, 환원·파운드리·클라우드 CAPEX가 반대편에서 받칩니다.",
            f"하이닉스: 40조 소각 + 25~27 FCF 385조×50%=192.5조. 환율 1,520→1,420 시 EPS −{a.SKH_EPS_FX_HIT:.1f}%. 키옥시아+가격+환율 = 60조 중후반.",
            "8/19 금리는 내렸는데 SOX −2.12%. 알리바바 Cloud +40% vs EBITA −84%, 울프스피드 매출 부합·EPS 하회 — 수요는 있고 이익은 늦게 옵니다.",
        ],
    )

    n.table(
        ["날짜", "원본 워드", "한 줄"],
        [
            ["8/18", "이란 장기전·엔캐리·트럼프·실리콘투", "소모전 + JGB 30년래 최고 + CVC 3,000억"],
            ["8/19", "미국장 / 매크로 / NOBUY / 환원 / 밸류 / 환율", "유가발 금리 쇼크 + 40조 소각 + 파운드리 인상"],
            ["8/20", "미국장 금리↓ vs OpenAI / 알리바바·울프스피드 / NVDA", "바이백 진정, SOX는 따로, CAPEX가 이익을 삼킴"],
        ],
        [2.4, 7.4, 7.8],
    )

    n.h1("매크로 — 이란 · 유가 · 엔캐리 · 바이백", "1.")
    n.flow(["이란 장기전", "유가", "미 장기금리", "고PER 할인", "AI·반도체"])
    n.img("kospi_turnover")
    n.img("foreign")
    n.bullet("코스피 7000 회복 vs 8월 대금 25.7조(5~6월 50조의 반토막). 개인은 빠지고 외국인 5일 +9.5조. 추가 상승 = 외국인 지속.")
    n.img("jgb")
    n.img("crash_2024")
    n.call(
        "엔캐리: 조건 형성 ≠ 이미 2024.8",
        [
            f"JGB 2년 {d.JGB['2y']}% / 5년 {d.JGB['5y']}% / 10년 {d.JGB['10y_now']}%. BOJ 9월 80%.",
            f"위험 5개 동시: 10년>3% + USD/JPY 급락 + 미 30년>5.3~5.4% + SOX 급락 + 입찰 부진. 현재 엔 {d.USDJPY_NOW}면 1차는 금리/밸류.",
            "노무라 경로 26.9/27.1/27.4 각 25bp, 최종 1.75%. NISA 개인국채 3.8조엔이 구조적 카드.",
        ],
        "note",
    )
    n.img("sess819")
    n.img("macro_levels")
    n.bullet("이란: 1년 원유 비축, 12일 단독작전, 저항의 축. FT는 유럽 미군 시설 ‘검토’이지 결정이 아님.")
    n.bullet("바이백 $20→40억 = 방어선. 재정·인플레·AI 회사채는 남음. 2024.9 50bp 인하를 지금 대입하지 말 것.")
    n.bullet("트럼프 한미훈련 축소: 대북 유화보다 대한 불만(이란 미지지·투자 미이행) 해석. 산업부 ‘1호=반도체’ 부인.")

    n.h1("환율 — 국내 수급 + 60조 바구니", "2.")
    n.flow(["법인세·설비", "수출 달러매도", "헤지 ↑", "달러-원 ↓", "잔여매도", "가속"])
    n.img("fx_ladder")
    n.img("fx_sensitivity")
    n.img("skh_hit60")
    n.call(
        "민감도",
        [
            f"삼성 β 0.4 / 하이닉스 β 0.9. 1,520→1,420 → SKH EPS −{a.SKH_EPS_FX_HIT:.1f}%, 27년 NI 18~24조.",
            "2H26 환율 16.3조. 키옥시아 지분평가 + 가격/이익 + 환율 = 60조 중후반. 특별 일회성은 여지.",
            "1,300원대 중반 조건부. DXY 99→96~97, 1,350 달러수요, 외국인 매도, 유가.",
        ],
        "blue",
    )

    n.h1("메모리 — NOBUY · 밸류 · SCA", "3.")
    n.img("per_hdd")
    n.img("adr_premium")
    n.img("hbm_net")
    n.img("sca_asp")
    n.img("gp_bit")
    n.bullet("메모리 PER 6~8배 vs WDC/STX 26배. 같은 스토리지가 아닙니다.")
    n.bullet("ADR 52%는 과함. 정상 +20%면 본주 ~190만, 30~35%면 169~175만. PER 6~7배면 208~242만.")
    n.bullet("SCA 50%·시장 −20%·Floor 90% → GP/bit −16.6%. Bit +20%면 총GP 유지. HBM 프리미엄 제외.")
    n.call("NOBUY", ["비싸서 수요 꺾임 = 타당성 낮음. 비쌀수록 효율화 유인 = 높음. SRAM≠대체. 26~28 초과이익, 28+ 양날."], "key")

    n.h1("주주환원", "4.")
    n.img("fcf_years")
    n.img("skh_return")
    n.img("sndk_bb")
    n.bullet("385조 = 25+150+210 (2025A~2027E 보수). 192.5조는 3년 누적 하한. 2028 102.5조는 정책 아님.")
    n.bullet("샌디스크 $140억은 8/5~7 −15.1%와 겹침. 본반등은 8/13 ID(장기 GM ~80%). 키옥시아 +7.4%.")

    n.h1("비메모리 · NVIDIA · OpenAI", "5.")
    n.img("mrvl_avgo")
    n.img("nvda_q2")
    n.img("nvda_openai")
    n.bullet("삼성 파운드리 SF4 중·미 10~15%. 평택 풀가동, 내년 흑자, 구글 4nm 협의.")
    n.bullet(f"마벨 워런트 {a.MRVL_WARRANT_M}만주 ${a.MRVL_STRIKE}, $5억마다 vest, FY27Q3~FY33.")
    n.bullet(f"NVDA 컨센 ${d.NVDA_Q2['cons_rev']}B(+{d.NVDA_Q2['cons_yoy']}%), 저자 $93.5B. Beat보다 CAPEX·Rubin·순환금융 설명.")

    n.h1("알리바바 · 울프스피드", "6.")
    n.img("baba")
    n.img("baba_mult")
    n.img("wolf")
    n.bullet(f"BABA 매출 {d.BABA_REV:,}억위안 +{d.BABA_REV_YOY}%, Cloud {d.BABA_CLOUD}억 +{d.BABA_CLOUD_YOY}%. 직전 EBITA {d.BABA_EBITA_DROP}%. CAPEX 3,800억 초과 가능.")
    n.bullet("PEG 0.5, 고점 −33%. 이익 극대화가 아니라 AI 선투자.")
    n.bullet("WOLF 매출 $149.6M 부합, EPS −$2.26 vs −$1.47, GM −20%, AI DC +20% QoQ. SiC = EV + AI 전원.")

    n.h1("유니트리 · 실리콘투", "7.")
    n.img("unitree")
    n.img("silicon2")
    n.bullet(f"유니트리 ¥{d.UNITREE_PX}(+{d.UNITREE_IPO_CHG}%), 시총 {d.UNITREE_MCAP:,}억위안, PSR 155배(60배의 2.6배), PER ~{d.UNITREE_PER}배. 다음 관문은 공장 ROI.")
    n.bullet(f"실리콘투 CVC {d.SIL_CVC:,}억, 발행가 {d.SIL_ISSUE_PX:,}원(+{d.SIL_PREM}%), 희석 {d.SIL_DILUTE}%. 핵심은 Douglas. 글랜우드 오버행은 잔존.")

    n.h1("클로징", "8.")
    n.call(
        "가져갈 네 문장",
        [
            "1) 사흘은 AI 수요 파괴가 아니라 할인율 + 차익실현 + 원화 강세입니다.",
            "2) 하이닉스 40조는 시작이고, 환율·키옥시아·가격을 합치면 이익에서 60조 중후반이 움직입니다.",
            "3) HBM 가격결정력을 영구로 두지 말 것. SRAM이 HBM을 지운다고 단정하지 말 것.",
            "4) 알리바바 클라우드, 울프스피드 AI DC, OpenAI 손실은 같은 그림 — 수요는 있고 이익은 늦게 옵니다.",
        ],
    )
    n.p("위 자료는 업로드된 11개 워드를 도표로 재구성한 참고용입니다. 매수·매도 추천이 아닙니다.", 9.5, color=GRAY)
    n.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print("Wrote", build(), "bytes pending")
