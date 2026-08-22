#!/usr/bin/env python3
"""강의노트(.docx) 공통 서식."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Mm, Pt, RGBColor

KR_FONT = "맑은 고딕"
EN_FONT = "Calibri"

NAVY = RGBColor(0x0F, 0x20, 0x43)
NAVY2 = RGBColor(0x1E, 0x40, 0x7C)
GOLD = RGBColor(0xB8, 0x94, 0x3A)
GRAY = RGBColor(0x4B, 0x55, 0x63)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x16, 0x65, 0x34)
RED = RGBColor(0x99, 0x1B, 0x1B)
AMBER = RGBColor(0x7A, 0x5C, 0x12)

NAVY_HEX = "0F2043"
NAVY2_HEX = "1E407C"
GOLD_HEX = "B8943A"
LIGHT_HEX = "EEF2F8"
GREEN_HEX = "E8F5E9"
RED_HEX = "FDECEA"
AMBER_HEX = "FFF8E7"
BLUE_HEX = "E8F1FB"
ROW_HEX = "F7F9FC"
WHITE_HEX = "FFFFFF"


def set_run_font(run, size=11, bold=False, color=DARK, italic=False, font=KR_FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), KR_FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(
        parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill}"/>')
    )


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(
        parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'<w:top w:w="{top}" w:type="dxa"/>'
            f'<w:left w:w="{left}" w:type="dxa"/>'
            f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
            f'<w:right w:w="{right}" w:type="dxa"/>'
            f"</w:tcMar>"
        )
    )


def set_table_borders(table, color="D0D7E2", sz="4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f"</w:tblBorders>"
    )
    tbl_pr.append(borders)


def set_left_accent(cell, color=NAVY_HEX, sz="24"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(
        parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="nil"/>'
            f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="nil"/>'
            f'<w:right w:val="nil"/>'
            f"</w:tcBorders>"
        )
    )


def prevent_row_split(row):
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()
    tr_pr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))


def cell_text(cell, text, size=10, bold=False, color=DARK, align="left", font=KR_FONT):
    cell.text = ""
    align_enum = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    lines = str(text).split("\n")
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = align_enum
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1 if i < len(lines) - 1 else 0)
        p.paragraph_format.line_spacing = 1.12
        run = p.add_run(line)
        set_run_font(run, size=size, bold=bold, color=color, font=font)
    set_cell_margins(cell)


def add_runs(paragraph, parts, size=11, color=DARK):
    for part in parts:
        if isinstance(part, str):
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)
        else:
            text, bold, *rest = part
            c = rest[0] if rest else color
            run = paragraph.add_run(text)
            set_run_font(run, size=size, bold=bold, color=c)


class Notes:
    def __init__(self, header, footer, title, author="준혁", subject=""):
        self.header_text = header
        self.footer_text = footer
        self.title = title
        self.author = author
        self.subject = subject
        self.doc = Document()
        self._setup()

    def _setup(self):
        sec = self.doc.sections[0]
        sec.page_width = Mm(210)
        sec.page_height = Mm(297)
        sec.left_margin = Mm(16)
        sec.right_margin = Mm(16)
        sec.top_margin = Mm(16)
        sec.bottom_margin = Mm(16)
        sec.header_distance = Mm(8)
        sec.footer_distance = Mm(8)

        normal = self.doc.styles["Normal"]
        normal.font.name = KR_FONT
        normal.font.size = Pt(11)
        normal.font.color.rgb = DARK
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), KR_FONT)
        pf = normal.paragraph_format
        pf.space_after = Pt(6)
        pf.space_before = Pt(0)
        pf.line_spacing = 1.18

        header = sec.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = hp.add_run(self.header_text)
        set_run_font(r, size=8.5, color=GRAY)

        footer = sec.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run(self.footer_text)
        set_run_font(r, size=8, color=GRAY)
        fld = parse_xml(
            f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE ">'
            f'<w:r><w:rPr><w:sz w:val="16"/><w:color w:val="4B5563"/>'
            f'<w:rFonts w:ascii="{KR_FONT}" w:hAnsi="{KR_FONT}" w:eastAsia="{KR_FONT}"/></w:rPr>'
            f"<w:t></w:t></w:r></w:fldSimple>"
        )
        fp._p.append(fld)

        core = self.doc.core_properties
        core.title = self.title
        core.author = self.author
        core.subject = self.subject

    def p(self, text, size=11, bold=False, color=DARK, space_after=6, space_before=0, align="left"):
        para = self.doc.add_paragraph()
        para.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }[align]
        para.paragraph_format.space_after = Pt(space_after)
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.line_spacing = 1.18
        run = para.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
        return para

    def rich(self, parts, size=11, space_after=6, space_before=0):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_after = Pt(space_after)
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.line_spacing = 1.18
        add_runs(para, parts, size=size)
        return para

    def h1(self, text, num=None):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing = 1.1
        if num:
            run = para.add_run(f"{num}  ")
            set_run_font(run, size=16, bold=True, color=GOLD)
        run = para.add_run(text)
        set_run_font(run, size=16, bold=True, color=NAVY)
        pPr = para._p.get_or_add_pPr()
        pPr.append(
            parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="12" w:space="4" w:color="{NAVY_HEX}"/>'
                f"</w:pBdr>"
            )
        )
        return para

    def h2(self, text):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(text)
        set_run_font(run, size=13, bold=True, color=NAVY2)
        return para

    def h3(self, text):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(3)
        run = para.add_run(text)
        set_run_font(run, size=11.5, bold=True, color=NAVY)
        return para

    def bullet(self, text, level=0, bold_lead=None, size=11):
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.55 + level * 0.45)
        para.paragraph_format.first_line_indent = Cm(-0.35)
        para.paragraph_format.space_after = Pt(2.5)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.line_spacing = 1.15
        mark = "• " if level == 0 else "– "
        run = para.add_run(mark)
        set_run_font(run, size=size, color=NAVY2 if level == 0 else GRAY)
        if bold_lead:
            run = para.add_run(bold_lead)
            set_run_font(run, size=size, bold=True, color=DARK)
            run = para.add_run(text)
            set_run_font(run, size=size, color=DARK)
        else:
            run = para.add_run(text)
            set_run_font(run, size=size, color=DARK)
        return para

    def flow(self, items, size=11):
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing = 1.2
        for i, item in enumerate(items):
            if i:
                run = para.add_run("   →   ")
                set_run_font(run, size=size, bold=True, color=GOLD)
            run = para.add_run(item)
            set_run_font(run, size=size, bold=True, color=NAVY)
        return para

    def callout(self, title, body, kind="key"):
        palette = {
            "key": (NAVY_HEX, LIGHT_HEX, NAVY),
            "bull": ("166534", GREEN_HEX, GREEN),
            "bear": ("991B1B", RED_HEX, RED),
            "note": (GOLD_HEX, AMBER_HEX, AMBER),
            "blue": (NAVY2_HEX, BLUE_HEX, NAVY2),
        }
        accent, fill, title_color = palette[kind]
        table = self.doc.add_table(rows=1, cols=1)
        table.autofit = True
        table.allow_autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        shade_cell(cell, fill)
        set_left_accent(cell, accent, sz="28")
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        p1.paragraph_format.space_before = Pt(0)
        r = p1.add_run(title)
        set_run_font(r, size=10, bold=True, color=title_color)
        if isinstance(body, str):
            body = [body]
        for line in body:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(line)
            set_run_font(r, size=10.5, color=DARK)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def table(self, headers, rows, col_widths=None, first_col_bold=True):
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        set_table_borders(table, color="D5DCE6", sz="4")
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            shade_cell(cell, NAVY_HEX)
            cell_text(cell, h, size=9.5, bold=True, color=WHITE, align="center")
        prevent_row_split(table.rows[0])
        for r_i, row in enumerate(rows):
            for c_i, val in enumerate(row):
                cell = table.rows[r_i + 1].cells[c_i]
                if r_i % 2 == 1:
                    shade_cell(cell, ROW_HEX)
                else:
                    shade_cell(cell, WHITE_HEX)
                align = "left" if c_i == 0 else "center"
                bold = first_col_bold and c_i == 0
                cell_text(cell, str(val), size=9.5, bold=bold, color=DARK, align=align)
            prevent_row_split(table.rows[r_i + 1])
        if col_widths:
            for row in table.rows:
                for i, w in enumerate(col_widths):
                    row.cells[i].width = Cm(w)
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(8)
        spacer.paragraph_format.space_before = Pt(2)
        return table

    def spacer(self, pt=4):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(pt)
        p.paragraph_format.space_before = Pt(0)

    def save(self, path: Path):
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))
